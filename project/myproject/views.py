# myapp/views.py

from django.http import HttpResponse
from django.contrib.auth import get_user_model, authenticate
from django.db import connection, IntegrityError
from rest_framework.views import APIView
from rest_framework.response import Response
from .serializers import UserSerializer
from rest_framework_simplejwt.views import TokenObtainPairView as JWTTokenObtainPairView, TokenRefreshView as JWTTokenRefreshView
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework_simplejwt.tokens import AccessToken
from rest_framework_simplejwt.exceptions import TokenError
from django.conf import settings

User = get_user_model()
try:
    from rolepermissions.decorators import has_role
except Exception:
    # rolepermissions may not import cleanly in this environment; provide
    # a no-op fallback so management commands still work.
    def has_role(*args, **kwargs):
        def _decorator(func):
            return func
        return _decorator
from rest_framework import status
from django.db import connection
from django.utils import timezone
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.hashers import make_password
import tempfile
import subprocess
import os
import logging
import traceback
import sys
import zipfile
import re
import shutil
import binascii
import time
from myproject.models import DownloadLog
from myproject.rbac import get_role_from_request, RolePermission

# module logger
logger = logging.getLogger(__name__)


def _normalize_for_json(obj):
    """Recursively convert Decimal and non-serializable objects to serializable types.
    - Decimal -> float (preserving value)
    - datetime/timestamp/date -> ISO string
    - bytes -> utf-8 string
    Leaves lists/dicts intact but normalizes their contents.
    """
    try:
        from decimal import Decimal
        import datetime as _dt
    except Exception:
        return obj

    if obj is None:
        return None
    if isinstance(obj, Decimal):
        try:
            return float(obj)
        except Exception:
            return str(obj)
    if isinstance(obj, (_dt.datetime, _dt.date, _dt.time)):
        try:
            return obj.isoformat()
        except Exception:
            return str(obj)
    if isinstance(obj, bytes):
        try:
            return obj.decode('utf-8')
        except Exception:
            return str(obj)
    if isinstance(obj, dict):
        return {k: _normalize_for_json(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [ _normalize_for_json(v) for v in obj ]
    return obj


def _resolve_username(request):
    """Return a best-effort username/full_name for the current request.
    Tries `request.user`, then token payload fields (`username`), then
    fallback lookup by `user_id` claim. Returns empty string when unknown.
    """
    try:
        # 1) request.user (normal DRF auth)
        auth_user = getattr(request, 'user', None)
        if auth_user and getattr(auth_user, 'is_authenticated', False):
            # prefer username then full_name; if both empty, continue to token fallback
            uname = getattr(auth_user, 'username', None) or getattr(auth_user, 'full_name', None)
            if uname:
                return uname

        # 2) token payload (try to decode Authorization header)
        auth_header = request.META.get('HTTP_AUTHORIZATION', '')
        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header.split(' ', 1)[1].strip()
            # Prefer SimpleJWT AccessToken decoding (verifies signature)
            try:
                payload = AccessToken(token)
                # prefer explicit username if present
                uname = payload.get('username') or payload.get('user') or payload.get('user_name')
                if uname:
                    return uname
                # otherwise try user_id lookup
                user_id = payload.get('user_id') or payload.get('uid') or payload.get('id')
                if user_id:
                    u = User.objects.filter(pk=user_id).first()
                    if u:
                        return getattr(u, 'username', None) or getattr(u, 'full_name', None) or ''
            except Exception:
                # Fallback: decode JWT without verification to extract user_id
                try:
                    import jwt
                    payload_raw = jwt.decode(token, options={"verify_signature": False})
                    uname = payload_raw.get('username') or payload_raw.get('user') or payload_raw.get('user_name')
                    if uname:
                        return uname
                    user_id = payload_raw.get('user_id') or payload_raw.get('uid') or payload_raw.get('id')
                    if user_id:
                        u = User.objects.filter(pk=user_id).first()
                        if u:
                            return getattr(u, 'username', None) or getattr(u, 'full_name', None) or ''
                except Exception:
                    pass
    except Exception:
        pass
    return ''


def _ensure_synthesized_pk(cursor, cols_meta, data_map, table_name):
    """If the table's primary key column is present but not auto_increment
    and not supplied in data_map, synthesize a new value using MAX(pk)+1
    and insert it into data_map. This helps support legacy schemas where
    PK defaults to 0 and would otherwise cause duplicate-key errors.
    """
    try:
        pk_row = next((r for r in cols_meta if (r[3] or '').upper() == 'PRI'), None)
        if not pk_row:
            # try common fallback names
            candidate = next((r for r in cols_meta if (r[0] or '').lower() in ('id','contract_id','uv_collateral_id','bl_collateral_id')), None)
            pk_row = candidate
        if not pk_row:
            return
        pk_name = pk_row[0]
        extra = (pk_row[5] or '').lower()
        if pk_name in data_map:
            return
        if 'auto_increment' in extra:
            return
        # compute next id
        try:
            cursor.execute(f"SELECT COALESCE(MAX({pk_name}), 0) + 1 FROM {table_name}")
            row = cursor.fetchone()
            if row and row[0] is not None:
                data_map[pk_name] = row[0]
        except Exception:
            return
    except Exception:
        return


def _safe_rmtree(path, retries=6, delay=0.25):
    """Remove directory tree with retries to handle Windows file locks."""
    last_exc = None
    for attempt in range(retries):
        try:
            shutil.rmtree(path)
            logger.info('Removed temporary path %s (attempt %d)', path, attempt+1)
            return
        except Exception as e:
            last_exc = e
            time.sleep(delay * (attempt + 1))
    # final attempt with onerror handler
    def _onerror(func, p, exc_info):
        try:
            os.chmod(p, 0o700)
            func(p)
        except Exception:
            pass
    try:
        shutil.rmtree(path, onerror=_onerror)
        logger.info('Removed temporary path %s with onerror handler', path)
    except Exception:
        logger.warning('Failed to remove path %s after retries: %s', path, str(last_exc))


def _convert_docx_to_pdf(docx_path, pdf_path, retries=2, min_size=2048):
    # Delegate to central helper in `myproject.common` so the configured
    # `SOFFICE_PATH` is respected across apps.
    try:
        from myproject.common import _convert_docx_to_pdf as _common_convert
        return _common_convert(docx_path, pdf_path)
    except Exception as e:
        return False, f'Conversion helper import failed: {e}'
 
def home(request):
    return HttpResponse('OK')


# Helper: format number with dot thousands separator, e.g. 1000000 -> '1.000.000'
def format_number_dot(val):
    if val is None:
        return val
    try:
        s = str(val).strip()
        if s == '':
            return ''
        s = s.replace(',', '.')
        num = float(s)
        if abs(num - int(num)) < 1e-9:
            n = int(num)
            return format(n, ',').replace(',', '.')
        whole = int(num)
        frac = abs(num - whole)
        frac_str = ('%.2f' % frac)[2:]
        return format(whole, ',').replace(',', '.') + '.' + frac_str
    except Exception:
        return val


def number_to_indonesian_words(val):
    units = ['', 'satu','dua','tiga','empat','lima','enam','tujuh','delapan','sembilan']

    def _sentence_case(s):
        try:
            if not s:
                return ''
            s = str(s).strip()
            return s[0].upper() + s[1:].lower() if len(s) > 1 else s.upper()
        except Exception:
            return s

    def spell_int(n):
        n = int(n)
        if n < 10:
            return units[n]
        if n < 20:
            if n == 10:
                return 'sepuluh'
            if n == 11:
                return 'sebelas'
            return units[n-10] + ' belas'
        if n < 100:
            tens = n // 10
            rest = n % 10
            return (units[tens] + ' puluh' + (' ' + units[rest] if rest else '')).strip()
        if n < 200:
            return 'seratus' + (' ' + spell_int(n-100) if n-100 else '')
        if n < 1000:
            hundreds = n // 100
            rest = n % 100
            return (units[hundreds] + ' ratus' + (' ' + spell_int(rest) if rest else '')).strip()
        if n < 2000:
            return 'seribu' + (' ' + spell_int(n-1000) if n-1000 else '')
        if n < 1000000:
            thousands = n // 1000
            rest = n % 1000
            return (spell_int(thousands) + ' ribu' + (' ' + spell_int(rest) if rest else '')).strip()
        if n < 1000000000:
            millions = n // 1000000
            rest = n % 1000000
            return (spell_int(millions) + ' juta' + (' ' + spell_int(rest) if rest else '')).strip()
        if n < 1000000000000:
            billions = n // 1000000000
            rest = n % 1000000000
            return (spell_int(billions) + ' miliar' + (' ' + spell_int(rest) if rest else '')).strip()
        return str(n)

    try:
        s = str(val).strip()
        if s == '':
            return ''
        s = s.replace(',', '.')
        res = ''
        if '.' in s:
            int_part, dec_part = s.split('.', 1)
            int_words = spell_int(int(int_part)) if int_part and int(int_part) != 0 else ('nol' if int_part == '0' else '')
            dec_words = ' '.join([units[int(d)] if d.isdigit() else d for d in dec_part])
            if int_words:
                res = (int_words + ' koma ' + dec_words).strip()
            else:
                res = ('koma ' + dec_words).strip()
        else:
            res = spell_int(int(float(s)))
        try:
            return str(res).upper()
        except Exception:
            return res
    except Exception:
        return str(val)


def date_to_indonesian_words(val, title_case=False, uppercase_month=False, uppercase_all=False):
    from datetime import datetime, date
    months = ['januari','februari','maret','april','mei','juni','juli','agustus','september','oktober','november','desember']
    if val is None:
        return ''
    try:
        if isinstance(val, (datetime,)):
            d = val.date()
        elif isinstance(val, date):
            d = val
        else:
            s = str(val).strip()
            if not s:
                return ''
            try:
                d = datetime.fromisoformat(s).date()
            except Exception:
                try:
                    d = datetime.strptime(s, '%Y-%m-%d').date()
                except Exception:
                    return ''
        day_word = number_to_indonesian_words(d.day, title_case=title_case)
        month_raw = months[d.month-1]
        year_word = number_to_indonesian_words(d.year, title_case=title_case)

        if uppercase_month:
            month = month_raw.upper()
        else:
            month = month_raw

        combined = f"{day_word} {month} {year_word}".strip()
        if uppercase_all:
            try:
                return combined.upper()
            except Exception:
                return combined
        try:
            if not combined:
                return ''
            return combined.upper()
        except Exception:
            return combined
    except Exception:
        return ''


def format_indonesian_date(val, uppercase_all=False):
    """Return date as '5 Januari 2025' (month name capitalized) for given value.
    Accepts date/datetime or ISO/date-like strings.
    """
    from datetime import datetime, date
    months = ['Januari','Februari','Maret','April','Mei','Juni','Juli','Agustus','September','Oktober','November','Desember']
    if val is None:
        return ''
    try:
        if isinstance(val, datetime):
            d = val.date()
        elif isinstance(val, date):
            d = val
        else:
            s = str(val).strip()
            if not s:
                return ''
            try:
                d = datetime.fromisoformat(s).date()
            except Exception:
                try:
                    d = datetime.strptime(s, '%Y-%m-%d').date()
                except Exception:
                    try:
                        d = datetime.strptime(s, '%d/%m/%Y').date()
                    except Exception:
                        return ''
        out = f"{d.day} {months[d.month-1]} {d.year}"
        return out.upper() if uppercase_all else out
    except Exception:
        return ''


class SimpleLoginView(APIView):
    """
    Simple login endpoint that checks credentials against `auth_user` table
    and returns JWT tokens. Kept minimal and raw-SQL to avoid ORM differences.
    """
    permission_classes = []

    def post(self, request):
        from django.contrib.auth.hashers import check_password
        from rest_framework_simplejwt.tokens import RefreshToken

        username = request.data.get('username')
        password = request.data.get('password')

        if not username or not password:
            return Response({'error': 'Username dan password harus diisi'}, status=status.HTTP_400_BAD_REQUEST)

        with connection.cursor() as cursor:
            cursor.execute('SELECT id, username, password, email, full_name, role, is_staff, branch_id, area_id, region_id FROM auth_user WHERE username=%s', [username])
            row = cursor.fetchone()

        if not row:
            return Response({'error': 'Username atau password salah'}, status=status.HTTP_401_UNAUTHORIZED)

        user_id, db_username, db_password, db_email, db_full_name, db_role, db_is_staff, db_branch_id, db_area_id, db_region_id = row

        if not check_password(password, db_password):
            return Response({'error': 'Username atau password salah'}, status=status.HTTP_401_UNAUTHORIZED)

        refresh = RefreshToken()
        refresh['user_id'] = user_id
        refresh['username'] = db_username
        access_token = refresh.access_token

        return Response({
            'access': str(access_token),
            'refresh': str(refresh),
            'user': {
                'id': user_id,
                'username': db_username,
                'email': db_email,
                'full_name': db_full_name,
                'role': db_role,
                'is_staff': bool(db_is_staff),
                'branch_id': db_branch_id,
                'area_id': db_area_id,
                'region_id': db_region_id,
                }
            })


class UserListCreateView(APIView):
    # Allow anonymous POSTs so frontend can create SP3 rows without requiring auth
    permission_classes = [AllowAny]
    # Bypass default authentication classes for this view to avoid 401 on anonymous POSTs
    authentication_classes = []

    def get(self, request):
        try:
            with connection.cursor() as cursor:
                # Include is_active so clients can display correct active/inactive status
                cursor.execute("SELECT id, username, email, full_name, role, is_active FROM auth_user ORDER BY username")
                cols = [c[0] for c in cursor.description] if cursor.description else []
                rows = cursor.fetchall()
                users = [dict(zip(cols, r)) for r in rows]
            return Response({'users': users}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class DownloadLogListView(APIView):
    permission_classes = [IsAuthenticated, RolePermission]
    # Only users with these roles may list/download logs
    required_roles = ['Admin', 'BOD']

    def get(self, request):
        qs = DownloadLog.objects.all().order_by('-timestamp')
        file_type = request.query_params.get('file_type')
        if file_type:
            qs = qs.filter(file_type=file_type)
        username = request.query_params.get('username')
        if username:
            qs = qs.filter(username__icontains=username)
        user_id = request.query_params.get('user_id')
        if user_id:
            try:
                qs = qs.filter(user_id=int(user_id))
            except Exception:
                pass
        contract = request.query_params.get('contract_number') or request.query_params.get('file_identifier')
        if contract:
            qs = qs.filter(file_identifier__icontains=str(contract))
        date_from = request.query_params.get('date_from')
        date_to = request.query_params.get('date_to')
        try:
            from django.utils.dateparse import parse_datetime, parse_date
            if date_from:
                dt = parse_datetime(date_from) or parse_date(date_from)
                if dt:
                    qs = qs.filter(timestamp__gte=dt)
            if date_to:
                dt2 = parse_datetime(date_to) or parse_date(date_to)
                if dt2:
                    qs = qs.filter(timestamp__lte=dt2)
        except Exception:
            pass

        # Pagination: support `page` (1-based) and `limit` (page size)
        page = request.query_params.get('page')
        limit = request.query_params.get('limit')
        try:
            page = int(page) if page else 1
        except Exception:
            page = 1
        try:
            limit = int(limit) if limit else 100
        except Exception:
            limit = 100

        start = (page - 1) * limit
        end = start + limit

        items = []
        for r in qs[start:end]:
            items.append({
                'id': r.id,
                'user_id': r.user_id,
                'username': r.username,
                'email': r.email,
                'file_type': r.file_type,
                'file_identifier': r.file_identifier,
                'filename': r.filename,
                'timestamp': r.timestamp.isoformat() if r.timestamp else None,
                'ip_address': r.ip_address,
                'user_agent': r.user_agent,
                'success': r.success,
                'file_size': r.file_size,
                'method': r.method,
            })

        return Response({'count': qs.count(), 'results': items}, status=status.HTTP_200_OK)
    def post(self, request):
        # create via serializer when available
        serializer = UserSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({'message': 'User created successfully'}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class UserDetailView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            with connection.cursor() as cursor:
                # Return richer contract objects similar to /api/bl-agreement/contracts/
                cursor.execute("""
                    SELECT DISTINCT contract_id, contract_number, name_of_debtor, nik_number_of_debtor, loan_amount
                    FROM (
                        SELECT id AS contract_id, contract_number, COALESCE(name_of_debtor, '') AS name_of_debtor,
                               COALESCE(nik_number_of_debtor, '') AS nik_number_of_debtor,
                               COALESCE(loan_amount, '') AS loan_amount
                        FROM contract
                        UNION
                        SELECT NULL AS contract_id, contract_number, COALESCE(name_of_debtor, '') AS name_of_debtor,
                               COALESCE(nik_number_of_debtor, '') AS nik_number_of_debtor,
                               COALESCE(loan_amount, '') AS loan_amount
                        FROM bl_agreement
                        UNION
                        SELECT NULL AS contract_id, contract_number, '' AS name_of_debtor, '' AS nik_number_of_debtor, '' AS loan_amount
                        FROM bl_collateral
                    ) t
                    ORDER BY contract_number
                """)

                rows = cursor.fetchall()
                contracts = []
                for r in rows:
                    contracts.append({
                        'contract_id': r[0],
                        'contract_number': r[1],
                        'name_of_debtor': r[2],
                        'nik_number_of_debtor': r[3],
                        'loan_amount': r[4],
                    })

                # include column names so frontend can render dynamic contract fields
                return Response({'contracts': contracts, 'columns': cols}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request):
        """Create a new director record."""
        name = request.data.get('name') or request.data.get('name_of_director')
        phone = request.data.get('phone_number_of_lolc') or request.data.get('phone')
        if not name:
            return Response({'error': 'Missing name'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                # Some DB schemas have director_id NOT NULL without AUTO_INCREMENT.
                # Try to generate a new director_id if needed by taking MAX(director_id)+1.
                new_id = None
                try:
                    cursor.execute("SELECT COALESCE(MAX(director_id), 0) + 1 FROM director")
                    row = cursor.fetchone()
                    if row:
                        new_id = row[0]
                except Exception:
                    new_id = None

                if new_id is not None:
                    cursor.execute("INSERT INTO director (director_id, name_of_director, phone_number_of_lolc) VALUES (%s, %s, %s)", [new_id, name, phone])
                else:
                    cursor.execute("INSERT INTO director (name_of_director, phone_number_of_lolc) VALUES (%s, %s)", [name, phone])
            return Response({'status': 'created'}, status=status.HTTP_201_CREATED)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def patch(self, request, pk=None):
        """Update an existing director by id (pk)."""
        if not pk:
            return Response({'error': 'Missing director id'}, status=status.HTTP_400_BAD_REQUEST)
        name = request.data.get('name') or request.data.get('name_of_director')
        phone = request.data.get('phone_number_of_lolc') or request.data.get('phone')
        if name is None and phone is None:
            return Response({'error': 'No fields to update'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                updates = []
                params = []
                if name is not None:
                    updates.append('name_of_director=%s')
                    params.append(name)
                if phone is not None:
                    updates.append('phone_number_of_lolc=%s')
                    params.append(phone)
                params.append(pk)
                sql = f"UPDATE director SET {', '.join(updates)} WHERE director_id=%s"
                cursor.execute(sql, params)
            return Response({'status': 'updated'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def delete(self, request, pk=None):
        """Soft-delete a director by is_active if available, else delete row."""
        if not pk:
            return Response({'error': 'Missing director id'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                # Try soft-delete first
                try:
                    cursor.execute("UPDATE director SET is_active=0 WHERE director_id=%s", [pk])
                except Exception:
                    # fallback to hard delete
                    cursor.execute("DELETE FROM director WHERE director_id=%s", [pk])
            return Response({'status': 'deleted'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def get(self, request, username):
        try:
            with connection.cursor() as cursor:
                cursor.execute('SELECT * FROM auth_user WHERE username=%s LIMIT 1', [username])
                cols = [c[0] for c in cursor.description] if cursor.description else []
                row = cursor.fetchone()
                user = dict(zip(cols, row)) if row else None
            if not user:
                return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
            return Response({'user': user}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def put(self, request, username):
        # basic update: allow email, full_name, password, role, phone
        data = request.data or {}
        allowed = ['email', 'full_name', 'password', 'role', 'phone']
        fields = []
        params = []
        try:
            if 'password' in data and data.get('password'):
                hashed = make_password(data.get('password'))
                fields.append('password=%s')
                params.append(hashed)
            for k in allowed:
                if k == 'password':
                    continue
                if k in data:
                    fields.append(f"{k}=%s")
                    params.append(data.get(k))
            if not fields:
                return Response({'message': 'No fields to update'}, status=status.HTTP_200_OK)
            params.append(username)
            sql = 'UPDATE auth_user SET ' + ', '.join(fields) + ' WHERE username=%s'
            with connection.cursor() as cursor:
                cursor.execute(sql, params)
                cursor.execute('SELECT id, username, email, full_name, role FROM auth_user WHERE username=%s', [username])
                cols = [c[0] for c in cursor.description] if cursor.description else []
                row = cursor.fetchone()
                user = dict(zip(cols, row)) if row else None
            return Response({'user': user, 'message': 'User updated'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def patch(self, request, username):
        """Support HTTP PATCH as an alias to PUT for partial updates."""
        return self.put(request, username)

    def patch(self, request, username):
        """Support HTTP PATCH as an alias to PUT for partial updates."""
        return self.put(request, username)


class CustomTokenObtainPairView(JWTTokenObtainPairView):
    permission_classes = []
    pass


class CustomTokenRefreshView(JWTTokenRefreshView):
    pass


class ProtectedView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = getattr(request, 'user', None)
        return Response({
            'message': 'Welcome to the dashboard',
            'user': getattr(user, 'username', None),
            'full_name': getattr(user, 'full_name', getattr(user, 'username', None)),
            'role': getattr(user, 'role', 'User'),
            'email': getattr(user, 'email', None),
        })


class DashboardSummaryView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Role-aware counts: Admin sees all, CSA sees rows they created,
        # branch/area/region roles see rows matching their branch/area/region.
        def safe_count_sql(sql, params=None):
            try:
                with connection.cursor() as cursor:
                    cursor.execute(sql, params or [])
                    row = cursor.fetchone()
                    return int(row[0]) if row and row[0] is not None else 0
            except Exception:
                return 0

        # Resolve user identity and role
        user_id = None
        username = None
        user_role = None
        user_branch = None
        user_area = None
        user_region = None
        try:
            if getattr(request, 'user', None) and getattr(request.user, 'is_authenticated', False):
                user_id = getattr(request.user, 'id', None)
                username = getattr(request.user, 'username', None)
            if not user_id:
                auth_header = request.META.get('HTTP_AUTHORIZATION', '')
                if auth_header and auth_header.startswith('Bearer '):
                    try:
                        payload = AccessToken(auth_header.split(' ', 1)[1].strip())
                        user_id = payload.get('user_id') or payload.get('uid') or payload.get('id')
                        username = payload.get('username') or username
                    except Exception:
                        pass
            if user_id:
                with connection.cursor() as cursor:
                    cursor.execute('SELECT username, role, branch_id, area_id, region_id FROM auth_user WHERE id=%s', [user_id])
                    row = cursor.fetchone()
                    if row:
                        username, user_role, user_branch, user_area, user_region = row
        except Exception:
            pass

        # Build WHERE clauses based on role
        bl_where = ''
        bl_params = []
        uv_where = ''
        uv_params = []

        if user_role:
            r = (user_role or '').strip().lower()
            if r == 'csa':
                # CSA must be scoped strictly to their branch_id. If branch_id
                # is not available, return zero-counts by using a false WHERE.
                if user_branch:
                    bl_where = 'WHERE branch_id=%s'
                    bl_params = [user_branch]
                    uv_where = 'WHERE branch_id=%s'
                    uv_params = [user_branch]
                else:
                    bl_where = 'WHERE 1=0'
                    uv_where = 'WHERE 1=0'
            elif r in ('bm', 'branchmanager', 'branch_manager'):
                if user_branch:
                    bl_where = 'WHERE branch_id=%s'
                    bl_params = [user_branch]
                    uv_where = 'WHERE branch_id=%s'
                    uv_params = [user_branch]
            elif r in ('area', 'areamanager', 'area_manager'):
                if user_area:
                    bl_where = 'WHERE area_id=%s'
                    bl_params = [user_area]
                    uv_where = 'WHERE area_id=%s'
                    uv_params = [user_area]
            elif r in ('region', 'regionmanager', 'region_manager'):
                if user_region:
                    bl_where = 'WHERE region_id=%s'
                    bl_params = [user_region]
                    uv_where = 'WHERE region_id=%s'
                    uv_params = [user_region]
            else:
                # default: admins and others see all
                bl_where = ''
                uv_where = ''

        # Fallback: if no role resolution, return global counts
        bl_sql = f"SELECT COUNT(*) FROM bl_agreement {bl_where}".strip()
        uv_sql = f"SELECT COUNT(*) FROM uv_agreement {uv_where}".strip()

        data = {
            'bl_agreement': safe_count_sql(bl_sql, bl_params),
            'uv_agreement': safe_count_sql(uv_sql, uv_params),
        }
        return Response(data)


class UVAgreementView(APIView):
    """
    Mirrors BLAgreementView behavior but operates on uv_agreement and uv_collateral tables.
    This keeps BL endpoints untouched while providing the same save/list/get semantics
    for UV agreements.
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = request.data
        contract_number = data.get('contract_number')
        branch_id = data.get('branch_id')
        director = data.get('director')
        bm_data = data.get('bm_data', {})
        contract_data = data.get('contract_data', {})
        collateral_data = data.get('collateral_data', {})
        header_fields = data.get('header_fields', {})

        try:
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM uv_agreement")
                cols_meta = cursor.fetchall()
                cols_info = [row[0] for row in cols_meta]
                # detect primary key column name (if any)
                try:
                    pk_row = next((r for r in cols_meta if (r[3] or '').upper() == 'PRI'), None)
                    pk_col_name = pk_row[0] if pk_row else None
                except Exception:
                    pk_col_name = None
                cols_lookup = {c.lower(): c for c in cols_info}

                data_map = {}
                if contract_number:
                    try:
                        data_map['contract_number'] = str(contract_number).upper()
                    except Exception:
                        data_map['contract_number'] = contract_number

                client_created_by = data.get('created_by')
                if client_created_by and 'created_by' in cols_info:
                    data_map['created_by'] = client_created_by

                user_full_name = None
                user_username = None
                if getattr(request, 'user', None) and getattr(request.user, 'is_authenticated', False):
                    user_full_name = getattr(request.user, 'full_name', None)
                    user_username = getattr(request.user, 'username', None)
                    if (not user_full_name or not user_username) and getattr(request.user, 'id', None):
                        try:
                            u = User.objects.filter(pk=request.user.id).first()
                            if u:
                                user_full_name = user_full_name or getattr(u, 'full_name', None)
                                user_username = user_username or getattr(u, 'username', None)
                        except Exception:
                            pass
                else:
                    auth_header = request.META.get('HTTP_AUTHORIZATION', '')
                    if auth_header and auth_header.startswith('Bearer '):
                        token = auth_header.split(' ', 1)[1].strip()
                        try:
                            payload = AccessToken(token)
                            user_id = payload.get('user_id') or payload.get('uid') or payload.get('id')
                            if user_id:
                                try:
                                    u = User.objects.filter(pk=user_id).first()
                                    if u:
                                        user_full_name = user_full_name or getattr(u, 'full_name', None)
                                        user_username = user_username or getattr(u, 'username', None)
                                except Exception:
                                    pass
                        except TokenError:
                            pass

                user_username = user_username or (user_full_name and user_full_name.replace(' ', '').lower()) or 'anonymous'
                now = timezone.now()

                if branch_id is not None and 'branch_id' in cols_lookup:
                    data_map[cols_lookup['branch_id']] = branch_id
                if director is not None:
                    if 'name_of_director' in cols_lookup:
                        data_map[cols_lookup['name_of_director']] = director
                    elif 'director' in cols_lookup:
                        data_map[cols_lookup['director']] = director

                branch_data = data.get('branch_data', {})
                for src in (bm_data, branch_data, contract_data, collateral_data, header_fields):
                    if not isinstance(src, dict):
                        continue
                    for k, v in src.items():
                        key = str(k).lower()
                        if key == 'name_of_director':
                            key = 'name_of_director'
                        if key in cols_lookup:
                            data_map[cols_lookup[key]] = v

                from datetime import datetime, date as _date
                field_type_map = {row[0]: row[1].lower() for row in cols_meta}
                for k in list(data_map.keys()):
                    ft = field_type_map.get(k, '')
                    if not any(t in ft for t in ('date', 'timestamp', 'datetime')):
                        continue
                    val = data_map.get(k)
                    if val is None or (isinstance(val, str) and val.strip() == ''):
                        data_map.pop(k, None)
                        continue
                    if isinstance(val, (_date, datetime)):
                        continue
                    s = str(val).strip()
                    parsed = None
                    try:
                        try:
                            parsed = datetime.fromisoformat(s)
                        except Exception:
                            import re
                            if re.match(r'^\s*\d{1,2}\s+\d{1,2}\s+\d{4}\s*$', s):
                                s = re.sub(r'\s+', '/', s.strip())
                            for fmt in ('%Y-%m-%d', '%d-%m-%Y', '%d/%m/%Y', '%Y/%m/%d', '%d %m %Y', '%d %b %Y', '%d %B %Y'):
                                try:
                                    parsed = datetime.strptime(s, fmt)
                                    break
                                except Exception:
                                    continue
                        if parsed is None:
                            data_map.pop(k, None)
                            continue
                        if 'date' in ft and 'time' not in ft and 'datetime' not in ft:
                            data_map[k] = parsed.date()
                        else:
                            if parsed.tzinfo is None:
                                parsed = timezone.make_aware(parsed)
                            data_map[k] = parsed
                    except Exception:
                        data_map.pop(k, None)

                    # Coerce empty-string values for numeric columns to 0 to avoid
                    # MySQL "Incorrect integer value: ''" errors when clients send
                    # empty strings for numeric fields (common from HTML inputs).
                    try:
                        for col_name, col_type in field_type_map.items():
                            if col_name in data_map:
                                v = data_map[col_name]
                                if isinstance(v, str) and v.strip() == '':
                                    if any(t in col_type for t in ('int', 'decimal', 'float', 'double')):
                                        # prefer integer where appropriate
                                        if 'decimal' in col_type or 'float' in col_type or 'double' in col_type:
                                            data_map[col_name] = 0.0
                                        else:
                                            data_map[col_name] = 0
                    except Exception:
                        pass

                if 'id' in data_map:
                    data_map.pop('id', None)

                field_type_map = {row[0]: row[1].lower() for row in cols_meta}
                for _f in ('admin_rate', 'tlo', 'life_insurance'):
                    if _f in cols_lookup and cols_lookup[_f] not in data_map:
                        ftype = field_type_map.get(_f, '')
                        if any(t in ftype for t in ('int', 'decimal', 'float', 'double')):
                            if 'decimal' in ftype or 'float' in ftype or 'double' in ftype:
                                data_map[cols_lookup[_f]] = 0.0
                            else:
                                data_map[cols_lookup[_f]] = 0
                        else:
                            data_map[cols_lookup[_f]] = '-'

                for col_row in cols_meta:
                    field_name = col_row[0]
                    field_type = col_row[1].lower()
                    is_nullable = col_row[2]
                    default_val = col_row[4]
                    if field_name in data_map:
                        continue
                    if field_name in ('id', 'created_by', 'created_at', 'update_at'):
                        continue
                    if is_nullable == 'NO' and default_val is None:
                        if any(t in field_type for t in ('int', 'decimal', 'float', 'double')):
                            data_map[field_name] = 0
                        elif any(t in field_type for t in ('date', 'timestamp', 'datetime')):
                            data_map[field_name] = timezone.now()
                        else:
                            data_map[field_name] = '-'

                existing_check_sql = "SELECT contract_number FROM uv_agreement WHERE contract_number=%s"
                exists = None
                if contract_number:
                    cursor.execute(existing_check_sql, [contract_number])
                    exists = cursor.fetchone()

                edit_only = bool(data.get('edit_only'))
                create_only = bool(data.get('create_only'))

                if exists:
                    if create_only:
                        return Response({'error': 'Record already exists (create_only specified)'}, status=status.HTTP_400_BAD_REQUEST)
                    if 'update_at' in cols_info:
                        data_map['update_at'] = now
                    data_map.pop('created_by', None)
                    data_map.pop('created_at', None)
                    set_cols = []
                    params = []
                    for col, val in data_map.items():
                        if col == 'contract_number':
                            continue
                        set_cols.append(f"{col}=%s")
                        params.append(val)
                    if set_cols:
                        sql = f"UPDATE uv_agreement SET {', '.join(set_cols)} WHERE contract_number=%s"
                        params.append(contract_number)
                        cursor.execute(sql, params)
                else:
                    if edit_only:
                        return Response({'error': 'Record not found for update (edit_only specified)'}, status=status.HTTP_400_BAD_REQUEST)
                    if 'created_by' in cols_info and user_username:
                        data_map['created_by'] = user_username
                    if 'created_at' in cols_info:
                        data_map[cols_lookup['created_at']] = now if 'created_at' in cols_lookup else now
                    cols = []
                    placeholders = []
                    params = []
                    for col, val in data_map.items():
                        cols.append(col)
                        placeholders.append('%s')
                        params.append(val)
                    sql = f"INSERT INTO uv_agreement ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
                    cursor.execute(sql, params)

            return Response({'message': 'Data UV berhasil disimpan'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def get(self, request):
        contract_number = request.query_params.get('contract_number', '').strip()
        try:
            with connection.cursor() as cursor:
                if not contract_number:
                    # Build a safe SELECT list from actual uv_agreement columns to avoid
                    # referencing missing columns (e.g. collateral_type may not exist).
                    cursor.execute("SHOW COLUMNS FROM uv_agreement")
                    cols_meta_short = cursor.fetchall()
                    available_cols = [r[0] for r in cols_meta_short]

                    # inspect uv_collateral for vehicle-like columns to coalesce
                    cursor.execute("SHOW COLUMNS FROM uv_collateral")
                    coll_meta = cursor.fetchall()
                    coll_cols = [r[0] for r in coll_meta]

                    preferred = ['agreement_date', 'contract_number', 'name_of_debtor', 'nik_number_of_debtor', 'vehicle_type', 'vehicle_types', 'vechile_types', 'collateral_type', 'uv_collateral_type', 'created_by']

                    select_parts = []
                    # Always include contract_number
                    if 'contract_number' in available_cols:
                        select_parts.append('uv_agreement.contract_number')

                    # include agreement_date and other preferred columns from uv_agreement if present
                    for p in preferred:
                        if p == 'contract_number':
                            continue
                        if p in available_cols and p not in ('vehicle_types', 'vechile_types', 'collateral_type', 'uv_collateral_type'):
                            select_parts.append(f'uv_agreement.{p}')

                    # Build a COALESCE for vehicle types using uv_agreement then uv_collateral
                    coalesce_sources = []
                    if 'vehicle_types' in available_cols:
                        coalesce_sources.append('uv_agreement.vehicle_types')
                    if 'vechile_types' in available_cols:
                        coalesce_sources.append('uv_agreement.vechile_types')
                    # from uv_collateral prefixed with alias uc
                    for candidate in ('vehicle_types', 'vechile_types', 'collateral_type', 'uv_collateral_type'):
                        if candidate in coll_cols:
                            coalesce_sources.append(f'uc.{candidate}')

                    if coalesce_sources:
                        select_parts.append('COALESCE(' + ', '.join(coalesce_sources) + ') AS vehicle_types')

                    # include created_by if present
                    if 'created_by' in available_cols:
                        select_parts.append('uv_agreement.created_by')

                    if not select_parts:
                        # fallback: select contract_number only
                        select_clause = 'uv_agreement.contract_number'
                    else:
                        select_clause = ', '.join(select_parts)

                    sql = f"SELECT {select_clause} FROM uv_agreement LEFT JOIN uv_collateral uc ON uv_agreement.contract_number = uc.contract_number ORDER BY COALESCE(uv_agreement.agreement_date, uv_agreement.created_at) DESC"
                    cursor.execute(sql)
                    cols = [c[0] for c in cursor.description] if cursor.description else []
                    rows = cursor.fetchall()
                    items = [dict(zip(cols, r)) for r in rows]
                    return Response({'agreements': _normalize_for_json(items)}, status=status.HTTP_200_OK)

                cursor.execute(
                    "SELECT * FROM uv_agreement WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                    [contract_number]
                )
                cols = [col[0] for col in cursor.description] if cursor.description else []
                uv_row = cursor.fetchone()
                uv_data = dict(zip(cols, uv_row)) if uv_row else None

                cursor.execute(
                    "SELECT * FROM uv_collateral WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                    [contract_number]
                )
                collateral_columns = [col[0] for col in cursor.description] if cursor.description else []
                uv_collateral_row = cursor.fetchone()
                collateral_data = dict(zip(collateral_columns, uv_collateral_row)) if uv_collateral_row else None

                if uv_data:
                    try:
                        if 'name_of_director' in uv_data and 'Name_of_director' not in uv_data:
                            uv_data['Name_of_director'] = uv_data.get('name_of_director')
                    except Exception:
                        pass
                    return Response({'debtor': _normalize_for_json(uv_data), 'collateral': _normalize_for_json(collateral_data)}, status=status.HTTP_200_OK)

                cursor.execute(
                    "SELECT * FROM contract WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                    [contract_number]
                )
                columns = [col[0] for col in cursor.description] if cursor.description else []
                contract_row = cursor.fetchone()
                debtor_data = dict(zip(columns, contract_row)) if contract_row else None

                if not debtor_data and not collateral_data:
                    return Response({'debtor': None, 'collateral': None, 'message': 'Data tidak ditemukan untuk nomor kontrak ini'}, status=status.HTTP_404_NOT_FOUND)

                return Response({'debtor': _normalize_for_json(debtor_data), 'collateral': _normalize_for_json(collateral_data)}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



class UVCollateralCreateView(APIView):
    """Create or update UV collateral rows (mirrors BL collateral handler semantics)."""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        """Allow frontend to lookup uv_collateral rows by contract_number via GET.
        Returns JSON: { "collateral": [ { ... }, ... ] }
        """
        contract_number = request.query_params.get('contract_number') or request.GET.get('contract_number')
        # If contract_number not provided, return column metadata so frontend
        # can render full uv_collateral form fields even when no rows exist yet.
        if not contract_number:
            try:
                with connection.cursor() as cursor:
                    cursor.execute("SHOW COLUMNS FROM uv_collateral")
                    cols_meta = cursor.fetchall()
                    cols_info = [row[0] for row in cols_meta]
                return Response({'collateral': [], 'columns': cols_info}, status=status.HTTP_200_OK)
            except Exception as e:
                logger.exception('UVCollateral columns lookup failed')
                return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        try:
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM uv_collateral")
                cols_meta = cursor.fetchall()
                cols_info = [row[0] for row in cols_meta]
                sql = f"SELECT {', '.join(cols_info)} FROM uv_collateral WHERE LOWER(contract_number)=LOWER(%s)"
                cursor.execute(sql, [contract_number])
                rows = cursor.fetchall()
                result = []
                for row in rows:
                    row_dict = {cols_info[i]: row[i] for i in range(len(cols_info))}
                    result.append(row_dict)
                return Response({'collateral': result, 'columns': cols_info}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('UVCollateral lookup failed')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request):
        data = request.data
        contract_number = data.get('contract_number')
        collateral = data.get('collateral', {})
        if not contract_number:
            return Response({'error': 'contract_number is required'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM uv_collateral")
                cols_meta = cursor.fetchall()
                cols_info = [r[0] for r in cols_meta]
                cols_lookup = {c.lower(): c for c in cols_info}

                try:
                    data_map = {'contract_number': str(contract_number).upper()}
                except Exception:
                    data_map = {'contract_number': contract_number}
                if isinstance(collateral, dict):
                    for k, v in collateral.items():
                        key = str(k).lower()
                        if key in cols_lookup:
                            data_map[cols_lookup[key]] = v

                for col_row in cols_meta:
                    field_name = col_row[0]
                    field_type = col_row[1].lower()
                    is_nullable = col_row[2]
                    default_val = col_row[4]
                    if field_name in data_map:
                        continue
                    if field_name in ('id', 'created_by', 'created_at', 'update_at'):
                        continue
                    if is_nullable == 'NO' and default_val is None:
                        if any(t in field_type for t in ('int', 'decimal', 'float', 'double')):
                            data_map[field_name] = 0
                        elif any(t in field_type for t in ('date', 'timestamp', 'datetime')):
                            data_map[field_name] = timezone.now()
                        else:
                            data_map[field_name] = '-'

                # Prevent duplicate uv_collateral rows for same contract_number (case-insensitive)
                try:
                    cursor.execute("SELECT 1 FROM uv_collateral WHERE LOWER(contract_number)=LOWER(%s) LIMIT 1", [contract_number])
                    if cursor.fetchone():
                        return Response({'error': 'Duplicate contract_number: collateral already exists'}, status=status.HTTP_409_CONFLICT)
                except Exception:
                    # if uniqueness check fails for any reason, continue to attempt insert
                    pass

                # No existing row found -> insert new UV collateral row
                
                    # server-side audit fields for new uv_collateral rows
                    now = timezone.now()
                    username = _resolve_username(request)
                    if 'created_by' in cols_info and 'created_by' not in data_map:
                        data_map['created_by'] = username or ''
                    if 'created_at' in cols_info and 'created_at' not in data_map:
                        data_map['created_at'] = now
                    if 'update_at' in cols_info and 'update_at' not in data_map:
                        data_map['update_at'] = now

                    # If primary key column is present but not auto_increment, synthesize a new PK
                    try:
                        # explicit fallback for common legacy PK name
                        if 'uv_collateral_id' in cols_info and 'uv_collateral_id' not in data_map:
                            try:
                                cursor.execute("SELECT COALESCE(MAX(uv_collateral_id),0)+1 FROM uv_collateral")
                                v = cursor.fetchone()
                                if v:
                                    data_map['uv_collateral_id'] = v[0]
                            except Exception:
                                pass
                        pk_col_row = next((r for r in cols_meta if (r[3] or '').upper() == 'PRI'), None)
                        if pk_col_row is not None:
                            pk_name = pk_col_row[0]
                            extra = (pk_col_row[5] or '').lower()
                            if 'auto_increment' not in extra and pk_name not in data_map:
                                cursor.execute(f"SELECT COALESCE(MAX({pk_name}), 0) + 1 FROM uv_collateral")
                                new_pk = cursor.fetchone()
                                if new_pk:
                                    data_map[pk_name] = new_pk[0]
                        else:
                            # fallback: pick a plausible PK column (e.g., *_id or id)
                            candidate = next((r[0] for r in cols_meta if (r[0] or '').lower().endswith('_id') or (r[0] or '').lower() == 'id'), None)
                            if candidate and candidate not in data_map:
                                try:
                                    cursor.execute(f"SELECT COALESCE(MAX({candidate}), 0) + 1 FROM uv_collateral")
                                    new_pk = cursor.fetchone()
                                    if new_pk:
                                        data_map[candidate] = new_pk[0]
                                except Exception:
                                    pass
                    except Exception:
                        pass
                    # Ensure PK synthesized for legacy uv_collateral schemas
                    _ensure_synthesized_pk(cursor, cols_meta, data_map, 'uv_collateral')
                    cols = []
                    placeholders = []
                    params = []
                    for col, val in data_map.items():
                        cols.append(col)
                        placeholders.append('%s')
                        params.append(val)
                    sql = f"INSERT INTO uv_collateral ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
                    try:
                        cursor.execute(sql, params)
                    except IntegrityError as ie:
                        # likely a UNIQUE constraint on contract_number (or similar)
                        return Response({'error': 'Duplicate contract_number: collateral already exists'}, status=status.HTTP_409_CONFLICT)

            return Response({'message': 'UV collateral saved'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    permission_classes = [IsAuthenticated]

    def put(self, request, username):
        # Resolve target user id
        try:
            if str(username).isdigit():
                target_id = int(username)
            else:
                with connection.cursor() as cursor:
                    cursor.execute('SELECT id FROM auth_user WHERE username=%s', [username])
                    row = cursor.fetchone()
                    target_id = row[0] if row else None
            if not target_id:
                return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception:
            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)

        # Collect fields to update
        fields = []
        params = []

        email = request.data.get('email')
        password = request.data.get('password')
        role = request.data.get('role')
        phone = request.data.get('phone')
        employee_id = request.data.get('employee_id')
        region_id = request.data.get('region_id')
        area_id = request.data.get('area_id')
        branch_id = request.data.get('branch_id')
        full_name = request.data.get('full_name')
        is_active = request.data.get('is_active')

        # Email duplicate check
        if email is not None:
            with connection.cursor() as cursor:
                cursor.execute('SELECT id FROM auth_user WHERE email=%s AND id<>%s', [email, target_id])
                if cursor.fetchone():
                    return Response({'error': 'Email already exists'}, status=status.HTTP_400_BAD_REQUEST)
            fields.append('email=%s')
            params.append(email)

        if password:
            hashed = make_password(password)
            fields.append('password=%s')
            params.append(hashed)

        if role is not None:
            fields.append('role=%s')
            params.append(role)
            is_staff_val = 1 if role == 'Administrator' else 0
            fields.append('is_staff=%s')
            params.append(is_staff_val)

        if full_name is not None:
            fields.append('full_name=%s')
            params.append(full_name)
        if phone is not None:
            fields.append('phone=%s')
            params.append(phone)
        if employee_id is not None:
            fields.append('employee_id=%s')
            params.append(employee_id)
        if region_id is not None:
            fields.append('region_id=%s')
            params.append(region_id)
        if area_id is not None:
            fields.append('area_id=%s')
            params.append(area_id)
        if branch_id is not None:
            fields.append('branch_id=%s')
            params.append(branch_id)
        if is_active is not None:
            fields.append('is_active=%s')
            params.append(1 if is_active else 0)

        if not fields:
            return Response({'message': 'No fields to update'}, status=status.HTTP_200_OK)

        update_sql = 'UPDATE auth_user SET ' + ', '.join(fields) + ' WHERE id=%s'
        params.append(target_id)

        try:
            with connection.cursor() as cursor:
                cursor.execute(update_sql, params)
                cursor.execute('''
                    SELECT id, username, password, email, phone, employee_id, role, 
                           region_id, area_id, branch_id, full_name, is_active, is_staff, 
                           last_login, is_superuser, date_joined
                    FROM auth_user WHERE id=%s
                ''', [target_id])
                cols = [c[0] for c in cursor.description] if cursor.description else []
                row = cursor.fetchone()
                user_data = dict(zip(cols, row)) if row else None

            return Response({'user': user_data, 'message': 'User updated successfully'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class BLAgreementView(APIView):
    def post(self, request):
            """
            Simpan data BL Agreement ke tabel bl_agreement.
            """
            data = request.data
            contract_number = data.get('contract_number')
            branch_id = data.get('branch_id')
            director = data.get('director')
            bm_data = data.get('bm_data', {})
            contract_data = data.get('contract_data', {})
            collateral_data = data.get('collateral_data', {})
            header_fields = data.get('header_fields', {})

            # contract_number is optional now; allow saving without it

            try:
                with connection.cursor() as cursor:
                    # Determine actual columns in bl_agreement table
                    # SHOW COLUMNS returns: Field, Type, Null, Key, Default, Extra
                    cursor.execute("SHOW COLUMNS FROM bl_agreement")
                    cols_meta = cursor.fetchall()
                    cols_info = [row[0] for row in cols_meta]
                    # detect primary key column name (if any)
                    try:
                        pk_row = next((r for r in cols_meta if (r[3] or '').upper() == 'PRI'), None)
                        pk_col_name = pk_row[0] if pk_row else None
                    except Exception:
                        pk_col_name = None
                    # map lowercase column name -> actual column name to handle case differences
                    cols_lookup = {c.lower(): c for c in cols_info}

                    # Build a map of values to store, only for columns that actually exist
                    data_map = {}
                    # include contract_number only if provided
                    if contract_number:
                        data_map['contract_number'] = contract_number

                    # If frontend provided a created_by field (usernameDisplay), accept it
                    client_created_by = data.get('created_by')
                    if client_created_by and 'created_by' in cols_info:
                        data_map['created_by'] = client_created_by

                    # determine current user full name for created_by
                    # Resolve both username and full_name; prefer username for created_by
                    user_full_name = None
                    user_username = None
                    if getattr(request, 'user', None) and getattr(request.user, 'is_authenticated', False):
                        # request.user may be a minimal CustomUser (created by RawSQLJWTAuthentication)
                        user_full_name = getattr(request.user, 'full_name', None)
                        user_username = getattr(request.user, 'username', None)
                        if (not user_full_name or not user_username) and getattr(request.user, 'id', None):
                            try:
                                u = User.objects.filter(pk=request.user.id).first()
                                if u:
                                    if not user_full_name:
                                        user_full_name = getattr(u, 'full_name', None)
                                    if not user_username:
                                        user_username = getattr(u, 'username', None)
                            except Exception:
                                user_full_name = user_full_name or None
                                user_username = user_username or None
                    else:
                        # Try to decode JWT token from Authorization header as a fallback
                        auth_header = request.META.get('HTTP_AUTHORIZATION', '')
                        if auth_header and auth_header.startswith('Bearer '):
                            token = auth_header.split(' ', 1)[1].strip()
                            try:
                                payload = AccessToken(token)
                                user_id = payload.get('user_id') or payload.get('uid') or payload.get('id')
                                if user_id:
                                    try:
                                        u = User.objects.filter(pk=user_id).first()
                                        if u:
                                            user_full_name = getattr(u, 'full_name', None)
                                            user_username = getattr(u, 'username', None)
                                    except Exception:
                                        user_full_name = user_full_name or None
                                        user_username = user_username or None
                            except TokenError:
                                user_full_name = user_full_name or None
                                user_username = user_username or None
                    if not user_username:
                        # fallback to username from full_name or anonymous
                        user_username = user_username or (user_full_name and user_full_name.replace(' ', '').lower()) or 'anonymous'

                    now = timezone.now()

                    # Map simple top-level fields if corresponding columns exist
                    if branch_id is not None and 'branch_id' in cols_lookup:
                        data_map[cols_lookup['branch_id']] = branch_id
                    if director is not None:
                        # prefer column name 'name_of_director' if present, otherwise 'director'
                        if 'name_of_director' in cols_lookup:
                            data_map[cols_lookup['name_of_director']] = director
                        elif 'director' in cols_lookup:
                            data_map[cols_lookup['director']] = director

                    # Merge nested dictionaries (bm_data, branch_data, contract_data, collateral_data, header_fields)
                    # include branch_data so fields like `street_name` are processed when sent from frontend
                    branch_data = data.get('branch_data', {})
                    for src in (bm_data, branch_data, contract_data, collateral_data, header_fields):
                        if not isinstance(src, dict):
                            continue
                        for k, v in src.items():
                            key = str(k).lower()
                            # normalize some header field names
                            if key == 'name_of_director':
                                key = 'name_of_director'
                            # if column exists (case-insensitive), set using actual column name
                            if key in cols_lookup:
                                data_map[cols_lookup[key]] = v

                    # Sanitize date/datetime/timestamp fields provided by client.
                    # If a client supplies an unparseable value (e.g. 'B'), remove it
                    # so the NOT NULL default handling below can fill a safe value.
                    from datetime import datetime, date as _date
                    field_type_map = {row[0]: row[1].lower() for row in cols_meta}
                    for k in list(data_map.keys()):
                        ft = field_type_map.get(k, '')
                        if not any(t in ft for t in ('date', 'timestamp', 'datetime')):
                            continue
                        val = data_map.get(k)
                        # empty strings -> treat as None
                        if val is None or (isinstance(val, str) and val.strip() == ''):
                            data_map.pop(k, None)
                            continue
                        # already a date/datetime instance -> keep
                        if isinstance(val, (_date, datetime)):
                            continue
                        s = str(val).strip()
                        parsed = None
                        try:
                            try:
                                parsed = datetime.fromisoformat(s)
                            except Exception:
                                # Normalize spaced numeric dates like '15 04 2026' -> '15/04/2026'
                                import re
                                if re.match(r'^\s*\d{1,2}\s+\d{1,2}\s+\d{4}\s*$', s):
                                    s = re.sub(r'\s+', '/', s.strip())
                                # try several common formats (including '%d %m %Y')
                                for fmt in ('%Y-%m-%d', '%d-%m-%Y', '%d/%m/%Y', '%Y/%m/%d', '%d %m %Y', '%d %b %Y', '%d %B %Y'):
                                    try:
                                        parsed = datetime.strptime(s, fmt)
                                        break
                                    except Exception:
                                        continue
                            if parsed is None:
                                # couldn't parse -> drop the value so defaulting logic runs
                                data_map.pop(k, None)
                                continue
                            # For pure DATE columns, store a date object
                            if 'date' in ft and 'time' not in ft and 'datetime' not in ft:
                                data_map[k] = parsed.date()
                            else:
                                # For datetime/timestamp, make aware if naive
                                if parsed.tzinfo is None:
                                    parsed = timezone.make_aware(parsed)
                                data_map[k] = parsed
                        except Exception:
                            # On any error, remove the problematic key so default handling applies
                            data_map.pop(k, None)

                    # Prevent accidental override/insertion of primary key 'id'
                    if 'id' in data_map:
                        # remove incoming id so INSERT will use auto-increment
                        data_map.pop('id', None)

                    # Ensure certain fields have sensible defaults when not provided by frontend
                    # For numeric columns (int/decimal) use 0/0.0; otherwise fall back to '-'.
                    field_type_map = {row[0]: row[1].lower() for row in cols_meta}
                    for _f in ('admin_rate', 'tlo', 'life_insurance'):
                        if _f in cols_lookup and cols_lookup[_f] not in data_map:
                            ftype = field_type_map.get(_f, '')
                            if any(t in ftype for t in ('int', 'decimal', 'float', 'double')):
                                if 'decimal' in ftype or 'float' in ftype or 'double' in ftype:
                                    data_map[cols_lookup[_f]] = 0.0
                                else:
                                    data_map[cols_lookup[_f]] = 0
                            else:
                                data_map[cols_lookup[_f]] = '-'

                    # If some NOT NULL columns have no value provided, fill with safe defaults
                    # This prevents MySQL error 1364 when INSERT omits required columns.
                    # cols_meta rows: (Field, Type, Null, Key, Default, Extra)
                    for col_row in cols_meta:
                        field_name = col_row[0]
                        field_type = col_row[1].lower()
                        is_nullable = col_row[2]
                        default_val = col_row[4]
                        # Skip auto-managed or already-provided fields. Also skip created_by/created_at/update_at
                        if field_name in data_map:
                            continue
                        # do not auto-fill the primary key column; let AUTO_INCREMENT or synthesized PK handle it
                        if field_name in ('id', 'created_by', 'created_at', 'update_at') or (pk_col_name and field_name == pk_col_name):
                            continue

                        # If column is NOT NULL and has no default, provide a safe fallback
                        if is_nullable == 'NO' and default_val is None:
                            if any(t in field_type for t in ('int', 'decimal', 'float', 'double')):
                                data_map[field_name] = 0
                            elif any(t in field_type for t in ('date', 'timestamp', 'datetime')):
                                # use current timestamp for date/time types
                                data_map[field_name] = timezone.now()
                            else:
                                # default to '-' for textual/other types when frontend didn't include the field
                                data_map[field_name] = '-'

                    # Prepare columns and values for upsert
                    # Format date and numeric fields according to requirements
                    # NOTE: Formatting (date/number -> localized strings) is intentionally
                    # NOT applied here. Store raw values in DB and perform localization
                    # only when rendering/exporting (e.g., DOCX). This preserves numeric
                    # integrity (integers/decimals) and lets clients format as needed.
                    # Debug: inspect data_map after populating defaults
                    # debug removed
                    existing_check_sql = "SELECT contract_number FROM bl_agreement WHERE contract_number=%s"
                    exists = None
                    if contract_number:
                        cursor.execute(existing_check_sql, [contract_number])
                        exists = cursor.fetchone()

                    # If frontend explicitly requests edit-only or create-only behavior, enforce accordingly
                    edit_only = bool(data.get('edit_only'))
                    create_only = bool(data.get('create_only'))

                    if exists:
                        # If create_only requested but row exists, return error instead of updating
                        if create_only:
                            return Response({'error': 'Record already exists (create_only specified)'}, status=status.HTTP_400_BAD_REQUEST)
                        # UPDATE: add update_at timestamp if column exists
                        if 'update_at' in cols_info:
                            # do not overwrite created_by/created_at on update
                            data_map['update_at'] = now

                        # ensure we don't accidentally change created_by/created_at on update
                        data_map.pop('created_by', None)
                        data_map.pop('created_at', None)

                        # UPDATE
                        set_cols = []
                        params = []
                        for col, val in data_map.items():
                            if col == 'contract_number':
                                continue
                            set_cols.append(f"{col}=%s")
                            params.append(val)
                        if set_cols:
                            sql = f"UPDATE bl_agreement SET {', '.join(set_cols)} WHERE contract_number=%s"
                            params.append(contract_number)
                            cursor.execute(sql, params)
                    else:
                        # If edit_only requested but row doesn't exist, return error instead of inserting
                        if edit_only:
                            return Response({'error': 'Record not found for update (edit_only specified)'}, status=status.HTTP_400_BAD_REQUEST)
                        # If create_only is false, proceed to insert; if create_only true we already handled exists case above
                        # INSERT: set created_by (prefer username) and created_at if available
                        if 'created_by' in cols_info and user_username:
                            data_map['created_by'] = user_username
                        if 'created_at' in cols_info:
                            data_map['created_at'] = now

                        # INSERT
                        # Synthesize PK for legacy schemas or remove explicit zero PK so
                        # AUTO_INCREMENT can assign a new id. This prevents Duplicate entry '0'.
                        try:
                            _ensure_synthesized_pk(cursor, cols_meta, data_map, 'bl_agreement')
                        except Exception:
                            pass
                        try:
                            pk_row = next((r for r in cols_meta if (r[3] or '').upper() == 'PRI'), None)
                            if pk_row:
                                pk_name = pk_row[0]
                                if pk_name in data_map and (data_map[pk_name] is None or str(data_map[pk_name]) in ('', '0')):
                                    data_map.pop(pk_name, None)
                        except Exception:
                            pass
                        # debug removed
                        cols = []
                        placeholders = []
                        params = []
                        for col, val in data_map.items():
                            cols.append(col)
                            placeholders.append('%s')
                            params.append(val)
                        sql = f"INSERT INTO bl_agreement ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
                        try:
                            # Debug: log SQL and params to diagnose Duplicate entry '0' issues
                            # debug removed
                            cursor.execute(sql, params)
                        except Exception as ie:
                            msg = str(ie)
                            # Handle duplicate-primary gracefully: try to perform an UPDATE when possible
                            if 'Duplicate entry' in msg and 'PRIMARY' in msg:
                                logger.warning('Duplicate primary on bl_agreement insert: %s', msg)
                                # debug removed
                                # If we have a contract_number, attempt to update the existing row instead
                                if contract_number:
                                    try:
                                        set_cols = []
                                        params2 = []
                                        for col, val in data_map.items():
                                            if col == 'contract_number':
                                                continue
                                            set_cols.append(f"{col}=%s")
                                            params2.append(val)
                                        if set_cols:
                                            update_sql = f"UPDATE bl_agreement SET {', '.join(set_cols)} WHERE contract_number=%s"
                                            params2.append(contract_number)
                                            cursor.execute(update_sql, params2)
                                    except Exception:
                                        logger.exception('Failed safe-update after duplicate insert into bl_agreement')
                                else:
                                    # No contract_number to target; log and continue without failing download
                                    logger.warning('Duplicate primary but no contract_number available; skipping insert')
                            else:
                                raise
                return Response({'message': 'Data berhasil disimpan'}, status=status.HTTP_200_OK)
            except Exception as e:
                return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def get(self, request):
        contract_number = request.query_params.get('contract_number', '').strip()
        # Diagnostic: log incoming query params to help debug format handling / 404s
        try:
            logs_dir = os.path.join(settings.BASE_DIR, 'logs')
            os.makedirs(logs_dir, exist_ok=True)
            req_log = os.path.join(logs_dir, 'bl_agreement_docx_requests.log')
            with open(req_log, 'a', encoding='utf-8') as rf:
                rf.write(f"[{timezone.now().isoformat()}] query_params={dict(request.query_params)}\n")
        except Exception:
            pass

        try:
            with connection.cursor() as cursor:
                if not contract_number:
                    # Return list of agreements from bl_agreement table
                    # Select commonly used columns for listing
                    cursor.execute(
                        "SELECT agreement_date, contract_number, name_of_debtor, nik_number_of_debtor, collateral_type, created_by FROM bl_agreement ORDER BY COALESCE(agreement_date, created_at) DESC"
                    )
                    cols = [c[0] for c in cursor.description] if cursor.description else []
                    rows = cursor.fetchall()
                    items = [dict(zip(cols, r)) for r in rows]
                    return Response({'agreements': _normalize_for_json(items)}, status=status.HTTP_200_OK)

                # If contract_number provided, try to fetch from bl_agreement first
                cursor.execute(
                    "SELECT * FROM bl_agreement WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                    [contract_number]
                )
                cols = [col[0] for col in cursor.description] if cursor.description else []
                bl_row = cursor.fetchone()
                bl_data = dict(zip(cols, bl_row)) if bl_row else None

                # Attempt to get collateral data from bl_collateral ONLY
                cursor.execute(
                    "SELECT * FROM bl_collateral WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                    [contract_number]
                )
                collateral_columns = [col[0] for col in cursor.description] if cursor.description else []
                bl_collateral_row = cursor.fetchone()
                collateral_data = dict(zip(collateral_columns, bl_collateral_row)) if bl_collateral_row else None

                if bl_data:
                    # Enrich bl_data with related info if certain header fields are missing
                    try:
                        # Alias Name_of_director for frontend consistency
                        if 'name_of_director' in bl_data and 'Name_of_director' not in bl_data:
                            bl_data['Name_of_director'] = bl_data.get('name_of_director')

                        # If phone_number_of_lolc missing but director name present, try director table
                        if (not bl_data.get('phone_number_of_lolc')) and bl_data.get('name_of_director'):
                            cursor.execute("SELECT phone_number_of_lolc FROM director WHERE name_of_director=%s LIMIT 1", [bl_data.get('name_of_director')])
                            row = cursor.fetchone()
                            if row and row[0]:
                                bl_data['phone_number_of_lolc'] = row[0]

                        # If sp3 info missing, try bl_sp3 table for the latest entry
                        if (not bl_data.get('sp3_number') or not bl_data.get('sp3_date')):
                            cursor.execute("SELECT sp3_number, sp3_date FROM bl_sp3 WHERE LOWER(contract_number)=LOWER(%s) ORDER BY COALESCE(sp3_date, created_at) DESC LIMIT 1", [contract_number])
                            sp3row = cursor.fetchone()
                            if sp3row:
                                if not bl_data.get('sp3_number') and sp3row[0]:
                                    bl_data['sp3_number'] = sp3row[0]
                                if not bl_data.get('sp3_date') and sp3row[1]:
                                    bl_data['sp3_date'] = sp3row[1]
                    except Exception:
                        # Non-fatal enrichment failure; continue returning base bl_data
                        pass

                    return Response({'debtor': _normalize_for_json(bl_data), 'collateral': _normalize_for_json(collateral_data)}, status=status.HTTP_200_OK)

                # Fallback to legacy `contract` table if bl_agreement not present
                cursor.execute(
                    "SELECT * FROM contract WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                    [contract_number]
                )
                columns = [col[0] for col in cursor.description] if cursor.description else []
                contract_row = cursor.fetchone()
                debtor_data = dict(zip(columns, contract_row)) if contract_row else None

                # If neither table has the contract, return 404
                if not debtor_data and not collateral_data:
                    return Response({
                        'debtor': None,
                        'collateral': None,
                        'message': 'Data tidak ditemukan untuk nomor kontrak ini'
                    }, status=status.HTTP_404_NOT_FOUND)

                return Response({'debtor': _normalize_for_json(debtor_data), 'collateral': _normalize_for_json(collateral_data)}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


        class UVAgreementView(APIView):
            """
            Mirrors BLAgreementView behavior but operates on uv_agreement and uv_collateral tables.
            This keeps BL endpoints untouched while providing the same save/list/get semantics
            for UV agreements.
            """
            permission_classes = [IsAuthenticated]

            def post(self, request):
                data = request.data
                contract_number = data.get('contract_number')
                branch_id = data.get('branch_id')
                director = data.get('director')
                bm_data = data.get('bm_data', {})
                contract_data = data.get('contract_data', {})
                collateral_data = data.get('collateral_data', {})
                header_fields = data.get('header_fields', {})

                try:
                    with connection.cursor() as cursor:
                        cursor.execute("SHOW COLUMNS FROM uv_agreement")
                        cols_meta = cursor.fetchall()
                        cols_info = [row[0] for row in cols_meta]
                        cols_lookup = {c.lower(): c for c in cols_info}

                        data_map = {}
                        if contract_number:
                            data_map['contract_number'] = contract_number

                        client_created_by = data.get('created_by')
                        if client_created_by and 'created_by' in cols_info:
                            data_map['created_by'] = client_created_by

                        user_full_name = None
                        user_username = None
                        if getattr(request, 'user', None) and getattr(request.user, 'is_authenticated', False):
                            user_full_name = getattr(request.user, 'full_name', None)
                            user_username = getattr(request.user, 'username', None)
                            if (not user_full_name or not user_username) and getattr(request.user, 'id', None):
                                try:
                                    u = User.objects.filter(pk=request.user.id).first()
                                    if u:
                                        user_full_name = user_full_name or getattr(u, 'full_name', None)
                                        user_username = user_username or getattr(u, 'username', None)
                                except Exception:
                                    pass
                        else:
                            auth_header = request.META.get('HTTP_AUTHORIZATION', '')
                            if auth_header and auth_header.startswith('Bearer '):
                                token = auth_header.split(' ', 1)[1].strip()
                                try:
                                    payload = AccessToken(token)
                                    user_id = payload.get('user_id') or payload.get('uid') or payload.get('id')
                                    if user_id:
                                        try:
                                            u = User.objects.filter(pk=user_id).first()
                                            if u:
                                                user_full_name = user_full_name or getattr(u, 'full_name', None)
                                                user_username = user_username or getattr(u, 'username', None)
                                        except Exception:
                                            pass
                                except TokenError:
                                    pass

                        user_username = user_username or (user_full_name and user_full_name.replace(' ', '').lower()) or 'anonymous'
                        now = timezone.now()

                        if branch_id is not None and 'branch_id' in cols_lookup:
                            data_map[cols_lookup['branch_id']] = branch_id
                        if director is not None:
                            if 'name_of_director' in cols_lookup:
                                data_map[cols_lookup['name_of_director']] = director
                            elif 'director' in cols_lookup:
                                data_map[cols_lookup['director']] = director

                        branch_data = data.get('branch_data', {})
                        for src in (bm_data, branch_data, contract_data, collateral_data, header_fields):
                            if not isinstance(src, dict):
                                continue
                            for k, v in src.items():
                                key = str(k).lower()
                                if key == 'name_of_director':
                                    key = 'name_of_director'
                                if key in cols_lookup:
                                    data_map[cols_lookup[key]] = v

                        from datetime import datetime, date as _date
                        field_type_map = {row[0]: row[1].lower() for row in cols_meta}
                        for k in list(data_map.keys()):
                            ft = field_type_map.get(k, '')
                            if not any(t in ft for t in ('date', 'timestamp', 'datetime')):
                                continue
                            val = data_map.get(k)
                            if val is None or (isinstance(val, str) and val.strip() == ''):
                                data_map.pop(k, None)
                                continue
                            if isinstance(val, (_date, datetime)):
                                continue
                            s = str(val).strip()
                            parsed = None
                            try:
                                try:
                                    parsed = datetime.fromisoformat(s)
                                except Exception:
                                    import re
                                    if re.match(r'^\s*\d{1,2}\s+\d{1,2}\s+\d{4}\s*$', s):
                                        s = re.sub(r'\s+', '/', s.strip())
                                    for fmt in ('%Y-%m-%d', '%d-%m-%Y', '%d/%m/%Y', '%Y/%m/%d', '%d %m %Y', '%d %b %Y', '%d %B %Y'):
                                        try:
                                            parsed = datetime.strptime(s, fmt)
                                            break
                                        except Exception:
                                            continue
                                if parsed is None:
                                    data_map.pop(k, None)
                                    continue
                                if 'date' in ft and 'time' not in ft and 'datetime' not in ft:
                                    data_map[k] = parsed.date()
                                else:
                                    if parsed.tzinfo is None:
                                        parsed = timezone.make_aware(parsed)
                                    data_map[k] = parsed
                            except Exception:
                                data_map.pop(k, None)

                        if 'id' in data_map:
                            data_map.pop('id', None)

                        field_type_map = {row[0]: row[1].lower() for row in cols_meta}
                        for _f in ('admin_rate', 'tlo', 'life_insurance'):
                            if _f in cols_lookup and cols_lookup[_f] not in data_map:
                                ftype = field_type_map.get(_f, '')
                                if any(t in ftype for t in ('int', 'decimal', 'float', 'double')):
                                    if 'decimal' in ftype or 'float' in ftype or 'double' in ftype:
                                        data_map[cols_lookup[_f]] = 0.0
                                    else:
                                        data_map[cols_lookup[_f]] = 0
                                else:
                                    data_map[cols_lookup[_f]] = '-'

                        for col_row in cols_meta:
                            field_name = col_row[0]
                            field_type = col_row[1].lower()
                            is_nullable = col_row[2]
                            default_val = col_row[4]
                            if field_name in data_map:
                                continue
                            if field_name in ('id', 'created_by', 'created_at', 'update_at'):
                                continue
                            if is_nullable == 'NO' and default_val is None:
                                if any(t in field_type for t in ('int', 'decimal', 'float', 'double')):
                                    data_map[field_name] = 0
                                elif any(t in field_type for t in ('date', 'timestamp', 'datetime')):
                                    data_map[field_name] = timezone.now()
                                else:
                                    data_map[field_name] = '-'

                        existing_check_sql = "SELECT contract_number FROM uv_agreement WHERE contract_number=%s"
                        exists = None
                        if contract_number:
                            cursor.execute(existing_check_sql, [contract_number])
                            exists = cursor.fetchone()

                        edit_only = bool(data.get('edit_only'))
                        create_only = bool(data.get('create_only'))

                        if exists:
                            if create_only:
                                return Response({'error': 'Record already exists (create_only specified)'}, status=status.HTTP_400_BAD_REQUEST)
                            if 'update_at' in cols_info:
                                data_map['update_at'] = now
                            data_map.pop('created_by', None)
                            data_map.pop('created_at', None)
                            set_cols = []
                            params = []
                            for col, val in data_map.items():
                                if col == 'contract_number':
                                    continue
                                set_cols.append(f"{col}=%s")
                                params.append(val)
                            if set_cols:
                                sql = f"UPDATE uv_agreement SET {', '.join(set_cols)} WHERE contract_number=%s"
                                params.append(contract_number)
                                cursor.execute(sql, params)
                        else:
                            if edit_only:
                                return Response({'error': 'Record not found for update (edit_only specified)'}, status=status.HTTP_400_BAD_REQUEST)
                            if 'created_by' in cols_info and user_username:
                                data_map['created_by'] = user_username
                            if 'created_at' in cols_info:
                                data_map['created_at'] = now
                            cols = []
                            placeholders = []
                            params = []
                            for col, val in data_map.items():
                                cols.append(col)
                                placeholders.append('%s')
                                params.append(val)
                            sql = f"INSERT INTO uv_agreement ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
                            cursor.execute(sql, params)

                    return Response({'message': 'Data UV berhasil disimpan'}, status=status.HTTP_200_OK)
                except Exception as e:
                    return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            def get(self, request):
                contract_number = request.query_params.get('contract_number', '').strip()
                try:
                    with connection.cursor() as cursor:
                        if not contract_number:
                            cursor.execute(
                                "SELECT agreement_date, contract_number, name_of_debtor, nik_number_of_debtor, vehicle_type, collateral_type, created_by FROM uv_agreement ORDER BY COALESCE(agreement_date, created_at) DESC"
                            )
                            cols = [c[0] for c in cursor.description] if cursor.description else []
                            rows = cursor.fetchall()
                            items = [dict(zip(cols, r)) for r in rows]
                            return Response({'agreements': items}, status=status.HTTP_200_OK)

                        cursor.execute(
                            "SELECT * FROM uv_agreement WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                            [contract_number]
                        )
                        cols = [col[0] for col in cursor.description] if cursor.description else []
                        uv_row = cursor.fetchone()
                        uv_data = dict(zip(cols, uv_row)) if uv_row else None

                        cursor.execute(
                            "SELECT * FROM uv_collateral WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                            [contract_number]
                        )
                        collateral_columns = [col[0] for col in cursor.description] if cursor.description else []
                        uv_collateral_row = cursor.fetchone()
                        collateral_data = dict(zip(collateral_columns, uv_collateral_row)) if uv_collateral_row else None

                        if uv_data:
                            try:
                                if 'name_of_director' in uv_data and 'Name_of_director' not in uv_data:
                                    uv_data['Name_of_director'] = uv_data.get('name_of_director')
                            except Exception:
                                pass
                            return Response({'debtor': uv_data, 'collateral': collateral_data}, status=status.HTTP_200_OK)

                        cursor.execute(
                            "SELECT * FROM contract WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                            [contract_number]
                        )
                        columns = [col[0] for col in cursor.description] if cursor.description else []
                        contract_row = cursor.fetchone()
                        debtor_data = dict(zip(columns, contract_row)) if contract_row else None

                        if not debtor_data and not collateral_data:
                            return Response({'debtor': None, 'collateral': None, 'message': 'Data tidak ditemukan untuk nomor kontrak ini'}, status=status.HTTP_404_NOT_FOUND)

                        return Response({'debtor': debtor_data, 'collateral': collateral_data}, status=status.HTTP_200_OK)
                except Exception as e:
                    return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


        class UVCollateralCreateView(APIView):
            """Create or update UV collateral rows (mirrors BL collateral handler semantics)."""
            permission_classes = [IsAuthenticated]

            def post(self, request):
                data = request.data
                contract_number = data.get('contract_number')
                collateral = data.get('collateral', {})
                if not contract_number:
                    return Response({'error': 'contract_number is required'}, status=status.HTTP_400_BAD_REQUEST)

                try:
                    with connection.cursor() as cursor:
                        cursor.execute("SHOW COLUMNS FROM uv_collateral")
                        cols_meta = cursor.fetchall()
                        cols_info = [r[0] for r in cols_meta]
                        cols_lookup = {c.lower(): c for c in cols_info}

                        try:
                            data_map = {'contract_number': str(contract_number).upper()}
                        except Exception:
                            data_map = {'contract_number': contract_number}
                        if isinstance(collateral, dict):
                            for k, v in collateral.items():
                                key = str(k).lower()
                                if key in cols_lookup:
                                    data_map[cols_lookup[key]] = v

                        for col_row in cols_meta:
                            field_name = col_row[0]
                            field_type = col_row[1].lower()
                            is_nullable = col_row[2]
                            default_val = col_row[4]
                            if field_name in data_map:
                                continue
                            if field_name in ('id', 'created_by', 'created_at', 'update_at'):
                                continue
                            if is_nullable == 'NO' and default_val is None:
                                if any(t in field_type for t in ('int', 'decimal', 'float', 'double')):
                                    data_map[field_name] = 0
                                elif any(t in field_type for t in ('date', 'timestamp', 'datetime')):
                                    data_map[field_name] = timezone.now()
                                else:
                                    data_map[field_name] = '-'

                        # Prevent duplicate uv_collateral rows for same contract_number (case-insensitive)
                        cursor.execute("SELECT 1 FROM uv_collateral WHERE LOWER(contract_number)=LOWER(%s) LIMIT 1", [contract_number])
                        if cursor.fetchone():
                            return Response({'error': 'Duplicate contract_number: collateral already exists'}, status=status.HTTP_409_CONFLICT)

                        cols = []
                        placeholders = []
                        params = []
                        for col, val in data_map.items():
                            cols.append(col)
                            placeholders.append('%s')
                            params.append(val)
                        sql = f"INSERT INTO uv_collateral ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
                        try:
                            cursor.execute(sql, params)
                        except IntegrityError:
                            return Response({'error': 'Duplicate contract_number: collateral already exists'}, status=status.HTTP_409_CONFLICT)

                    return Response({'message': 'UV collateral saved'}, status=status.HTTP_200_OK)
                except Exception as e:
                    return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class BLAgreementContractListView(APIView):
    """
    API endpoint untuk mengambil daftar contract numbers dari tabel bl_agreement dan bl_collateral.
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            with connection.cursor() as cursor:
                # Return richer contract objects. Combine possible contract sources
                # `contract`, `bl_agreement`, and `bl_collateral` using UNION so frontend
                # receives objects with requested fields even when some tables only
                # contain the contract_number.
                cursor.execute("""
                    SELECT DISTINCT contract_id, contract_number, name_of_debtor, nik_number_of_debtor, loan_amount
                    FROM (
                        SELECT id AS contract_id, contract_number, COALESCE(name_of_debtor, '') AS name_of_debtor,
                               COALESCE(nik_number_of_debtor, '') AS nik_number_of_debtor,
                               COALESCE(loan_amount, '') AS loan_amount
                        FROM contract
                        UNION
                        SELECT NULL AS contract_id, contract_number, COALESCE(name_of_debtor, '') AS name_of_debtor,
                               COALESCE(nik_number_of_debtor, '') AS nik_number_of_debtor,
                               COALESCE(loan_amount, '') AS loan_amount
                        FROM bl_agreement
                        UNION
                        SELECT NULL AS contract_id, contract_number, '' AS name_of_debtor, '' AS nik_number_of_debtor, '' AS loan_amount
                        FROM bl_collateral
                    ) t
                    ORDER BY contract_number
                """)

                rows = cursor.fetchall()
                contracts = []
                for r in rows:
                    contracts.append({
                        'contract_id': r[0],
                        'contract_number': r[1],
                        'name_of_debtor': r[2],
                        'nik_number_of_debtor': r[3],
                        'loan_amount': r[4],
                    })

                return Response({'contracts': contracts}, status=status.HTTP_200_OK)
        
        except Exception as e:
            return Response(
                {'error': f'Error: {str(e)}'}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class ContractCreateView(APIView):
    """
    API endpoint to create a row in the legacy `contract` table.
    Frontend expects POST to `/api/contracts/` with contract fields (excluding *_in_word/_by_word).
    This endpoint will insert only columns that exist in the `contract` table and will
    populate `created_by`, `created_at`, `updated_at` when available.
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = request.data or {}
        try:
            with connection.cursor() as cursor:
                # discover contract table columns (map lowercase -> actual name)
                cursor.execute("SHOW COLUMNS FROM contract")
                cols_meta = cursor.fetchall()
                cols_info = [row[0] for row in cols_meta]
                cols_lookup = {c.lower(): c for c in cols_info}

                # build data map only for columns that exist (use actual column names)
                data_map = {}
                for k, v in data.items():
                    key = str(k).lower()
                    # skip helper fields
                    if key.endswith('_in_word') or key.endswith('_by_word'):
                        continue
                    if key in cols_lookup:
                        data_map[cols_lookup[key]] = v

                # Remove any client-supplied timestamp fields (may be ISO strings with Z)
                # and set server-side timestamps in DB-friendly format.
                if 'created_at' in data_map:
                    data_map.pop('created_at', None)
                if 'updated_at' in data_map:
                    data_map.pop('updated_at', None)

                # Ensure certain optional string fields are stored with a visible
                # placeholder when the frontend leaves them empty (requested behavior).
                # Use twelve underscores as the placeholder so it's easy to spot.
                placeholder = '______________'
                try:
                    vak = cols_lookup.get('virtual_account_number')
                    if vak:
                        val = data_map.get(vak)
                        if val is None or (isinstance(val, str) and val.strip() == ''):
                            data_map[vak] = placeholder
                except Exception:
                    pass
                try:
                    tpk = cols_lookup.get('topup_contract')
                    if tpk:
                        val2 = data_map.get(tpk)
                        if val2 is None or (isinstance(val2, str) and val2.strip() == ''):
                            data_map[tpk] = placeholder
                except Exception:
                    pass

                # set audit fields if available
                now = timezone.now()
                username = _resolve_username(request)
                if 'created_by' in cols_lookup and cols_lookup['created_by'] not in data_map:
                    data_map[cols_lookup['created_by']] = username or ''
                # Always set timestamps server-side to avoid timezone/format issues from client
                if 'created_at' in cols_lookup:
                    data_map[cols_lookup['created_at']] = now
                if 'updated_at' in cols_lookup:
                    data_map[cols_lookup['updated_at']] = now

                # Ensure admin_rate, tlo, life_insurance have sensible defaults when not provided
                # Use numeric 0 for numeric columns (int/decimal/float), otherwise '-' for text
                field_type_map = {row[0].lower(): row[1].lower() for row in cols_meta}
                for _f in ('admin_rate', 'tlo', 'life_insurance'):
                    if _f in cols_lookup and cols_lookup[_f] not in data_map:
                        ftype = field_type_map.get(_f, '')
                        if any(t in ftype for t in ('int', 'decimal', 'float', 'double')):
                            # decimal -> use 0.0, integers -> 0
                            if 'decimal' in ftype or 'float' in ftype or 'double' in ftype:
                                data_map[cols_lookup[_f]] = 0.0
                            else:
                                data_map[cols_lookup[_f]] = 0
                        else:
                            data_map[cols_lookup[_f]] = '-'

                # If some NOT NULL columns have no value provided, fill with safe defaults
                # This prevents MySQL error 1364 when INSERT omits required columns.
                # cols_meta rows: (Field, Type, Null, Key, Default, Extra)
                for col_row in cols_meta:
                    field_name = col_row[0]
                    field_type = col_row[1].lower()
                    is_nullable = col_row[2]
                    default_val = col_row[4]
                    # Skip auto-managed or already-provided fields. Also skip created_by/created_at/updated_at
                    if field_name in data_map:
                        continue
                    if field_name in ('id', 'created_by', 'created_at', 'updated_at'):
                        continue

                    if is_nullable == 'NO' and default_val is None:
                        if any(t in field_type for t in ('int', 'decimal', 'float', 'double')):
                            # decimal -> 0.0, integers -> 0
                            if 'decimal' in field_type or 'float' in field_type or 'double' in field_type:
                                data_map[field_name] = 0.0
                            else:
                                data_map[field_name] = 0
                        elif any(t in field_type for t in ('date', 'timestamp', 'datetime')):
                            data_map[field_name] = timezone.now()
                        else:
                            data_map[field_name] = '-'

                # Ensure primary key is synthesized for legacy PK columns (non-auto_increment)
                # Allow caller to request skipping server-side normalization (e.g., modal Add-Contract)
                skip_normalization = False
                try:
                    # backend accepts both snake_case and camelCase flags
                    if data.get('create_only') or data.get('createOnly'):
                        skip_normalization = True
                except Exception:
                    pass
                try:
                    if data.get('skip_normalization'):
                        skip_normalization = True
                except Exception:
                    pass

                # Previously we applied Title Case to certain fields. That behavior
                # is no longer desired. Preserve caller-provided casing by default.
                # If the caller explicitly requests skipping normalization (modal
                # create flows), uppercase all textual fields so modal saves persist
                # in UPPERCASE server-side.
                try:
                    if skip_normalization:
                        for _k in list(data_map.keys()):
                            # Skip dates and numeric columns
                            ft = field_type_map.get(_k.lower(), '')
                            if any(t in ft for t in ('date', 'timestamp', 'datetime')):
                                continue
                            if any(t in ft for t in ('int', 'decimal', 'float', 'double', 'numeric')):
                                continue
                            v = data_map.get(_k)
                            if isinstance(v, str):
                                data_map[_k] = v.strip().upper()
                except Exception:
                    pass

                _ensure_synthesized_pk(cursor, cols_meta, data_map, 'contract')
                # Explicit fallback: if contract_id column exists and not provided, synthesize now
                try:
                    if 'contract_id' in [c[0] for c in cols_meta] and 'contract_id' not in data_map:
                        cursor.execute("SELECT COALESCE(MAX(contract_id), 0) + 1 FROM contract")
                        r = cursor.fetchone()
                        if r and r[0] is not None:
                            data_map['contract_id'] = r[0]
                except Exception:
                    pass

                if not data_map:
                    return Response({'error': 'No valid contract fields provided'}, status=status.HTTP_400_BAD_REQUEST)

                # Prevent duplicate contract_number entries: check existing contract_number (case-insensitive)
                try:
                    contract_col = cols_lookup.get('contract_number')
                    if contract_col:
                        # value may be in data_map under actual column name or in original payload
                        contract_val = data_map.get(contract_col) or (data.get('contract_number') if isinstance(data, dict) else None)
                        if contract_val:
                            try:
                                # Only force uppercase when normalization is enabled.
                                if not skip_normalization:
                                    data_map[contract_col] = str(contract_val).upper()
                                else:
                                    data_map[contract_col] = contract_val
                            except Exception:
                                data_map[contract_col] = contract_val
                            cursor.execute("SELECT 1 FROM contract WHERE LOWER(contract_number)=LOWER(%s) LIMIT 1", [data_map[contract_col]])
                            if cursor.fetchone():
                                return Response({'error': 'Duplicate contract_number: contract already exists'}, status=status.HTTP_409_CONFLICT)
                except Exception:
                    # don't block insert if uniqueness check fails for any reason; fallback to DB constraint handling
                    pass

                cols = []
                placeholders = []
                params = []
                for col, val in data_map.items():
                    cols.append(col)
                    placeholders.append('%s')
                    params.append(val)

                sql = f"INSERT INTO contract ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
                try:
                    cursor.execute(sql, params)
                except Exception as e:
                    # If duplicate PK due to legacy default PK value, try to synthesize PK and retry
                    msg = str(e.args[0]) if hasattr(e, 'args') and e.args else str(e)
                    if '1062' in msg or 'Duplicate entry' in msg:
                        try:
                            # detect primary key column from cols_meta
                            pk_row = next((r for r in cols_meta if (r[3] or '').upper() == 'PRI'), None)
                            if pk_row:
                                pk_name = pk_row[0]
                            else:
                                # fallback to common names
                                pk_name = next((c for c in ('contract_id','id') if c in [r[0] for r in cols_meta]), None)
                            if pk_name:
                                cursor.execute(f'SELECT COALESCE(MAX({pk_name}), 0) + 1 FROM contract')
                                new_id_row = cursor.fetchone()
                                new_id = new_id_row[0] if new_id_row else None
                                if new_id is not None:
                                    if pk_name not in data_map:
                                        data_map[pk_name] = new_id
                                    cols = []
                                    placeholders = []
                                    params = []
                                    for col, val in data_map.items():
                                        cols.append(col)
                                        placeholders.append('%s')
                                        params.append(val)
                                    sql2 = f"INSERT INTO contract ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
                                    cursor.execute(sql2, params)
                                    # success on retry
                                    pass
                                else:
                                    raise
                            else:
                                raise
                        except Exception:
                            # re-raise original for visibility
                            raise
                    else:
                        raise

            return Response({'message': 'Contract saved'}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('Failed to save contract')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def get(self, request):
        """
        Allow frontend to lookup a contract row by contract_number using GET /api/contracts/?contract_number=...
        Returns JSON: { "contract": { ... } }
        """
        contract_number = request.GET.get('contract_number') or request.query_params.get('contract_number')
        if not contract_number:
            return Response({'error': 'contract_number query parameter required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM contract")
                cols_meta = cursor.fetchall()
                cols_info = [row[0] for row in cols_meta]
                # Build SELECT with explicit columns to preserve order
                sql = f"SELECT {', '.join(cols_info)} FROM contract WHERE contract_number = %s LIMIT 1"
                cursor.execute(sql, [contract_number])
                row = cursor.fetchone()
                if not row:
                    return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)
                row_dict = {cols_info[i]: row[i] for i in range(len(cols_info))}
                return Response({'contract': row_dict}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('Contract lookup failed')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ContractCRUDView(APIView):
    """Detail endpoint for contract mutations: PATCH and DELETE.
    Keep create/lookup on /api/contracts/ untouched.
    """
    permission_classes = [IsAuthenticated]

    def _detect_pk_column(self, cursor):
        cursor.execute("SHOW COLUMNS FROM contract")
        cols_meta = cursor.fetchall()
        cols = [r[0] for r in cols_meta]
        for candidate in ('contract_id', 'id', 'id_contract'):
            if candidate in cols:
                return candidate, cols_meta
        # fallback to first column
        return cols[0] if cols else None, cols_meta

    def patch(self, request, pk=None):
        if not pk:
            return Response({'error': 'Missing contract id'}, status=status.HTTP_400_BAD_REQUEST)
        data = request.data or {}
        try:
            with connection.cursor() as cursor:
                pk_col, cols_meta = self._detect_pk_column(cursor)
                if not pk_col:
                    return Response({'error': 'Contract table has no detectable PK'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                # map lowercase -> actual column names
                cols_info = [r[0] for r in cols_meta]
                cols_lookup = {c.lower(): c for c in cols_info}

                updates = []
                params = []
                for k, v in data.items():
                    key = str(k).lower()
                    if key in cols_lookup and cols_lookup[key] != pk_col:
                        updates.append(f"{cols_lookup[key]}=%s")
                        params.append(v)

                if not updates:
                    return Response({'error': 'No updatable fields'}, status=status.HTTP_400_BAD_REQUEST)

                # set updated_at if available
                if 'updated_at' in cols_lookup and 'updated_at' not in [u.split('=')[0] for u in updates]:
                    updates.append('updated_at=%s')
                    params.append(timezone.now())

                params.append(pk)
                sql = f"UPDATE contract SET {', '.join(updates)} WHERE {pk_col}=%s"
                cursor.execute(sql, params)
            return Response({'status': 'updated'}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('Contract update failed')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def delete(self, request, pk=None):
        if not pk:
            return Response({'error': 'Missing contract id'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                pk_col, cols_meta = self._detect_pk_column(cursor)
                if not pk_col:
                    return Response({'error': 'Contract table has no detectable PK'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                cols = [r[0] for r in cols_meta]
                # try soft-delete
                if 'is_active' in cols:
                    cursor.execute(f"UPDATE contract SET is_active=0 WHERE {pk_col}=%s", [pk])
                else:
                    cursor.execute(f"DELETE FROM contract WHERE {pk_col}=%s", [pk])
            return Response({'status': 'deleted'}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('Contract delete failed')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ContractLookupView(APIView):
    """
    Dedicated lookup endpoint returning a contract row for a given contract_number.
    This is used by frontend create-modal lookups when GET on /api/contracts/ is not desired.
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        contract_number = request.query_params.get('contract_number')
        if not contract_number:
            return Response({'error': 'contract_number required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM contract")
                cols_meta = cursor.fetchall()
                cols_info = [row[0] for row in cols_meta]
                sql = f"SELECT {', '.join(cols_info)} FROM contract WHERE LOWER(contract_number)=LOWER(%s) LIMIT 1"
                cursor.execute(sql, [contract_number])
                row = cursor.fetchone()
                if not row:
                    return Response({'error': 'Not found'}, status=status.HTTP_404_NOT_FOUND)
                row_dict = {cols_info[i]: row[i] for i in range(len(cols_info))}
                return Response({'contract': row_dict}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('ContractLookup failed')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ContractsListView(APIView):
    """
    New read-only endpoint that returns richer contract objects aggregated from
    `contract`, `bl_agreement`, and `bl_collateral` without changing existing
    behavior of `/api/contracts/` which is used for create/lookup.
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            with connection.cursor() as cursor:
                # detect primary/id-like column in contract table to avoid "Unknown column 'id'" errors
                cursor.execute("SHOW COLUMNS FROM contract")
                contract_cols_meta = cursor.fetchall()
                contract_cols = [r[0] for r in contract_cols_meta]
                id_col = None
                for candidate in ('id', 'contract_id', 'id_contract'):
                    if candidate in contract_cols:
                        id_col = candidate
                        break

                if id_col:
                    contract_id_select = f"{id_col} AS contract_id"
                else:
                    contract_id_select = "NULL AS contract_id"

                sql = f"""
                    SELECT DISTINCT contract_id, contract_number, name_of_debtor, nik_number_of_debtor, loan_amount
                    FROM (
                        SELECT {contract_id_select}, contract_number, COALESCE(name_of_debtor, '') AS name_of_debtor,
                               COALESCE(nik_number_of_debtor, '') AS nik_number_of_debtor,
                               COALESCE(loan_amount, '') AS loan_amount
                        FROM contract
                        UNION
                        SELECT NULL AS contract_id, contract_number, COALESCE(name_of_debtor, '') AS name_of_debtor,
                               COALESCE(nik_number_of_debtor, '') AS nik_number_of_debtor,
                               COALESCE(loan_amount, '') AS loan_amount
                        FROM bl_agreement
                        UNION
                        SELECT NULL AS contract_id, contract_number, '' AS name_of_debtor, '' AS nik_number_of_debtor, '' AS loan_amount
                        FROM bl_collateral
                    ) t
                    ORDER BY contract_number
                """

                cursor.execute(sql)
                rows = cursor.fetchall()
                contracts = []
                for r in rows:
                    contracts.append({
                        'contract_id': r[0],
                        'contract_number': r[1],
                        'name_of_debtor': r[2],
                        'nik_number_of_debtor': r[3],
                        'loan_amount': r[4],
                    })

                return Response({'contracts': contracts}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('Contracts list failed')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ContractsTableView(APIView):
    """Return rows directly from the `contract` table (no UNION, no duplicates).
    Fields returned: `contract_id` (or id), `contract_number`, `name_of_debtor`,
    `nik_number_of_debtor`, `loan_amount` when available in the table.
    """
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            with connection.cursor() as cursor:
                # discover columns from the contract table
                cursor.execute("SHOW COLUMNS FROM contract")
                cols_meta = cursor.fetchall()
                cols = [r[0] for r in cols_meta]

                # exclude timestamp-like columns from the returned payload
                excluded = ('created_at', 'updated_at')

                # decide id mapping: prefer contract_id, else alias id -> contract_id
                id_col = None
                for candidate in ('contract_id', 'id', 'id_contract'):
                    if candidate in cols:
                        id_col = candidate
                        break

                select_list = []
                if id_col == 'id':
                    select_list.append('id AS contract_id')
                elif id_col == 'contract_id':
                    select_list.append('contract_id')
                else:
                    select_list.append("NULL AS contract_id")

                # include all other columns (preserve DB order), skipping id alias and timestamps
                for c in cols:
                    if c in ('id', 'contract_id'):
                        continue
                    if c in excluded:
                        continue
                    select_list.append(c)

                sql = f"SELECT DISTINCT {', '.join(select_list)} FROM contract ORDER BY contract_number"
                cursor.execute(sql)
                rows = cursor.fetchall()

                # build returned column names list to send as metadata
                returned_cols = ['contract_id'] + [c for c in cols if c not in ('id', 'contract_id') and c not in excluded]

                contracts = []
                for row in rows:
                    obj = {}
                    for i, colname in enumerate(returned_cols):
                        obj[colname] = row[i]
                    contracts.append(obj)

                return Response({'contracts': contracts, 'columns': returned_cols}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('Contracts table failed')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class BLCollateralCreateView(APIView):
    """
    Endpoint to insert a bl_collateral row. Frontend will POST contract_number and collateral fields.
    """
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = request.data or {}
        try:
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM bl_collateral")
                cols_meta = cursor.fetchall()
                cols_info = [r[0] for r in cols_meta]

                data_map = {}
                for k, v in data.items():
                    key = str(k).lower()
                    if key.endswith('_in_word') or key.endswith('_by_word'):
                        continue
                    if key in cols_info:
                        data_map[key] = v

                # server-side timestamps and audit
                now = timezone.now()
                username = _resolve_username(request)
                if 'created_by' in cols_info and 'created_by' not in data_map:
                    data_map['created_by'] = username or ''
                if 'created_at' in cols_info:
                    data_map['created_at'] = now
                if 'update_at' in cols_info:
                    data_map['update_at'] = now

                if 'id' in data_map:
                    data_map.pop('id', None)

                if not data_map:
                    return Response({'error': 'No valid collateral fields provided'}, status=status.HTTP_400_BAD_REQUEST)

                # Prevent duplicate collateral rows for same contract_number
                contract_number = data.get('contract_number') or data_map.get('contract_number')
                try:
                    if contract_number:
                        cursor.execute("SELECT 1 FROM bl_collateral WHERE LOWER(contract_number)=LOWER(%s) LIMIT 1", [contract_number])
                        if cursor.fetchone():
                            return Response({'error': 'Duplicate contract_number: collateral already exists'}, status=status.HTTP_409_CONFLICT)
                except Exception:
                    # if uniqueness check fails for any reason, continue to attempt insert
                    pass

                cols = []
                placeholders = []
                params = []
                for col, val in data_map.items():
                    cols.append(col)
                    placeholders.append('%s')
                    params.append(val)

                sql = f"INSERT INTO bl_collateral ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
                cursor.execute(sql, params)

            return Response({'message': 'Collateral saved'}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('Failed to save bl_collateral')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def get(self, request):
        """Return list or columns for BL collateral. If contract_number provided, return matching rows."""
        contract_number = request.query_params.get('contract_number') or request.GET.get('contract_number')
        try:
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM bl_collateral")
                cols_meta = cursor.fetchall()
                cols_info = [r[0] for r in cols_meta]
                if not contract_number:
                    return Response({'collateral': [], 'columns': cols_info}, status=status.HTTP_200_OK)
                sql = f"SELECT {', '.join(cols_info)} FROM bl_collateral WHERE LOWER(contract_number)=LOWER(%s)"
                cursor.execute(sql, [contract_number])
                rows = cursor.fetchall()
                result = [dict(zip(cols_info, row)) for row in rows]
                return Response({'collateral': result, 'columns': cols_info}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('BLCollateral lookup failed')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class WhoAmIView(APIView):
    """Debug endpoint: return resolved full_name for the current request's token/user."""
    permission_classes = []

    def get(self, request):
        # Try request.user first
        full_name = None
        if getattr(request, 'user', None) and getattr(request.user, 'is_authenticated', False):
            full_name = getattr(request.user, 'full_name', None) or getattr(request.user, 'username', None)
            if (not full_name) and getattr(request.user, 'id', None):
                u = User.objects.filter(pk=request.user.id).first()
                if u:
                    full_name = getattr(u, 'full_name', None) or getattr(u, 'username', None)

        # Fallback: decode token
        if not full_name:
            auth_header = request.META.get('HTTP_AUTHORIZATION', '')
            if auth_header and auth_header.startswith('Bearer '):
                token = auth_header.split(' ', 1)[1].strip()
                try:
                    payload = AccessToken(token)
                    user_id = payload.get('user_id') or payload.get('uid') or payload.get('id')
                    if user_id:
                        u = User.objects.filter(pk=user_id).first()
                        if u:
                            full_name = getattr(u, 'full_name', None) or getattr(u, 'username', None)
                except Exception:
                    full_name = None

        # Also return username when available for frontend display
        username = None
        if getattr(request, 'user', None) and getattr(request.user, 'is_authenticated', False):
            username = getattr(request.user, 'username', None)
            if (not username) and getattr(request.user, 'id', None):
                u = User.objects.filter(pk=request.user.id).first()
                if u:
                    username = getattr(u, 'username', None)
        if not username:
            # try token payload
            auth_header = request.META.get('HTTP_AUTHORIZATION', '')
            if auth_header and auth_header.startswith('Bearer '):
                try:
                    payload = AccessToken(auth_header.split(' ', 1)[1].strip())
                    user_id = payload.get('user_id') or payload.get('uid') or payload.get('id')
                    if user_id:
                        u = User.objects.filter(pk=user_id).first()
                        if u:
                            username = getattr(u, 'username', None)
                except Exception:
                    username = None

        # Also include helper-resolved username for debugging
        try:
            helper_uname = _resolve_username(request) or ''
        except Exception:
            helper_uname = 'error'

        # Fetch role, branch_id, area_id, region_id for the current user
        user_role = None
        user_branch_id = None
        user_area_id = None
        user_region_id = None
        try:
            uid = None
            if getattr(request, 'user', None) and getattr(request.user, 'is_authenticated', False):
                uid = getattr(request.user, 'id', None)
            if not uid:
                auth_header = request.META.get('HTTP_AUTHORIZATION', '')
                if auth_header and auth_header.startswith('Bearer '):
                    try:
                        payload = AccessToken(auth_header.split(' ', 1)[1].strip())
                        uid = payload.get('user_id') or payload.get('uid') or payload.get('id')
                    except Exception:
                        pass
            if uid:
                with connection.cursor() as cursor:
                    cursor.execute('SELECT role, branch_id, area_id, region_id FROM auth_user WHERE id=%s', [uid])
                    row = cursor.fetchone()
                    if row:
                        user_role, user_branch_id, user_area_id, user_region_id = row
        except Exception:
            pass

        # Return single response including helper result for inspection
        return Response({
            'full_name': full_name or 'anonymous',
            'username': username or 'anonymous',
            'helper_resolved': helper_uname,
            'role': user_role,
            'branch_id': user_branch_id,
            'area_id': user_area_id,
            'region_id': user_region_id,
        }, status=status.HTTP_200_OK)


class BranchListView(APIView):
    """Return distinct branch cities (city_of_bm) from branch manager table."""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT DISTINCT city_of_bm FROM branch_manager ORDER BY city_of_bm")
                rows = [row[0] for row in cursor.fetchall()]
                return Response({'branches': rows}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class RegionListView(APIView):
    """Return list of regions with id and name."""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            with connection.cursor() as cursor:
                # include code column so frontend can show region code
                cursor.execute("SELECT id, region_name, code FROM regions WHERE is_active=1 ORDER BY region_name")
                cols = ['id', 'name', 'code']
                rows = [dict(zip(cols, row)) for row in cursor.fetchall()]
                return Response({'regions': rows}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request):
        """Create a new region record."""
        name = request.data.get('name') or request.data.get('region_name')
        code = request.data.get('code')
        if not name:
            return Response({'error': 'Missing name'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                cursor.execute("INSERT INTO regions (region_name, code, is_active) VALUES (%s, %s, 1)", [name, code])
            return Response({'status': 'created'}, status=status.HTTP_201_CREATED)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def patch(self, request, pk=None):
        """Update an existing region by id (pk)."""
        if not pk:
            return Response({'error': 'Missing region id'}, status=status.HTTP_400_BAD_REQUEST)
        name = request.data.get('name') or request.data.get('region_name')
        code = request.data.get('code')
        if name is None and code is None:
            return Response({'error': 'No fields to update'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                # build update dynamically
                updates = []
                params = []
                if name is not None:
                    updates.append('region_name=%s')
                    params.append(name)
                if code is not None:
                    updates.append('code=%s')
                    params.append(code)
                params.append(pk)
                sql = f"UPDATE regions SET {', '.join(updates)} WHERE id=%s"
                cursor.execute(sql, params)
            return Response({'status': 'updated'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def delete(self, request, pk=None):
        """Soft-delete a region by marking is_active=0."""
        if not pk:
            return Response({'error': 'Missing region id'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                cursor.execute("UPDATE regions SET is_active=0 WHERE id=%s", [pk])
            return Response({'status': 'deleted'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class AreaListView(APIView):
    """Return list of areas with id, name, and region_id."""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        region_id = request.query_params.get('region_id')
        try:
            with connection.cursor() as cursor:
                if region_id:
                    cursor.execute("SELECT id, name, region_id, code FROM areas WHERE region_id=%s AND is_active=1 ORDER BY name", [region_id])
                else:
                    cursor.execute("SELECT id, name, region_id, code FROM areas WHERE is_active=1 ORDER BY name")
                cols = ['id', 'name', 'region_id', 'code']
                rows = [dict(zip(cols, row)) for row in cursor.fetchall()]
                return Response({'areas': rows}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request):
        """Create a new area record."""
        name = request.data.get('name')
        region_id = request.data.get('region_id')
        code = request.data.get('code')
        if not name:
            return Response({'error': 'Missing name'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                cursor.execute("INSERT INTO areas (name, region_id, code, is_active) VALUES (%s, %s, %s, 1)", [name, region_id, code])
            return Response({'status': 'created'}, status=status.HTTP_201_CREATED)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def patch(self, request, pk=None):
        """Update an existing area by id (pk)."""
        if not pk:
            return Response({'error': 'Missing area id'}, status=status.HTTP_400_BAD_REQUEST)
        name = request.data.get('name')
        region_id = request.data.get('region_id')
        code = request.data.get('code')
        if name is None and region_id is None and code is None:
            return Response({'error': 'No fields to update'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                updates = []
                params = []
                if name is not None:
                    updates.append('name=%s')
                    params.append(name)
                if region_id is not None:
                    updates.append('region_id=%s')
                    params.append(region_id)
                if code is not None:
                    updates.append('code=%s')
                    params.append(code)
                params.append(pk)
                sql = f"UPDATE areas SET {', '.join(updates)} WHERE id=%s"
                cursor.execute(sql, params)
            return Response({'status': 'updated'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def delete(self, request, pk=None):
        """Soft-delete an area by marking is_active=0."""
        if not pk:
            return Response({'error': 'Missing area id'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                cursor.execute("UPDATE areas SET is_active=0 WHERE id=%s", [pk])
            return Response({'status': 'deleted'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class BranchListView(APIView):
    """Return list of branches with id, name, and area_id.

    This endpoint is intentionally left without authentication to avoid
    user-model related DB queries (some environments have a custom
    `auth_user` table missing standard fields). Frontend requires branch
    data to populate filters, so keep this public.
    """
    permission_classes = []

    def get(self, request):
        area_id = request.query_params.get('area_id')
        try:
            with connection.cursor() as cursor:
                # Use dynamic selection of branch and BM columns (b.* and bm.*)
                # This avoids referencing columns that may not exist in some databases
                if area_id:
                    cursor.execute(
                        "SELECT b.*, bm.* FROM branches b LEFT JOIN branch_manager bm ON b.bm_id = bm.bm_id WHERE b.area_id=%s AND b.is_active=1 ORDER BY b.name",
                        [area_id]
                    )
                else:
                    cursor.execute(
                        "SELECT b.*, bm.* FROM branches b LEFT JOIN branch_manager bm ON b.bm_id = bm.bm_id WHERE b.is_active=1 ORDER BY b.name"
                    )
                cols = [c[0] for c in cursor.description] if cursor.description else []
                rows = [dict(zip(cols, row)) for row in cursor.fetchall()]
                return Response({'branches': rows}, status=status.HTTP_200_OK)
        except Exception as e:
            # Return a clear error message for frontend and log server-side
            err_msg = str(e)
            return Response({'error': f'Error fetching branches: {err_msg}'}, status=status.HTTP_200_OK)
    
    def post(self, request):
        """Create a new branch record."""
        # accept either 'name' or 'branch_name'
        area_id = request.data.get('area_id')
        bm_id = request.data.get('bm_id')
        name = request.data.get('name') or request.data.get('branch_name')
        code = request.data.get('code')
        # accept phone number for branches (frontend may send 'phone_number_branch' or 'phone')
        phone = request.data.get('phone_number_branch') or request.data.get('phone') or ''
        street_name = request.data.get('street_name')
        subdistrict = request.data.get('subdistrict')
        district = request.data.get('district')
        city = request.data.get('city')
        province = request.data.get('province')
        if not name:
            return Response({'error': 'Missing branch name'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                # include phone_number_branch column; use empty string fallback to avoid NOT NULL errors
                cursor.execute(
                    "INSERT INTO branches (area_id, bm_id, name, code, street_name, subdistrict, district, city, province, phone_number_branch, is_active) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 1)",
                    [area_id, bm_id, name, code, street_name, subdistrict, district, city, province, phone]
                )
            return Response({'status': 'created'}, status=status.HTTP_201_CREATED)
        except Exception as e:
            return Response({'error': f'Error creating branch: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def patch(self, request, pk=None):
        """Update an existing branch by id (pk)."""
        if not pk:
            return Response({'error': 'Missing branch id'}, status=status.HTTP_400_BAD_REQUEST)
        # gather updatable fields
        area_id = request.data.get('area_id')
        bm_id = request.data.get('bm_id')
        name = request.data.get('name') or request.data.get('branch_name')
        code = request.data.get('code')
        phone = request.data.get('phone_number_branch') or request.data.get('phone')
        street_name = request.data.get('street_name')
        subdistrict = request.data.get('subdistrict')
        district = request.data.get('district')
        city = request.data.get('city')
        province = request.data.get('province')
        if all(v is None for v in [area_id, bm_id, name, code, street_name, subdistrict, district, city, province]):
            return Response({'error': 'No fields to update'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                updates = []
                params = []
                if area_id is not None:
                    updates.append('area_id=%s'); params.append(area_id)
                if bm_id is not None:
                    updates.append('bm_id=%s'); params.append(bm_id)
                if name is not None:
                    updates.append('name=%s'); params.append(name)
                if code is not None:
                    updates.append('code=%s'); params.append(code)
                if street_name is not None:
                    updates.append('street_name=%s'); params.append(street_name)
                if subdistrict is not None:
                    updates.append('subdistrict=%s'); params.append(subdistrict)
                if district is not None:
                    updates.append('district=%s'); params.append(district)
                if city is not None:
                    updates.append('city=%s'); params.append(city)
                if phone is not None:
                    updates.append('phone_number_branch=%s'); params.append(phone)
                if province is not None:
                    updates.append('province=%s'); params.append(province)
                params.append(pk)
                sql = f"UPDATE branches SET {', '.join(updates)} WHERE id=%s"
                cursor.execute(sql, params)
            return Response({'status': 'updated'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error updating branch: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def delete(self, request, pk=None):
        """Soft-delete a branch by marking is_active=0."""
        if not pk:
            return Response({'error': 'Missing branch id'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                cursor.execute("UPDATE branches SET is_active=0 WHERE id=%s", [pk])
            return Response({'status': 'deleted'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error deleting branch: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class BranchManagerByCityView(APIView):
    """Return branch manager (BM) data for a given city_of_bm."""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        bm_id = request.query_params.get('bm_id')
        city = request.query_params.get('city', '').strip()
        # If no parameters provided, return full list of branch managers
        if not bm_id and not city:
            try:
                with connection.cursor() as cursor:
                    # branch_manager table does not have is_active column in some schemas
                    # avoid filtering by is_active to prevent SQL error
                    cursor.execute("SELECT * FROM branch_manager ORDER BY name_of_bm")
                    cols = [c[0] for c in cursor.description] if cursor.description else []
                    rows = [dict(zip(cols, row)) for row in cursor.fetchall()]
                    return Response({'bm': rows}, status=status.HTTP_200_OK)
            except Exception as e:
                return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        try:
            with connection.cursor() as cursor:
                if bm_id:
                    cursor.execute("SELECT * FROM branch_manager WHERE bm_id = %s LIMIT 1", [bm_id])
                else:
                    # Fallback: try to match by name_of_bm if city provided (less reliable)
                    cursor.execute("SELECT * FROM branch_manager WHERE name_of_bm = %s LIMIT 1", [city])
                cols = [c[0] for c in cursor.description] if cursor.description else []
                row = cursor.fetchone()
                if not row:
                    return Response({'bm': None}, status=status.HTTP_200_OK)
                data = dict(zip(cols, row))
                return Response({'bm': data}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class BranchManagerCRUDView(APIView):
    """Separate CRUD API for branch_manager table to avoid touching legacy lookup view."""
    permission_classes = [IsAuthenticated]

    def post(self, request):
        bm_id = request.data.get('bm_id')
        branches_id = request.data.get('branches_id')
        name = request.data.get('name_of_bm') or request.data.get('name')
        nik = request.data.get('nik_number_of_bm')
        phone = request.data.get('phone_number_of_bm')
        place_birth = request.data.get('place_birth_of_bm') or request.data.get('place_of_birth_of_bm')
        date_birth = request.data.get('date_birth_of_bm') or request.data.get('date_of_birth_of_bm')
        street = request.data.get('street_name_of_bm') or request.data.get('street_of_bm') or request.data.get('street_name')
        subdistrict = request.data.get('subdistrict_of_bm') or request.data.get('subdistrict')
        district = request.data.get('district_of_bm') or request.data.get('district')
        city = request.data.get('city_of_bm') or request.data.get('city') or request.data.get('name')
        province = request.data.get('province_of_bm') or request.data.get('province')
        if not name:
            return Response({'error': 'Missing name_of_bm'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                if bm_id:
                    cursor.execute(
                        "INSERT INTO branch_manager (bm_id, branches_id, name_of_bm, place_birth_of_bm, date_birth_of_bm, nik_number_of_bm, phone_number_of_bm, street_name_of_bm, subdistrict_of_bm, district_of_bm, city_of_bm, province_of_bm) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                        [bm_id, branches_id, name, place_birth, date_birth, nik, phone, street, subdistrict, district, city, province]
                    )
                else:
                    cursor.execute(
                        "INSERT INTO branch_manager (branches_id, name_of_bm, place_birth_of_bm, date_birth_of_bm, nik_number_of_bm, phone_number_of_bm, street_name_of_bm, subdistrict_of_bm, district_of_bm, city_of_bm, province_of_bm) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                        [branches_id, name, place_birth, date_birth, nik, phone, street, subdistrict, district, city, province]
                    )
            return Response({'status': 'created'}, status=status.HTTP_201_CREATED)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def patch(self, request, pk=None):
        if not pk:
            return Response({'error': 'Missing bm id'}, status=status.HTTP_400_BAD_REQUEST)
        branches_id = request.data.get('branches_id')
        name = request.data.get('name_of_bm') or request.data.get('name')
        nik = request.data.get('nik_number_of_bm')
        phone = request.data.get('phone_number_of_bm')
        place_birth = request.data.get('place_birth_of_bm') or request.data.get('place_of_birth_of_bm')
        date_birth = request.data.get('date_birth_of_bm') or request.data.get('date_of_birth_of_bm')
        street = request.data.get('street_name_of_bm') or request.data.get('street_of_bm') or request.data.get('street_name')
        subdistrict = request.data.get('subdistrict_of_bm') or request.data.get('subdistrict')
        district = request.data.get('district_of_bm') or request.data.get('district')
        city = request.data.get('city_of_bm') or request.data.get('city') or request.data.get('name')
        province = request.data.get('province_of_bm') or request.data.get('province')
        if name is None and branches_id is None and nik is None and phone is None:
            return Response({'error': 'No fields to update'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                updates = []
                params = []
                if branches_id is not None:
                    updates.append('branches_id=%s'); params.append(branches_id)
                if name is not None:
                    updates.append('name_of_bm=%s'); params.append(name)
                if nik is not None:
                    updates.append('nik_number_of_bm=%s'); params.append(nik)
                if phone is not None:
                    updates.append('phone_number_of_bm=%s'); params.append(phone)
                if street is not None:
                    updates.append('street_name_of_bm=%s'); params.append(street)
                if place_birth is not None:
                    updates.append('place_birth_of_bm=%s'); params.append(place_birth)
                if date_birth is not None:
                    updates.append('date_birth_of_bm=%s'); params.append(date_birth)
                if subdistrict is not None:
                    updates.append('subdistrict_of_bm=%s'); params.append(subdistrict)
                if district is not None:
                    updates.append('district_of_bm=%s'); params.append(district)
                if city is not None:
                    updates.append('city_of_bm=%s'); params.append(city)
                if province is not None:
                    updates.append('province_of_bm=%s'); params.append(province)
                params.append(pk)
                sql = f"UPDATE branch_manager SET {', '.join(updates)} WHERE bm_id=%s"
                cursor.execute(sql, params)
            return Response({'status': 'updated'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def delete(self, request, pk=None):
        if not pk:
            return Response({'error': 'Missing bm id'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                try:
                    cursor.execute("UPDATE branch_manager SET is_active=0 WHERE bm_id=%s", [pk])
                except Exception:
                    cursor.execute("DELETE FROM branch_manager WHERE bm_id=%s", [pk])
            return Response({'status': 'deleted'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class DirectorListView(APIView):
    """Return list of director names from director table."""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # If a specific director name is provided, return full director row (useful to fetch phone_number_of_lolc)
        name = request.query_params.get('name')
        try:
            with connection.cursor() as cursor:
                if name:
                    cursor.execute("SELECT * FROM director WHERE name_of_director = %s LIMIT 1", [name])
                    cols = [c[0] for c in cursor.description] if cursor.description else []
                    row = cursor.fetchone()
                    if not row:
                        return Response({'director': None}, status=status.HTTP_200_OK)
                    data = dict(zip(cols, row))
                    return Response({'director': data}, status=status.HTTP_200_OK)
                else:
                    # Return full director objects so frontend can show id/name/phone
                    cursor.execute("SELECT director_id, name_of_director, phone_number_of_lolc FROM director ORDER BY name_of_director")
                    cols = [c[0] for c in cursor.description] if cursor.description else []
                    raw_rows = cursor.fetchall()
                    rows = [dict(zip(cols, r)) for r in raw_rows]
                    return Response({'directors': rows}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def post(self, request):
        """Create a new director record."""
        name = request.data.get('name') or request.data.get('name_of_director')
        phone = request.data.get('phone_number_of_lolc') or request.data.get('phone')
        if not name:
            return Response({'error': 'Missing name'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                # Some DB schemas have director_id NOT NULL without AUTO_INCREMENT.
                # Try to generate a new director_id if needed by taking MAX(director_id)+1.
                new_id = None
                try:
                    cursor.execute("SELECT COALESCE(MAX(director_id), 0) + 1 FROM director")
                    row = cursor.fetchone()
                    if row:
                        new_id = row[0]
                except Exception:
                    new_id = None

                if new_id is not None:
                    cursor.execute("INSERT INTO director (director_id, name_of_director, phone_number_of_lolc) VALUES (%s, %s, %s)", [new_id, name, phone])
                else:
                    cursor.execute("INSERT INTO director (name_of_director, phone_number_of_lolc) VALUES (%s, %s)", [name, phone])
            return Response({'status': 'created'}, status=status.HTTP_201_CREATED)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def patch(self, request, pk=None):
        """Update an existing director by id (pk)."""
        if not pk:
            return Response({'error': 'Missing director id'}, status=status.HTTP_400_BAD_REQUEST)
        name = request.data.get('name') or request.data.get('name_of_director')
        phone = request.data.get('phone_number_of_lolc') or request.data.get('phone')
        if name is None and phone is None:
            return Response({'error': 'No fields to update'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                updates = []
                params = []
                if name is not None:
                    updates.append('name_of_director=%s')
                    params.append(name)
                if phone is not None:
                    updates.append('phone_number_of_lolc=%s')
                    params.append(phone)
                params.append(pk)
                sql = f"UPDATE director SET {', '.join(updates)} WHERE director_id=%s"
                cursor.execute(sql, params)
            return Response({'status': 'updated'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def delete(self, request, pk=None):
        """Soft-delete a director by is_active if available, else delete row."""
        if not pk:
            return Response({'error': 'Missing director id'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            with connection.cursor() as cursor:
                try:
                    cursor.execute("UPDATE director SET is_active=0 WHERE director_id=%s", [pk])
                except Exception:
                    cursor.execute("DELETE FROM director WHERE director_id=%s", [pk])
            return Response({'status': 'deleted'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class BLAgreementDocxDownloadView(APIView):
    """Generate a DOCX from template for a given contract_number and return as attachment."""
    permission_classes = []

    def get(self, request):
        contract_number = request.query_params.get('contract_number', '').strip()
        if not contract_number:
            return Response({'error': 'contract_number parameter required'}, status=status.HTTP_400_BAD_REQUEST)

        # allow client to request an alternate template by filename (basename only)
        req_template = (request.query_params.get('template') or '').strip()
        if req_template:
            # sanitize to basename and ensure .docx extension
            req_template = os.path.basename(req_template)
            if not req_template.lower().endswith('.docx'):
                req_template = req_template + '.docx'
            template_path = os.path.join(settings.BASE_DIR, 'templates', 'docx', req_template)
        else:
            # default template
            template_path = os.path.join(settings.BASE_DIR, 'templates', 'docx', 'bl_agreement_template.docx')

        if not os.path.exists(template_path):
            return Response({'error': f'Template not found at {template_path}. Please place your .docx template there.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        try:
            # Fetch data for the contract. If a BL SP3 template is requested, pull
            # the latest row from `bl_sp3` instead of `bl_agreement` so the SP3
            # template gets the expected fields.
            req_template_basename = os.path.basename(template_path).lower()
            is_sp3_template = req_template_basename.startswith('bl_sp3')

            with connection.cursor() as cursor:
                if is_sp3_template:
                    # select latest sp3 row for this contract
                    cursor.execute("SELECT * FROM bl_sp3 WHERE LOWER(contract_number)=LOWER(%s) ORDER BY COALESCE(sp3_date, created_at) DESC LIMIT 1", [contract_number])
                    agreement_row = cursor.fetchone()
                    agreement_cols = [c[0] for c in cursor.description] if cursor.description else []
                    agreement = dict(zip(agreement_cols, agreement_row)) if agreement_row else {}
                else:
                    cursor.execute('SELECT * FROM bl_agreement WHERE contract_number=%s LIMIT 1', [contract_number])
                    agreement_row = cursor.fetchone()
                    agreement_cols = [c[0] for c in cursor.description] if cursor.description else []
                    agreement = dict(zip(agreement_cols, agreement_row)) if agreement_row else {}

                # collateral rows are still useful for SP3 templates (if present)
                cursor.execute('SELECT * FROM bl_collateral WHERE contract_number=%s LIMIT 1', [contract_number])
                coll_row = cursor.fetchone()
                coll_cols = [c[0] for c in cursor.description] if cursor.description else []
                collateral = dict(zip(coll_cols, coll_row)) if coll_row else {}

            # Build context for the template
            # Flatten agreement and collateral dicts into top-level variables expected by the .docx template
            ctx = {}
            # agreement fields take precedence
            if isinstance(agreement, dict):
                for k, v in agreement.items():
                    ctx[k] = v
            if isinstance(collateral, dict):
                for k, v in collateral.items():
                    # do not override agreement fields
                    if k not in ctx:
                        ctx[k] = v
            # always include contract_number
            ctx['contract_number'] = contract_number

            # Derived / formatted fields for template convenience
            # Numeric fields: provide formatted with dot thousands and _in_word
            numeric_keys = ['loan_amount', 'admin_fee', 'net_amount', 'notaris_fee', 'mortgage_amount']
            for nk in numeric_keys:
                val = ctx.get(nk)
                try:
                    # formatted number with dot thousands
                    ctx[nk] = format_number_dot(val) if val is not None else ''
                except Exception:
                    ctx[nk] = val
                # words
                try:
                    ctx[nk + '_in_word'] = number_to_indonesian_words(val) if val is not None else ''
                except Exception:
                    ctx[nk + '_in_word'] = ''

            # Date fields: provide human-readable Indonesian date strings
            date_keys = ['agreement_date', 'date_birth_of_debtor', 'date_birth_of_bm', 'sp3_date', 'date_of_delegated']
            for dk in date_keys:
                v = ctx.get(dk)
                try:
                    # human-readable date (e.g., '6 September 2026')
                    # Render human-readable date uppercased for documents
                    ctx[dk] = format_indonesian_date(v, uppercase_all=True) if v else ''
                    # spelled-out date words entirely UPPERCASE
                    ctx[dk + '_in_word'] = date_to_indonesian_words(v, uppercase_month=True, uppercase_all=True) if v else ''
                    # display with parentheses e.g. '(6 SEPTEMBER 2026)'
                    ctx[dk + '_display'] = f"({format_indonesian_date(v, uppercase_all=True)})" if v else ''
                except Exception:
                    ctx[dk + '_in_word'] = ''

            # Ensure sentence-case for any helper fields generated for words
            try:
                for key in list(ctx.keys()):
                    if isinstance(key, str) and (key.endswith('_in_word') or key.endswith('_by_word')):
                        val = ctx.get(key)
                        if val is None:
                            continue
                        s = str(val).strip()
                        if not s:
                            continue
                        # Sentence case: first char uppercase, rest lowercase
                        try:
                            ctx[key] = s[0].upper() + s[1:].lower() if len(s) > 1 else s.upper()
                        except Exception:
                            ctx[key] = s
            except Exception:
                pass

            # Uppercase names as requested
            try:
                if ctx.get('name_of_debtor'):
                    ctx['name_of_debtor'] = str(ctx.get('name_of_debtor')).upper()
            except Exception:
                pass
            try:
                if ctx.get('name_of_bm'):
                    ctx['name_of_bm'] = str(ctx.get('name_of_bm')).upper()
            except Exception:
                pass

            # agreement day name in Indonesian (e.g., 'Senin') based on agreement_date
            # Preserve any value already present (e.g., from DB). Only compute when absent.
            try:
                if ctx.get('agreement_day_in_word'):
                    # already provided by DB or earlier logic; keep it
                    pass
                else:
                    raw_date = None
                    # prefer original agreement dict value if available
                    try:
                        if isinstance(agreement, dict):
                            raw_date = agreement.get('agreement_date')
                    except Exception:
                        raw_date = None

                    # fallback to ctx value (may be ISO string) if raw_date not available
                    if not raw_date:
                        raw_date = ctx.get('agreement_date')

                    if raw_date:
                        from datetime import datetime
                        s = str(raw_date)
                        try:
                            d = datetime.fromisoformat(s).date()
                        except Exception:
                            try:
                                d = datetime.strptime(s, '%Y-%m-%d').date()
                            except Exception:
                                d = None
                        if d:
                            weekday_map = {
                                0: 'Senin',
                                1: 'Selasa',
                                2: 'Rabu',
                                3: 'Kamis',
                                4: 'Jumat',
                                5: 'Sabtu',
                                6: 'Minggu'
                            }
                            ctx['agreement_day_in_word'] = weekday_map.get(d.weekday(), '')
                        else:
                            ctx['agreement_day_in_word'] = ctx.get('agreement_day_in_word', '')
                    else:
                        ctx['agreement_day_in_word'] = ctx.get('agreement_day_in_word', '')
            except Exception:
                ctx['agreement_day_in_word'] = ctx.get('agreement_day_in_word', '')
            # Ensure phone_number_of_lolc is set (may be present as director phone)
            try:
                ctx['phone_number_of_lolc'] = ctx.get('phone_number_of_lolc') or ctx.get('phone_number_of_director') or ''
            except Exception:
                pass

            # Render DOCX using docxtpl
            try:
                from docxtpl import DocxTemplate
            except Exception as imp_e:
                # Log diagnostic info to help identify which Python/environment is running
                try:
                    logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                    os.makedirs(logs_dir, exist_ok=True)
                    diag_path = os.path.join(logs_dir, 'bl_agreement_docx_import_diag.log')
                    with open(diag_path, 'a', encoding='utf-8') as df:
                        df.write(f"[{timezone.now().isoformat()}] docxtpl import failed: {str(imp_e)}\n")
                        df.write(f"executable: {sys.executable}\n")
                        df.write(f"python_version: {sys.version}\n")
                        df.write(f"sys.path:\n")
                        for p in sys.path:
                            df.write(f"  {p}\n")
                        df.write("\n")
                except Exception:
                    pass
                logger.error('docxtpl import failed: %s', str(imp_e))
                if getattr(settings, 'DEBUG', False):
                    return Response({'error': 'docxtpl import failed', 'detail': str(imp_e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                return Response({'error': 'docxtpl not installed. Please install with `pip install docxtpl`.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            # Preprocess template to repair Jinja tags that were split across
            # Word XML runs (common when template was edited in Word).
            def _repair_docx_jinja_tags(src_path, contract_no=None):
                try:
                    tmpdir = tempfile.mkdtemp(prefix='docx_repair_')
                    with zipfile.ZipFile(src_path, 'r') as zin:
                        zin.extractall(tmpdir)

                    word_dir = os.path.join(tmpdir, 'word')
                    if not os.path.exists(word_dir):
                        shutil.rmtree(tmpdir)
                        return src_path

                    changed_files = []

                    # process all xml files under word/ (document, headers, footers, notes, comments)
                    for root, dirs, files in os.walk(word_dir):
                        for fname in files:
                            if not fname.lower().endswith('.xml'):
                                continue
                            fullpath = os.path.join(root, fname)
                            with open(fullpath, 'r', encoding='utf-8') as f:
                                data = f.read()
                            # Normalize problematic Unicode characters that may
                            # break Jinja lexing when templates are embedded inside
                            # Word XML runs. Replace with safe ASCII equivalents.
                            if '\u2026' in data:
                                data = data.replace('\u2026', '...')
                            # common problematic characters: NBSP, ZWSP, BOM, soft hyphen, fancy quotes, dashes
                            replacements = {
                                '\u00A0': ' ', '\u200B': '', '\u200C': '', '\u200D': '', '\u00AD': '', '\ufeff': '',
                                '\u2018': "'", '\u2019': "'", '\u201C': '"', '\u201D': '"',
                                '\u2013': '-', '\u2014': '-', '\u2010': '-', '\u2022': '*'
                            }
                            for k, v in replacements.items():
                                if k in data:
                                    data = data.replace(k, v)

                            if '{{' not in data or '}}' not in data:
                                continue

                            orig = data
                            new = []
                            pos = 0
                            L = len(data)
                            while pos < L:
                                i = data.find('{{', pos)
                                if i == -1:
                                    new.append(data[pos:])
                                    break
                                # append before
                                new.append(data[pos:i])
                                j = data.find('}}', i)
                                if j == -1:
                                    # unterminated, append rest and break
                                    new.append(data[i:])
                                    break
                                segment = data[i:j+2]
                                # remove XML tags inside the segment
                                cleaned = re.sub(r'<[^>]+>', '', segment)
                                # normalize whitespace inside the braces
                                if cleaned.count('{{') and cleaned.count('}}'):
                                    inner_start = cleaned.find('{{') + 2
                                    inner_end = cleaned.rfind('}}')
                                    inner = cleaned[inner_start:inner_end]
                                    # remove all whitespace inside the jinja expression to
                                    # fix cases where Word split tokens across runs and
                                    # inserted spaces (e.g. "flate_rate _by_word").
                                    inner = re.sub(r'\s+', '', inner)
                                    cleaned_full = '{{' + inner + '}}'
                                    new.append(cleaned_full)
                                else:
                                    # fallback to raw cleaned (if braces disappeared)
                                    cleaned2 = re.sub(r'<[^>]+>', '', segment)
                                    new.append(cleaned2)
                                pos = j+2

                            repaired = ''.join(new)
                            if repaired != orig:
                                with open(fullpath, 'w', encoding='utf-8') as f:
                                    f.write(repaired)
                                rel = os.path.relpath(fullpath, tmpdir)
                                changed_files.append(rel.replace('\\', '/'))

                    if changed_files:
                        # create repaired docx
                        repaired_fd, repaired_path = tempfile.mkstemp(suffix='.docx')
                        os.close(repaired_fd)
                        with zipfile.ZipFile(repaired_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                            for root, dirs, files in os.walk(tmpdir):
                                for file in files:
                                    full = os.path.join(root, file)
                                    arcname = os.path.relpath(full, tmpdir)
                                    zout.write(full, arcname)

                        # log repair info
                        try:
                            logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                            os.makedirs(logs_dir, exist_ok=True)
                            repair_log = os.path.join(logs_dir, 'bl_agreement_docx_repair.log')
                            with open(repair_log, 'a', encoding='utf-8') as lf:
                                lf.write(f"[{timezone.now().isoformat()}] repaired={repaired_path} contract={contract_no} files={changed_files}\n")
                        except Exception:
                            pass

                        shutil.rmtree(tmpdir)
                        return repaired_path
                    else:
                        shutil.rmtree(tmpdir)
                        return src_path
                except Exception:
                    try:
                        shutil.rmtree(tmpdir)
                    except Exception:
                        pass
                    return src_path

            # Provide nested header_fields dict for templates that expect header_fields.<key>
            try:
                hdr = {}
                for key in ('agreement_date','agreement_date_in_word','agreement_date_display','agreement_day_in_word','sp3_date','sp3_date_in_word','sp3_date_display','date_of_delegated','date_of_delegated_in_word','date_of_delegated_display','name_of_debtor','phone_number_of_lolc','sp3_number','contract_number'):
                    hdr[key] = ctx.get(key, '')
                ctx['header_fields'] = hdr
            except Exception:
                pass

            repaired_template = _repair_docx_jinja_tags(template_path, contract_no=contract_number)
            try:
                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                os.makedirs(logs_dir, exist_ok=True)
                with open(os.path.join(logs_dir, 'bl_agreement_docx_repair.log'), 'a', encoding='utf-8') as lf:
                    lf.write(f"[{timezone.now().isoformat()}] using_template={repaired_template} contract={contract_number}\n")
            except Exception:
                pass
            tpl = DocxTemplate(repaired_template)
            fallback_docx_path = None
            try:
                tpl.render(ctx)
            except Exception as render_e:
                logger.error('docxtpl render failed for %s: %s', repaired_template, str(render_e))
                try:
                    from docx import Document as _DocxDocument
                    doc_fallback = _DocxDocument()
                    doc_fallback.add_heading(f'BL Agreement - {contract_number}', level=1)
                    for k in ('sp3_number', 'sp3_date', 'name_of_debtor', 'contract_number'):
                        v = ctx.get(k, '')
                        doc_fallback.add_paragraph(f"{k}: {v}")
                    try:
                        hdr = ctx.get('header_fields') or ctx.get('header') or {}
                        if isinstance(hdr, dict) and hdr:
                            doc_fallback.add_heading('Header Fields', level=2)
                            for hk, hv in hdr.items():
                                doc_fallback.add_paragraph(f"{hk}: {hv}")
                    except Exception:
                        pass
                    fd, fbpath = tempfile.mkstemp(suffix='.docx')
                    os.close(fd)
                    doc_fallback.save(fbpath)
                    fallback_docx_path = fbpath
                    tpl = None
                except Exception as docx_e:
                    logger.error('Fallback DOCX generation failed: %s', str(docx_e))
                    raise

            # save to temporary file and return as attachment
            tmpdir = tempfile.mkdtemp(prefix='docx_out_')
            try:
                # sanitize contract_number to avoid path-separators causing nested dirs
                try:
                    import re as _re
                    safe_cn = _re.sub(r'[^A-Za-z0-9._-]', '_', str(contract_number or ''))
                    if not safe_cn:
                        safe_cn = timezone.now().strftime('no_contract_%Y%m%d%H%M%S')
                except Exception:
                    safe_cn = str(contract_number or 'contract')
                docx_path = os.path.join(tmpdir, f'bl_agreement_{safe_cn}.docx')
                if tpl is not None:
                    tpl.save(docx_path)
                else:
                    if fallback_docx_path and os.path.exists(fallback_docx_path):
                        shutil.copyfile(fallback_docx_path, docx_path)
                        try:
                            os.remove(fallback_docx_path)
                        except Exception:
                            pass
                    else:
                        raise
                # Support optional PDF conversion when requested via ?download=pdf
                download = request.query_params.get('download', '').strip().lower()
                if download == 'pdf':
                    pdf_path = os.path.join(tmpdir, f'bl_agreement_{safe_cn}.pdf')
                    try:
                        import shutil as _sh
                        if _sh.which('soffice') is None:
                            return Response({'error': 'LibreOffice (soffice) not found on server'}, status=status.HTTP_503_SERVICE_UNAVAILABLE)
                    except Exception:
                        return Response({'error': 'Server environment check failed (soffice detection)'}, status=status.HTTP_503_SERVICE_UNAVAILABLE)

                    ok, err = _convert_docx_to_pdf(docx_path, pdf_path)
                    if ok:
                        try:
                            with open(pdf_path, 'rb') as fh:
                                pdf_bytes = fh.read()
                            # Centralized document log entry
                            try:
                                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                                os.makedirs(logs_dir, exist_ok=True)
                                central = os.path.join(logs_dir, 'document_generation.log')
                                username = getattr(request, 'user', None)
                                uname = username.username if getattr(username, 'is_authenticated', False) else 'anonymous'
                                size = len(pdf_bytes)
                                with open(central, 'a', encoding='utf-8') as cf:
                                    cf.write(f"[{timezone.now().isoformat()}] user={uname} type=pdf endpoint=bl_agreement contract={contract_number} file=bl_agreement_{safe_cn}.pdf size={size}\n")
                                logger.info('PDF generated: bl_agreement_%s.pdf size=%s user=%s', safe_cn, size, uname)
                            except Exception:
                                logger.exception('Failed to write centralized document log (BL PDF) %s', contract_number)
                            response = HttpResponse(pdf_bytes, content_type='application/pdf')
                            response['Content-Disposition'] = f'attachment; filename="bl_agreement_{safe_cn}.pdf"'
                            return response
                        except Exception as e2:
                            logger.error('Failed to return PDF for %s: %s', contract_number, str(e2))
                            return Response({'error': 'Failed to return PDF.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


                # contracts list class was moved to top-level to avoid nesting inside
                # the BL agreement document generation handler.
                    else:
                        logger.error('PDF conversion failed for %s: %s', contract_number, str(err))
                        return Response({'error': 'PDF conversion failed', 'detail': str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                else:
                    try:
                        with open(docx_path, 'rb') as fh:
                            docx_bytes = fh.read()
                        # Centralized document log entry
                        try:
                            logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                            os.makedirs(logs_dir, exist_ok=True)
                            central = os.path.join(logs_dir, 'document_generation.log')
                            username = getattr(request, 'user', None)
                            uname = username.username if getattr(username, 'is_authenticated', False) else 'anonymous'
                            size = len(docx_bytes)
                            with open(central, 'a', encoding='utf-8') as cf:
                                cf.write(f"[{timezone.now().isoformat()}] user={uname} type=docx endpoint=bl_agreement contract={contract_number} file=bl_agreement_{safe_cn}.docx size={size}\n")
                            logger.info('DOCX generated: bl_agreement_%s.docx size=%s user=%s', safe_cn, size, uname)
                        except Exception:
                            logger.exception('Failed to write centralized document log (BL DOCX) %s', contract_number)
                        response = HttpResponse(docx_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = f'attachment; filename="bl_agreement_{safe_cn}.docx"'
                        return response
                    except Exception as e2:
                        logger.error('Failed to return DOCX for %s: %s', contract_number, str(e2))
                        return Response({'error': 'Failed to return DOCX.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            finally:
                try:
                    _safe_rmtree(tmpdir)
                except Exception:
                    pass
                # cleaned up temporary dir; do not attempt to read files here

        except Exception as e:
            # Capture full traceback and write to server log file for debugging
            tb = traceback.format_exc()
            try:
                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                os.makedirs(logs_dir, exist_ok=True)
                log_path = os.path.join(logs_dir, 'bl_agreement_docx_errors.log')
                with open(log_path, 'a', encoding='utf-8') as lf:
                    lf.write(f"[{timezone.now().isoformat()}] Error generating docx for {contract_number}\n")
                    lf.write(tb + "\n")
            except Exception:
                # ignore logging errors
                pass

            logger.error('Error generating DOCX for %s: %s', contract_number, tb)

            if getattr(settings, 'DEBUG', False):
                return Response({'error': str(e), 'traceback': tb}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            return Response({'error': 'Internal server error while generating DOCX. Check server logs.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class UVAgreementDocxDownloadView(APIView):
    """Generate a DOCX from UV template for a given contract_number and return as attachment."""
    permission_classes = []

    def get(self, request):
        contract_number = request.query_params.get('contract_number', '').strip()
        if not contract_number:
            return Response({'error': 'contract_number parameter required'}, status=status.HTTP_400_BAD_REQUEST)

        # locate UV template file
        template_path = os.path.join(settings.BASE_DIR, 'templates', 'docx', 'uv_agreement_template.docx')
        if not os.path.exists(template_path):
            return Response({'error': f'UV template not found at {template_path}. Please place your .docx template there.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        try:
            # Fetch data for the contract from uv_agreement and uv_collateral (if available)
            with connection.cursor() as cursor:
                cursor.execute('SELECT * FROM uv_agreement WHERE contract_number=%s LIMIT 1', [contract_number])
                agreement_row = cursor.fetchone()
                agreement_cols = [c[0] for c in cursor.description] if cursor.description else []
                agreement = dict(zip(agreement_cols, agreement_row)) if agreement_row else {}

                cursor.execute('SELECT * FROM uv_collateral WHERE contract_number=%s LIMIT 1', [contract_number])
                coll_row = cursor.fetchone()
                coll_cols = [c[0] for c in cursor.description] if cursor.description else []
                collateral = dict(zip(coll_cols, coll_row)) if coll_row else {}

            ctx = {}
            if isinstance(agreement, dict):
                for k, v in agreement.items():
                    ctx[k] = v
            if isinstance(collateral, dict):
                for k, v in collateral.items():
                    if k not in ctx:
                        ctx[k] = v
            ctx['contract_number'] = contract_number

            # Derived fields: numeric formatting and in-word
            numeric_keys = ['loan_amount', 'admin_fee', 'net_amount', 'notaris_fee', 'mortgage_amount']
            for nk in numeric_keys:
                val = ctx.get(nk)
                try:
                    ctx[nk] = format_number_dot(val) if val is not None else ''
                except Exception:
                    ctx[nk] = val
                try:
                    ctx[nk + '_in_word'] = number_to_indonesian_words(val) if val is not None else ''
                except Exception:
                    ctx[nk + '_in_word'] = ''

            date_keys = ['agreement_date', 'date_birth_of_debtor', 'sp3_date']
            for dk in date_keys:
                v = ctx.get(dk)
                try:
                    ctx[dk] = format_indonesian_date(v, uppercase_all=True) if v else ''
                    ctx[dk + '_in_word'] = date_to_indonesian_words(v, uppercase_month=True, uppercase_all=True) if v else ''
                    ctx[dk + '_display'] = f"({format_indonesian_date(v, uppercase_all=True)})" if v else ''
                except Exception:
                    ctx[dk + '_in_word'] = ''

            try:
                if ctx.get('name_of_debtor'):
                    ctx['name_of_debtor'] = str(ctx.get('name_of_debtor')).upper()
            except Exception:
                pass
            try:
                ctx['phone_number_of_lolc'] = ctx.get('phone_number_of_lolc') or ctx.get('phone_number_of_director') or ''
            except Exception:
                pass

            try:
                from docxtpl import DocxTemplate
            except Exception as imp_e:
                logger.error('docxtpl import failed for UV: %s', str(imp_e))
                if getattr(settings, 'DEBUG', False):
                    return Response({'error': 'docxtpl import failed', 'detail': str(imp_e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                return Response({'error': 'docxtpl not installed. Please install with `pip install docxtpl`.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            def _repair_docx_jinja_tags_uv(src_path, contract_no=None):
                # reuse same repair logic as BL but keep function name distinct
                try:
                    tmpdir = tempfile.mkdtemp(prefix='docx_repair_uv_')
                    with zipfile.ZipFile(src_path, 'r') as zin:
                        zin.extractall(tmpdir)

                    word_dir = os.path.join(tmpdir, 'word')
                    if not os.path.exists(word_dir):
                        shutil.rmtree(tmpdir)
                        return src_path

                    changed_files = []
                    for root, dirs, files in os.walk(word_dir):
                        for fname in files:
                            if not fname.lower().endswith('.xml'):
                                continue
                            fullpath = os.path.join(root, fname)
                            with open(fullpath, 'r', encoding='utf-8') as f:
                                data = f.read()
                            # Normalize problematic Unicode characters
                            if '\u2026' in data:
                                data = data.replace('\u2026', '...')
                            replacements = {
                                '\u00A0': ' ', '\u200B': '', '\u200C': '', '\u200D': '', '\u00AD': '', '\ufeff': '',
                                '\u2018': "'", '\u2019': "'", '\u201C': '"', '\u201D': '"',
                                '\u2013': '-', '\u2014': '-', '\u2010': '-', '\u2022': '*'
                            }
                            for k, v in replacements.items():
                                if k in data:
                                    data = data.replace(k, v)
                            if '{{' not in data or '}}' not in data:
                                continue
                            orig = data
                            new = []
                            pos = 0
                            L = len(data)
                            while pos < L:
                                i = data.find('{{', pos)
                                if i == -1:
                                    new.append(data[pos:])
                                    break
                                new.append(data[pos:i])
                                j = data.find('}}', i)
                                if j == -1:
                                    new.append(data[i:])
                                    break
                                segment = data[i:j+2]
                                cleaned = re.sub(r'<[^>]+>', '', segment)
                                if cleaned.count('{{') and cleaned.count('}}'):
                                    inner_start = cleaned.find('{{') + 2
                                    inner_end = cleaned.rfind('}}')
                                    inner = cleaned[inner_start:inner_end]
                                    inner = re.sub(r'\s+', '', inner)
                                    cleaned_full = '{{' + inner + '}}'
                                    new.append(cleaned_full)
                                else:
                                    cleaned2 = re.sub(r'<[^>]+>', '', segment)
                                    new.append(cleaned2)
                                pos = j+2
                            repaired = ''.join(new)
                            if repaired != orig:
                                with open(fullpath, 'w', encoding='utf-8') as f:
                                    f.write(repaired)
                                rel = os.path.relpath(fullpath, tmpdir)
                                changed_files.append(rel.replace('\\', '/'))

                    if changed_files:
                        repaired_fd, repaired_path = tempfile.mkstemp(suffix='.docx')
                        os.close(repaired_fd)
                        with zipfile.ZipFile(repaired_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                            for root, dirs, files in os.walk(tmpdir):
                                for file in files:
                                    full = os.path.join(root, file)
                                    arcname = os.path.relpath(full, tmpdir)
                                    zout.write(full, arcname)
                        try:
                            logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                            os.makedirs(logs_dir, exist_ok=True)
                            repair_log = os.path.join(logs_dir, 'uv_agreement_docx_repair.log')
                            with open(repair_log, 'a', encoding='utf-8') as lf:
                                lf.write(f"[{timezone.now().isoformat()}] repaired={repaired_path} contract={contract_no} files={changed_files}\n")
                        except Exception:
                            pass
                        shutil.rmtree(tmpdir)
                        return repaired_path
                    else:
                        shutil.rmtree(tmpdir)
                        return src_path
                except Exception:
                    try:
                        shutil.rmtree(tmpdir)
                    except Exception:
                        pass
                    return src_path

            repaired_template = _repair_docx_jinja_tags_uv(template_path, contract_no=contract_number)
            # Ensure nested header_fields present for UV templates
            try:
                hdr_uv = {}
                for key in ('agreement_date','agreement_date_in_word','agreement_date_display','agreement_day_in_word','sp3_date','sp3_date_in_word','sp3_date_display','name_of_debtor','phone_number_of_lolc','sp3_number','contract_number'):
                    hdr_uv[key] = ctx.get(key, '')
                ctx['header_fields'] = hdr_uv
            except Exception:
                pass
            tpl = DocxTemplate(repaired_template)
            fallback_docx_path = None
            try:
                tpl.render(ctx)
            except Exception as render_e:
                logger.error('docxtpl render failed for UV %s: %s', repaired_template, str(render_e))
                try:
                    from docx import Document as _DocxDocument
                    doc_fallback = _DocxDocument()
                    doc_fallback.add_heading(f'UV Agreement - {contract_number}', level=1)
                    for k in ('sp3_number', 'sp3_date', 'name_of_debtor', 'contract_number'):
                        v = ctx.get(k, '')
                        doc_fallback.add_paragraph(f"{k}: {v}")
                    try:
                        hdr = ctx.get('header_fields') or ctx.get('header') or {}
                        if isinstance(hdr, dict) and hdr:
                            doc_fallback.add_heading('Header Fields', level=2)
                            for hk, hv in hdr.items():
                                doc_fallback.add_paragraph(f"{hk}: {hv}")
                    except Exception:
                        pass
                    fd, fbpath = tempfile.mkstemp(suffix='.docx')
                    os.close(fd)
                    doc_fallback.save(fbpath)
                    fallback_docx_path = fbpath
                    tpl = None
                except Exception as docx_e:
                    logger.error('Fallback DOCX generation failed for UV: %s', str(docx_e))
                    raise

            tmpdir = tempfile.mkdtemp(prefix='docx_out_uv_')
            try:
                # sanitize contract_number to avoid path-separators causing nested dirs
                try:
                    import re as _re
                    safe_cn = _re.sub(r'[^A-Za-z0-9._-]', '_', str(contract_number or ''))
                    if not safe_cn:
                        safe_cn = timezone.now().strftime('no_contract_%Y%m%d%H%M%S')
                except Exception:
                    safe_cn = str(contract_number or 'contract')
                base_prefix = 'uv_agreement'
                docx_path = os.path.join(tmpdir, f'{base_prefix}_{safe_cn}.docx')
                if tpl is not None:
                    tpl.save(docx_path)
                else:
                    if fallback_docx_path and os.path.exists(fallback_docx_path):
                        shutil.copyfile(fallback_docx_path, docx_path)
                        try:
                            os.remove(fallback_docx_path)
                        except Exception:
                            pass
                    else:
                        raise
                # Support optional PDF conversion when requested via ?download=pdf
                download = request.query_params.get('download', '').strip().lower()
                if download == 'pdf':
                    pdf_path = os.path.join(tmpdir, f'{base_prefix}_{safe_cn}.pdf')
                    try:
                        import shutil as _sh
                        if _sh.which('soffice') is None:
                            return Response({'error': 'LibreOffice (soffice) not found on server'}, status=status.HTTP_503_SERVICE_UNAVAILABLE)
                    except Exception:
                        return Response({'error': 'Server environment check failed (soffice detection)'}, status=status.HTTP_503_SERVICE_UNAVAILABLE)

                    ok, err = _convert_docx_to_pdf(docx_path, pdf_path)
                    if ok:
                        try:
                            with open(pdf_path, 'rb') as fh:
                                pdf_bytes = fh.read()
                            # Centralized document log entry
                            try:
                                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                                os.makedirs(logs_dir, exist_ok=True)
                                central = os.path.join(logs_dir, 'document_generation.log')
                                username = getattr(request, 'user', None)
                                uname = username.username if getattr(username, 'is_authenticated', False) else 'anonymous'
                                size = len(pdf_bytes)
                                with open(central, 'a', encoding='utf-8') as cf:
                                    cf.write(f"[{timezone.now().isoformat()}] user={uname} type=pdf endpoint=uv_agreement contract={contract_number} file=uv_agreement_{safe_cn}.pdf size={size}\n")
                                logger.info('PDF generated: uv_agreement_%s.pdf size=%s user=%s', safe_cn, size, uname)
                            except Exception:
                                logger.exception('Failed to write centralized document log (UV PDF) %s', contract_number)
                            response = HttpResponse(pdf_bytes, content_type='application/pdf')
                            download_filename = f'{base_prefix}_{safe_cn}.pdf'.upper()
                            response['Content-Disposition'] = f'attachment; filename="{download_filename}"'
                            try:
                                ip = request.META.get('HTTP_X_FORWARDED_FOR') or request.META.get('REMOTE_ADDR') or ''
                                if ip and ',' in ip:
                                    ip = ip.split(',')[0].strip()
                                ua = request.META.get('HTTP_USER_AGENT', '')
                                user = getattr(request, 'user', None)
                                uid = getattr(user, 'id', None) if user and getattr(user, 'is_authenticated', False) else None
                                uname = getattr(user, 'username', None) if user and getattr(user, 'is_authenticated', False) else _resolve_username(request) or ''
                                email = getattr(user, 'email', None) if user and getattr(user, 'is_authenticated', False) else ''
                                DownloadLog.objects.create(
                                    user_id=uid,
                                    username=uname,
                                    email=email,
                                    file_type='uv',
                                    file_identifier=str(contract_number),
                                    filename=download_filename,
                                    ip_address=ip,
                                    user_agent=ua,
                                    success=True,
                                    file_size=len(pdf_bytes),
                                    method='stream',
                                )
                            except Exception:
                                logger.exception('Failed to write DownloadLog for UV PDF %s', contract_number)
                            return response
                        except Exception as e2:
                            logger.error('Failed to return UV PDF for %s: %s', contract_number, str(e2))
                            return Response({'error': 'Failed to return PDF.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                    else:
                        logger.error('UV PDF conversion failed for %s: %s', contract_number, str(err))
                        return Response({'error': 'PDF conversion failed', 'detail': str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                else:
                    try:
                        with open(docx_path, 'rb') as fh:
                            docx_bytes = fh.read()
                        # Centralized document log entry
                        try:
                            logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                            os.makedirs(logs_dir, exist_ok=True)
                            central = os.path.join(logs_dir, 'document_generation.log')
                            username = getattr(request, 'user', None)
                            uname = username.username if getattr(username, 'is_authenticated', False) else 'anonymous'
                            size = len(docx_bytes)
                            with open(central, 'a', encoding='utf-8') as cf:
                                cf.write(f"[{timezone.now().isoformat()}] user={uname} type=docx endpoint=uv_agreement contract={contract_number} file=uv_agreement_{safe_cn}.docx size={size}\n")
                            logger.info('DOCX generated: uv_agreement_%s.docx size=%s user=%s', safe_cn, size, uname)
                        except Exception:
                            logger.exception('Failed to write centralized document log (UV DOCX) %s', contract_number)
                        response = HttpResponse(docx_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        download_filename = f'{base_prefix}_{safe_cn}.docx'.upper()
                        response['Content-Disposition'] = f'attachment; filename="{download_filename}"'
                        try:
                            ip = request.META.get('HTTP_X_FORWARDED_FOR') or request.META.get('REMOTE_ADDR') or ''
                            if ip and ',' in ip:
                                ip = ip.split(',')[0].strip()
                            ua = request.META.get('HTTP_USER_AGENT', '')
                            user = getattr(request, 'user', None)
                            uid = getattr(user, 'id', None) if user and getattr(user, 'is_authenticated', False) else None
                            uname = getattr(user, 'username', None) if user and getattr(user, 'is_authenticated', False) else _resolve_username(request) or ''
                            email = getattr(user, 'email', None) if user and getattr(user, 'is_authenticated', False) else ''
                            DownloadLog.objects.create(
                                user_id=uid,
                                username=uname,
                                email=email,
                                file_type='uv',
                                file_identifier=str(contract_number),
                                filename=download_filename,
                                ip_address=ip,
                                user_agent=ua,
                                success=True,
                                file_size=len(docx_bytes),
                                method='stream',
                            )
                        except Exception:
                            logger.exception('Failed to write DownloadLog for UV DOCX %s', contract_number)
                        return response
                    except Exception as e2:
                        logger.error('Failed to return UV DOCX for %s: %s', contract_number, str(e2))
                        return Response({'error': 'Failed to return DOCX.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            finally:
                try:
                    _safe_rmtree(tmpdir)
                except Exception:
                    pass

        except Exception as e:
            tb = traceback.format_exc()
            try:
                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                os.makedirs(logs_dir, exist_ok=True)
                log_path = os.path.join(logs_dir, 'uv_agreement_docx_errors.log')
                with open(log_path, 'a', encoding='utf-8') as lf:
                    lf.write(f"[{timezone.now().isoformat()}] Error generating UV docx for {contract_number}\n")
                    lf.write(tb + "\n")
            except Exception:
                pass
            logger.error('Error generating UV DOCX for %s: %s', contract_number, tb)
            if getattr(settings, 'DEBUG', False):
                return Response({'error': str(e), 'traceback': tb}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            return Response({'error': 'Internal server error while generating UV DOCX. Check server logs.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class UVSP3DocxDownloadView(APIView):
    """Generate a DOCX from uv_sp3 template for a given contract_number and return as attachment."""
    permission_classes = []

    def get(self, request):
        contract_number = request.query_params.get('contract_number', '').strip()
        if not contract_number:
            return Response({'error': 'contract_number parameter required'}, status=status.HTTP_400_BAD_REQUEST)

        template_path = os.path.join(settings.BASE_DIR, 'templates', 'docx', 'uv_sp3_template.docx')
        if not os.path.exists(template_path):
            return Response({'error': f'UV SP3 template not found at {template_path}. Please place your .docx template there.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        try:
            # Fetch the row from uv_sp3
            with connection.cursor() as cursor:
                cursor.execute('SELECT * FROM uv_sp3 WHERE contract_number=%s LIMIT 1', [contract_number])
                row = cursor.fetchone()
                cols = [c[0] for c in cursor.description] if cursor.description else []
                data = dict(zip(cols, row)) if row else {}

            ctx = {}
            if isinstance(data, dict):
                for k, v in data.items():
                    ctx[k] = v

            # If some fields are nested dicts or JSON strings (common when frontend
            # stores header_fields/contract_data/collateral_data), merge them into
            # top-level ctx so templates using flat placeholders like
            # {{place_of_agreement}} or {{name_of_debtor}} will resolve.
            try:
                import json
                nested_keys = ['header_fields', 'header', 'contract_data', 'debtor', 'collateral_data', 'collateral']
                for nk in nested_keys:
                    if nk in ctx and ctx[nk]:
                        val = ctx[nk]
                        if isinstance(val, str):
                            try:
                                val_parsed = json.loads(val)
                                if isinstance(val_parsed, dict):
                                    for kk, vv in val_parsed.items():
                                        if kk not in ctx or ctx.get(kk) in (None, ''):
                                            ctx[kk] = vv
                                continue
                            except Exception:
                                pass
                        if isinstance(val, dict):
                            for kk, vv in val.items():
                                if kk not in ctx or ctx.get(kk) in (None, ''):
                                    ctx[kk] = vv
                # If collateral fields were provided as a list (multiple collateral rows),
                # prefer the first item and merge its keys to top-level so template placeholders
                # like {{vehicle_model}} resolve correctly.
                try:
                    for list_key in ('collateral', 'collateral_data'):
                        if list_key in ctx and isinstance(ctx[list_key], list) and len(ctx[list_key]) > 0:
                            first = ctx[list_key][0]
                            if isinstance(first, dict):
                                for kk, vv in first.items():
                                    if kk not in ctx or ctx.get(kk) in (None, ''):
                                        ctx[kk] = vv
                except Exception:
                    pass
            except Exception:
                pass

            ctx['contract_number'] = contract_number

            try:
                from docxtpl import DocxTemplate
            except Exception as imp_e:
                logger.error('docxtpl import failed for UV SP3: %s', str(imp_e))
                if getattr(settings, 'DEBUG', False):
                    return Response({'error': 'docxtpl import failed', 'detail': str(imp_e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                return Response({'error': 'docxtpl not installed. Please install with `pip install docxtpl`.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            try:
                # Repair potential broken Jinja tags inside DOCX XML parts (common when templates
                # are edited in Word and contain extra tags/HTML). This mirrors the robust
                # repair logic used by UVAgreementDocxDownloadView.
                def _repair_docx_jinja_tags_sp3(src_path, contract_no=None):
                    try:
                        tmpdir = tempfile.mkdtemp(prefix='docx_repair_uv_sp3_')
                        with zipfile.ZipFile(src_path, 'r') as zin:
                            zin.extractall(tmpdir)

                        word_dir = os.path.join(tmpdir, 'word')
                        if not os.path.exists(word_dir):
                            shutil.rmtree(tmpdir)
                            return src_path

                        changed_files = []
                        for root, dirs, files in os.walk(word_dir):
                            for fname in files:
                                if not fname.lower().endswith('.xml'):
                                    continue
                                fullpath = os.path.join(root, fname)
                                with open(fullpath, 'r', encoding='utf-8') as f:
                                    data = f.read()
                                if '{{' not in data or '}}' not in data:
                                    continue
                                orig = data
                                new = []
                                pos = 0
                                L = len(data)
                                while pos < L:
                                    i = data.find('{{', pos)
                                    if i == -1:
                                        new.append(data[pos:])
                                        break
                                    new.append(data[pos:i])
                                    j = data.find('}}', i)
                                    if j == -1:
                                        new.append(data[i:])
                                        break
                                    segment = data[i:j+2]
                                    cleaned = re.sub(r'<[^>]+>', '', segment)
                                    if cleaned.count('{{') and cleaned.count('}}'):
                                        inner_start = cleaned.find('{{') + 2
                                        inner_end = cleaned.rfind('}}')
                                        inner = cleaned[inner_start:inner_end]
                                        inner = re.sub(r'\s+', '', inner)
                                        cleaned_full = '{{' + inner + '}}'
                                        new.append(cleaned_full)
                                    else:
                                        cleaned2 = re.sub(r'<[^>]+>', '', segment)
                                        new.append(cleaned2)
                                    pos = j+2
                                repaired = ''.join(new)
                                if repaired != orig:
                                    with open(fullpath, 'w', encoding='utf-8') as f:
                                        f.write(repaired)
                                    rel = os.path.relpath(fullpath, tmpdir)
                                    changed_files.append(rel.replace('\\', '/'))

                        if changed_files:
                            repaired_fd, repaired_path = tempfile.mkstemp(suffix='.docx')
                            os.close(repaired_fd)
                            with zipfile.ZipFile(repaired_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                                for root, dirs, files in os.walk(tmpdir):
                                    for file in files:
                                        full = os.path.join(root, file)
                                        arcname = os.path.relpath(full, tmpdir)
                                        zout.write(full, arcname)
                            try:
                                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                                os.makedirs(logs_dir, exist_ok=True)
                                repair_log = os.path.join(logs_dir, 'uv_sp3_docx_repair.log')
                                with open(repair_log, 'a', encoding='utf-8') as lf:
                                    lf.write(f"[{timezone.now().isoformat()}] repaired={repaired_path} contract={contract_no} files={changed_files}\n")
                            except Exception:
                                pass
                            shutil.rmtree(tmpdir)
                            return repaired_path
                        else:
                            shutil.rmtree(tmpdir)
                            return src_path
                    except Exception:
                        try:
                            shutil.rmtree(tmpdir)
                        except Exception:
                            pass
                        return src_path

                # Provide derived fields commonly used by templates (dates, numeric formatting, header grouping)
                try:
                    # numeric/date formatting similar to UVAgreement
                    numeric_keys = ['loan_amount', 'admin_fee', 'net_amount', 'notaris_fee', 'mortgage_amount', 'admin_rate', 'tlo', 'life_insurance', 'total_amount']
                    for nk in numeric_keys:
                        val = ctx.get(nk)
                        try:
                            ctx[nk] = format_number_dot(val) if val is not None else ''
                        except Exception:
                            ctx[nk] = val
                        try:
                            ctx[nk + '_in_word'] = number_to_indonesian_words(val) if val is not None else ''
                        except Exception:
                            ctx[nk + '_in_word'] = ''

                    date_keys = ['agreement_date', 'sp3_date', 'date_of_delegated']
                    for dk in date_keys:
                        v = ctx.get(dk)
                        try:
                            ctx[dk] = format_indonesian_date(v, uppercase_all=True) if v else ''
                            ctx[dk + '_in_word'] = date_to_indonesian_words(v, uppercase_month=True, uppercase_all=True) if v else ''
                            ctx[dk + '_display'] = f"({format_indonesian_date(v, uppercase_all=True)})" if v else ''
                        except Exception:
                            ctx[dk + '_in_word'] = ''

                    # header grouping
                    hdr = {}
                    for key in ('agreement_date','agreement_date_in_word','agreement_date_display','agreement_day_in_word','sp3_date','sp3_date_in_word','sp3_date_display','name_of_debtor','phone_number_of_lolc','sp3_number','contract_number'):
                        hdr[key] = ctx.get(key, '')
                    ctx['header_fields'] = hdr
                except Exception:
                    pass

                repaired_template = _repair_docx_jinja_tags_sp3(template_path, contract_no=contract_number)
                tpl = DocxTemplate(repaired_template)
                # Debug: write ctx keys and full context to log for template troubleshooting
                try:
                    import json
                    logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                    os.makedirs(logs_dir, exist_ok=True)
                    ctx_log = os.path.join(logs_dir, 'uv_sp3_ctx.log')
                    with open(ctx_log, 'a', encoding='utf-8') as lf:
                        lf.write(f"[{timezone.now().isoformat()}] Rendering template {repaired_template}\n")
                        try:
                            lf.write('keys: ' + ','.join([str(k) for k in ctx.keys()]) + '\n')
                        except Exception:
                            pass
                        try:
                            lf.write(json.dumps(ctx, default=str, ensure_ascii=False) + '\n')
                        except Exception:
                            try:
                                lf.write(str(ctx) + '\n')
                            except Exception:
                                pass
                except Exception:
                    pass
                try:
                    tpl.render(ctx)
                except Exception as render_e:
                    # Log render failure and context to aid debugging
                    try:
                        logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                        os.makedirs(logs_dir, exist_ok=True)
                        err_log = os.path.join(logs_dir, 'uv_sp3_docx_errors.log')
                        with open(err_log, 'a', encoding='utf-8') as lf:
                            lf.write(f"[{timezone.now().isoformat()}] docxtpl render failed for {repaired_template}: {str(render_e)}\n")
                            try:
                                import traceback as _tb
                                lf.write(_tb.format_exc() + '\n')
                            except Exception:
                                pass
                            try:
                                import json as _json
                                lf.write(_json.dumps(ctx, default=str, ensure_ascii=False) + '\n')
                            except Exception:
                                try:
                                    lf.write(str(ctx) + '\n')
                                except Exception:
                                    pass
                    except Exception:
                        pass
                    # fall through to existing fallback behavior below
                    raise
                fd, outpath = tempfile.mkstemp(suffix='.docx')
                os.close(fd)
                tpl.save(outpath)
            except Exception as render_e:
                logger.error('docxtpl render failed for UV SP3 %s: %s', template_path, str(render_e))
                try:
                    from docx import Document as _DocxDocument
                    doc_fallback = _DocxDocument()
                    doc_fallback.add_heading(f'UV SP3 - {contract_number}', level=1)
                    for k, v in ctx.items():
                        doc_fallback.add_paragraph(f"{k}: {v}")
                    fd, outpath = tempfile.mkstemp(suffix='.docx')
                    os.close(fd)
                    doc_fallback.save(outpath)
                except Exception as docx_e:
                    logger.error('Fallback DOCX generation failed for UV SP3: %s', str(docx_e))
                    return Response({'error': 'Failed to generate DOCX.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            try:
                # Support optional PDF conversion via ?download=pdf
                download = request.query_params.get('download', '').strip().lower()
                if download == 'pdf':
                    pdf_path = os.path.join(tempfile.mkdtemp(prefix='pdf_out_'), f'uv_sp3_{contract_number}.pdf')
                    ok, err = _convert_docx_to_pdf(outpath, pdf_path)
                    if ok:
                        try:
                            with open(pdf_path, 'rb') as fh:
                                pdf_bytes = fh.read()
                            # Centralized document log entry
                            try:
                                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                                os.makedirs(logs_dir, exist_ok=True)
                                central = os.path.join(logs_dir, 'document_generation.log')
                                username = getattr(request, 'user', None)
                                uname = username.username if getattr(username, 'is_authenticated', False) else 'anonymous'
                                size = len(pdf_bytes)
                                with open(central, 'a', encoding='utf-8') as cf:
                                    cf.write(f"[{timezone.now().isoformat()}] user={uname} type=pdf endpoint=uv_sp3 contract={contract_number} file=uv_sp3_{contract_number}.pdf size={size}\n")
                                logger.info('PDF generated: uv_sp3_%s.pdf size=%s user=%s', contract_number, size, uname)
                            except Exception:
                                logger.exception('Failed to write centralized document log (UVSP3 PDF) %s', contract_number)
                            # cleanup PDF temp dir before returning
                            try:
                                pdf_dir = os.path.dirname(pdf_path)
                                if pdf_dir and os.path.exists(pdf_dir):
                                    _safe_rmtree(pdf_dir)
                            except Exception as cleanup_e:
                                logger.warning('Failed to cleanup pdf temp dir %s: %s', pdf_path, str(cleanup_e))
                            response = HttpResponse(pdf_bytes, content_type='application/pdf')
                            download_filename = f'uv_sp3_{contract_number}.pdf'.upper()
                            response['Content-Disposition'] = f'attachment; filename="{download_filename}"'
                            try:
                                ip = request.META.get('HTTP_X_FORWARDED_FOR') or request.META.get('REMOTE_ADDR') or ''
                                if ip and ',' in ip:
                                    ip = ip.split(',')[0].strip()
                                ua = request.META.get('HTTP_USER_AGENT', '')
                                user = getattr(request, 'user', None)
                                uid = getattr(user, 'id', None) if user and getattr(user, 'is_authenticated', False) else None
                                uname = getattr(user, 'username', None) if user and getattr(user, 'is_authenticated', False) else _resolve_username(request) or ''
                                email = getattr(user, 'email', None) if user and getattr(user, 'is_authenticated', False) else ''
                                DownloadLog.objects.create(
                                    user_id=uid,
                                    username=uname,
                                    email=email,
                                    file_type='uv',
                                    file_identifier=str(contract_number),
                                    filename=download_filename,
                                    ip_address=ip,
                                    user_agent=ua,
                                    success=True,
                                    file_size=len(pdf_bytes),
                                    method='stream',
                                )
                            except Exception:
                                logger.exception('Failed to write DownloadLog for UV SP3 PDF %s', contract_number)
                            return response
                        except Exception as e2:
                            logger.error('Failed to return UV SP3 PDF for %s: %s', contract_number, str(e2))
                            return Response({'error': 'Failed to return PDF.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                    else:
                        logger.error('UV SP3 PDF conversion failed for %s: %s', contract_number, str(err))
                        return Response({'error': 'PDF conversion failed', 'detail': str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                else:
                    try:
                        with open(outpath, 'rb') as fh:
                            docx_bytes = fh.read()
                        # Log DOCX generation
                        try:
                            logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                            os.makedirs(logs_dir, exist_ok=True)
                            docx_log = os.path.join(logs_dir, 'docx_generation.log')
                            username = getattr(request, 'user', None)
                            uname = username.username if getattr(username, 'is_authenticated', False) else 'anonymous'
                            size = len(docx_bytes)
                            with open(docx_log, 'a', encoding='utf-8') as lf:
                                lf.write(f"[{timezone.now().isoformat()}] user={uname} endpoint=uv_sp3 contract={contract_number} file=uv_suratp3_template_{contract_number}.docx size={size}\n")
                            logger.info('DOCX generated: uv_suratp3_template_%s.docx size=%s user=%s', contract_number, size, uname)
                            try:
                                central = os.path.join(logs_dir, 'document_generation.log')
                                with open(central, 'a', encoding='utf-8') as cf:
                                    cf.write(f"[{timezone.now().isoformat()}] user={uname} type=docx endpoint=uv_sp3 contract={contract_number} file=uv_suratp3_template_{contract_number}.docx size={size}\n")
                            except Exception:
                                logger.exception('Failed to write centralized document log (UVSP3 DOCX) %s', contract_number)
                        except Exception:
                            logger.exception('Failed to write DOCX generation log for UV SP3 %s', contract_number)
                        response = HttpResponse(docx_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        download_filename = f'uv_sp3_{contract_number}.docx'.upper()
                        response['Content-Disposition'] = f'attachment; filename="{download_filename}"'
                        try:
                            ip = request.META.get('HTTP_X_FORWARDED_FOR') or request.META.get('REMOTE_ADDR') or ''
                            if ip and ',' in ip:
                                ip = ip.split(',')[0].strip()
                            ua = request.META.get('HTTP_USER_AGENT', '')
                            user = getattr(request, 'user', None)
                            uid = getattr(user, 'id', None) if user and getattr(user, 'is_authenticated', False) else None
                            uname = getattr(user, 'username', None) if user and getattr(user, 'is_authenticated', False) else _resolve_username(request) or ''
                            email = getattr(user, 'email', None) if user and getattr(user, 'is_authenticated', False) else ''
                            DownloadLog.objects.create(
                                user_id=uid,
                                username=uname,
                                email=email,
                                file_type='uv',
                                file_identifier=str(contract_number),
                                filename=download_filename,
                                ip_address=ip,
                                user_agent=ua,
                                success=True,
                                file_size=len(docx_bytes),
                                method='stream',
                            )
                        except Exception:
                            logger.exception('Failed to write DownloadLog for UV SP3 DOCX %s', contract_number)
                        return response
                    except Exception as e2:
                        logger.error('Failed to return UV SP3 DOCX for %s: %s', contract_number, str(e2))
                        return Response({'error': 'Failed to return DOCX.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            finally:
                try:
                    if outpath and os.path.exists(outpath):
                        os.remove(outpath)
                except Exception:
                    pass

        except Exception as e:
            tb = traceback.format_exc()
            logger.error('Error generating UV SP3 DOCX for %s: %s', contract_number, tb)
            if getattr(settings, 'DEBUG', False):
                return Response({'error': str(e), 'traceback': tb}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            return Response({'error': 'Internal server error while generating UV SP3 DOCX. Check server logs.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class UVSP3ListView(APIView):
    """Return rows from `uv_sp3` table for the frontend UV SP3 list/table."""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM uv_sp3")
                cols_meta = cursor.fetchall()
                cols = [r[0] for r in cols_meta]
                # Select explicit columns to preserve order
                # Choose a safe column to ORDER BY based on available columns
                preferred_order_cols = ['created_at', 'created', 'updated_at', 'id']
                order_col = next((c for c in preferred_order_cols if c in cols), None)
                if not order_col:
                    order_col = cols[0] if cols else '1'
                sql = f"SELECT {', '.join(cols)} FROM uv_sp3 ORDER BY {order_col} DESC"
                cursor.execute(sql)
                rows = cursor.fetchall()
                results = []
                for r in rows:
                    results.append({cols[i]: r[i] for i in range(len(cols))})
            return Response({'rows': results, 'columns': cols}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('uv_sp3 list failed')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    def post(self, request):
        """Create a new row in `uv_sp3` by inserting any provided keys that match
        the table columns. This allows the frontend form to save into uv_sp3
        without requiring a strict schema mapping here.
        """
        try:
            payload = request.data or {}
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM uv_sp3")
                cols_meta = cursor.fetchall()
                cols = [r[0] for r in cols_meta]

                # Build a data_map only for actual table columns
                data_map = {}
                cols_lookup = {c.lower(): c for c in cols}

                # Accept direct top-level payload keys that exactly match column names
                for c in cols:
                    if c in payload:
                        val = payload[c]
                        # normalize ISO datetime strings for created_at/updated_at
                        if isinstance(val, str) and ('T' in val and (val.endswith('Z') or '+' in val or '-' in val[10:])) and c in ('created_at', 'updated_at'):
                            try:
                                from datetime import datetime
                                s = val.replace('Z', '+00:00')
                                dt = datetime.fromisoformat(s)
                                val = dt.strftime('%Y-%m-%d %H:%M:%S')
                            except Exception:
                                try:
                                    val = val.split('T')[0] + ' ' + val.split('T')[1].split('+')[0].split('-')[0].split('Z')[0]
                                except Exception:
                                    pass
                        data_map[c] = val

                # Merge nested payload sections (frontend sends many fields nested under
                # contract_data, header_fields, debtor, collateral_data, form_state, extra_fields)
                nested_sections = ['bm_data', 'branch_data', 'contract_data', 'debtor', 'debtor_data', 'collateral_data', 'header_fields', 'form_state', 'extra_fields']
                for section in nested_sections:
                    src = payload.get(section)
                    if isinstance(src, dict):
                        for k, v in src.items():
                            key = str(k).lower()
                            if key in cols_lookup:
                                data_map[cols_lookup[key]] = v

                # Accept frontend-provided created_by if present
                client_created_by = payload.get('created_by')
                if client_created_by and 'created_by' in cols and 'created_by' not in data_map:
                    data_map['created_by'] = client_created_by

                # If created_at exists but not provided, set to now
                if 'created_at' in cols and 'created_at' not in data_map:
                    data_map['created_at'] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')

                # Ensure created_by is populated server-side when not provided by frontend
                now = timezone.now()
                username = _resolve_username(request)
                if 'created_by' in cols and 'created_by' not in data_map:
                    data_map['created_by'] = username or ''

                # Coerce empty-string values for numeric columns to 0 to avoid
                # MySQL "Incorrect integer value: ''" errors when clients send
                # empty strings for numeric fields (common from HTML inputs).
                field_type_map = {row[0]: row[1].lower() for row in cols_meta}
                try:
                    for col_name, col_type in field_type_map.items():
                        if col_name in data_map:
                            v = data_map[col_name]
                            if isinstance(v, str) and v.strip() == '':
                                if any(t in col_type for t in ('int', 'decimal', 'float', 'double')):
                                    # prefer float for decimal-like types
                                    if any(t in col_type for t in ('decimal', 'float', 'double')):
                                        data_map[col_name] = 0.0
                                    else:
                                        data_map[col_name] = 0
                except Exception:
                    # swallow coercion errors; we'll still apply default-filling below
                    pass

                # Ensure NOT NULL columns without defaults receive a safe fallback to avoid MySQL 1364
                field_type_map = {row[0]: row[1].lower() for row in cols_meta}
                for col_row in cols_meta:
                    field_name = col_row[0]
                    field_type = col_row[1].lower()
                    is_nullable = col_row[2]
                    default_val = col_row[4]
                    if field_name in data_map:
                        continue
                    # skip auto-managed or intentionally omitted fields
                    if field_name in ('id', 'created_by', 'created_at', 'update_at'):
                        continue
                    if is_nullable == 'NO' and default_val is None:
                        if any(t in field_type for t in ('int', 'decimal', 'float', 'double')):
                            # prefer float for decimal-like types
                            if any(t in field_type for t in ('decimal', 'float', 'double')):
                                data_map[field_name] = 0.0
                            else:
                                data_map[field_name] = 0
                        elif any(t in field_type for t in ('date', 'timestamp', 'datetime')):
                            data_map[field_name] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
                        else:
                            # use empty string for text-like fields to avoid insertion errors
                            data_map[field_name] = ''

                if not data_map:
                    return Response({'error': 'No valid columns provided for uv_sp3 insert.'}, status=status.HTTP_400_BAD_REQUEST)

                cols_sql = ', '.join([f"`{c}`" for c in data_map.keys()])
                placeholders = ', '.join(['%s'] * len(data_map))
                params = list(data_map.values())
                sql = f"INSERT INTO uv_sp3 ({cols_sql}) VALUES ({placeholders})"
                cursor.execute(sql, params)
                new_id = getattr(cursor, 'lastrowid', None)
                # fetch the newly created row if possible
                if new_id and 'id' in cols:
                    cursor.execute(f"SELECT {', '.join(cols)} FROM uv_sp3 WHERE id=%s", [new_id])
                    row = cursor.fetchone()
                else:
                    order_col = 'created_at' if 'created_at' in cols else cols[0]
                    cursor.execute(f"SELECT {', '.join(cols)} FROM uv_sp3 ORDER BY {order_col} DESC LIMIT 1")
                    row = cursor.fetchone()
                result = {cols[i]: row[i] for i in range(len(cols))} if row else {}
            return Response({'row': result}, status=status.HTTP_201_CREATED)
        except Exception as e:
            logger.exception('uv_sp3 insert failed')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)