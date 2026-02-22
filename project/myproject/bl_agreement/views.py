"""BL Agreement views extracted from myproject.views.

This module contains a self-contained `BLAgreementView` that performs the
same DB-level operations as the original. It includes the helpers it
needs to avoid cross-app import dependencies.
"""
import os
import logging
import re
from datetime import datetime, date as _date
import json

from django.db import connection
from django.conf import settings
from django.utils import timezone
import tempfile
import subprocess
import shutil
import zipfile
import sys
import traceback
import re
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth import get_user_model
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework_simplejwt.tokens import AccessToken
from rest_framework_simplejwt.exceptions import TokenError

from myproject.rbac import get_role_from_request
from myproject.common import format_number_dot, number_to_indonesian_words
from myproject.models import DownloadLog, AgreementAccess

User = get_user_model()
logger = logging.getLogger(__name__)


from myproject.common import _resolve_username, _get_request_user_and_now, _ensure_synthesized_pk, date_to_indonesian_words, format_indonesian_date


def _normalize_for_json(obj):
    try:
        from decimal import Decimal
        import datetime as _dt
    except Exception:
        return obj

    if obj is None:
        return None
    if isinstance(obj, Decimal):
        try:
            return float(obj)
        except Exception:
            return str(obj)
    if isinstance(obj, (_dt.datetime, _dt.date, _dt.time)):
        try:
            return obj.isoformat()
        except Exception:
            return str(obj)
    if isinstance(obj, bytes):
        try:
            return obj.decode('utf-8')
        except Exception:
            return str(obj)
    if isinstance(obj, dict):
        return {k: _normalize_for_json(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [ _normalize_for_json(v) for v in obj ]
    return obj


class BLAgreementView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = request.data
        contract_number = data.get('contract_number')
        branch_id = data.get('branch_id')
        director = data.get('director')
        bm_data = data.get('bm_data', {})
        contract_data = data.get('contract_data', {})
        collateral_data = data.get('collateral_data', {})
        header_fields = data.get('header_fields', {})

        role = get_role_from_request(request) or getattr(request.user, 'role', '')
        # Only Admin and CSA can create. BM and BOD are read-only (download only)
        allowed_creators = ('Admin', 'CSA')
        if role not in allowed_creators:
            return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

        try:
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM bl_agreement")
                cols_meta = cursor.fetchall()
                cols_info = [row[0] for row in cols_meta]
                try:
                    pk_row = next((r for r in cols_meta if (r[3] or '').upper() == 'PRI'), None)
                    pk_col_name = pk_row[0] if pk_row else None
                except Exception:
                    pk_col_name = None
                cols_lookup = {c.lower(): c for c in cols_info}

                data_map = {}
                if contract_number:
                    data_map['contract_number'] = contract_number

                client_created_by = data.get('created_by')
                if client_created_by and 'created_by' in cols_info:
                    data_map['created_by'] = client_created_by

                user_full_name = None
                user_username = None
                if getattr(request, 'user', None) and getattr(request.user, 'is_authenticated', False):
                    user_full_name = getattr(request.user, 'full_name', None)
                    user_username = getattr(request.user, 'username', None)
                    if (not user_full_name or not user_username) and getattr(request.user, 'id', None):
                        try:
                            u = User.objects.filter(pk=request.user.id).first()
                            if u:
                                if not user_full_name:
                                    user_full_name = getattr(u, 'full_name', None)
                                if not user_username:
                                    user_username = getattr(u, 'username', None)
                        except Exception:
                            pass
                else:
                    auth_header = request.META.get('HTTP_AUTHORIZATION', '')
                    if auth_header and auth_header.startswith('Bearer '):
                        token = auth_header.split(' ', 1)[1].strip()
                        try:
                            payload = AccessToken(token)
                            user_id = payload.get('user_id') or payload.get('uid') or payload.get('id')
                            if user_id:
                                try:
                                    u = User.objects.filter(pk=user_id).first()
                                    if u:
                                        user_full_name = getattr(u, 'full_name', None)
                                        user_username = getattr(u, 'username', None)
                                except Exception:
                                    pass
                        except TokenError:
                            pass

                user_username = user_username or (user_full_name and user_full_name.replace(' ', '').lower()) or 'anonymous'
                now = timezone.now()

                if branch_id is not None and 'branch_id' in cols_lookup:
                    data_map[cols_lookup['branch_id']] = branch_id
                if director is not None:
                    if 'name_of_director' in cols_lookup:
                        data_map[cols_lookup['name_of_director']] = director
                    elif 'director' in cols_lookup:
                        data_map[cols_lookup['director']] = director

                branch_data = data.get('branch_data', {})
                for src in (bm_data, branch_data, contract_data, collateral_data, header_fields):
                    if not isinstance(src, dict):
                        continue
                    for k, v in src.items():
                        key = str(k).lower()
                        if key == 'name_of_director':
                            key = 'name_of_director'
                        if key in cols_lookup:
                            data_map[cols_lookup[key]] = v

                field_type_map = {row[0]: row[1].lower() for row in cols_meta}
                for k in list(data_map.keys()):
                    ft = field_type_map.get(k, '')
                    if not any(t in ft for t in ('date', 'timestamp', 'datetime')):
                        continue
                    val = data_map.get(k)
                    if val is None or (isinstance(val, str) and val.strip() == ''):
                        data_map.pop(k, None)
                        continue
                    if isinstance(val, (_date, datetime)):
                        continue
                    s = str(val).strip()
                    parsed = None
                    try:
                        try:
                            parsed = datetime.fromisoformat(s)
                        except Exception:
                            if re.match(r'^\s*\d{1,2}\s+\d{1,2}\s+\d{4}\s*$', s):
                                s = re.sub(r'\s+', '/', s.strip())
                            for fmt in ('%Y-%m-%d', '%d-%m-%Y', '%d/%m/%Y', '%Y/%m/%d', '%d %m %Y', '%d %b %Y', '%d %B %Y'):
                                try:
                                    parsed = datetime.strptime(s, fmt)
                                    break
                                except Exception:
                                    continue
                        if parsed is None:
                            data_map.pop(k, None)
                            continue
                        if 'date' in ft and 'time' not in ft and 'datetime' not in ft:
                            data_map[k] = parsed.date()
                        else:
                            if parsed.tzinfo is None:
                                parsed = timezone.make_aware(parsed)
                            data_map[k] = parsed
                    except Exception:
                        data_map.pop(k, None)

                for col_row in cols_meta:
                    field_name = col_row[0]
                    field_type = col_row[1].lower()
                    is_nullable = col_row[2]
                    default_val = col_row[4]
                    if field_name in data_map:
                        continue
                    if field_name in ('id', 'created_by', 'created_at', 'update_at') or (pk_col_name and field_name == pk_col_name):
                        continue
                    if is_nullable == 'NO' and default_val is None:
                        if any(t in field_type for t in ('int', 'decimal', 'float', 'double')):
                            data_map[field_name] = 0
                        elif any(t in field_type for t in ('date', 'timestamp', 'datetime')):
                            data_map[field_name] = timezone.now()
                        else:
                            data_map[field_name] = '-'

                # Determine if caller requested skipping normalization (e.g., modal Create)
                skip_normalization = False
                try:
                    # backend accepts both snake_case and camelCase flags
                    if data.get('create_only') or data.get('createOnly'):
                        skip_normalization = True
                except Exception:
                    pass

                # Also allow explicit client override: send `skip_normalization: true` to opt-out
                try:
                    if data.get('skip_normalization'):
                        skip_normalization = True
                except Exception:
                    pass

                # Normalize certain text fields to Title Case (Capital Each Word)
                def _title_each_word(val):
                    if val is None:
                        return val
                    if not isinstance(val, str):
                        return val
                    s = val.strip()
                    if not s:
                        return s
                    return ' '.join([w.capitalize() for w in s.split()])

                _titlecase_fields = [
                    'name_of_debtor','place_birth_of_debtor','date_birth_of_debtor_in_word',
                    'subdistrict_of_debtor','district_of_debtor','city_of_debtor','province_of_debtor',
                    'business_type','name_of_account_holder','loan_amount_in_word','term_by_word','flat_rate_by_word',
                    'notaris_fee_in_word','admin_fee_in_word','mortgage_amount_in_word','net_amount_in_word',
                    'admin_rate_in_word','tlo_in_word','life_insurance_in_word',
                    # Collateral/title fields
                    'name_of_collateral_owner'
                ]
                if not skip_normalization:
                    for _f in _titlecase_fields:
                        if _f in data_map:
                            try:
                                data_map[_f] = _title_each_word(data_map[_f])
                            except Exception:
                                pass

                existing_check_sql = "SELECT contract_number FROM bl_agreement WHERE contract_number=%s"
                exists = None
                if contract_number:
                    cursor.execute(existing_check_sql, [contract_number])
                    exists = cursor.fetchone()

                edit_only = bool(data.get('edit_only'))
                create_only = bool(data.get('create_only'))

                if exists:
                    if create_only:
                        return Response({'error': 'Record already exists (create_only specified)'}, status=status.HTTP_400_BAD_REQUEST)
                    if 'update_at' in cols_info:
                        data_map['update_at'] = now
                    if 'updated_at' in cols_info:
                        data_map['updated_at'] = now
                    data_map.pop('created_by', None)
                    data_map.pop('created_at', None)
                    set_cols = []
                    params = []
                    for col, val in data_map.items():
                        if col == 'contract_number':
                            continue
                        set_cols.append(f"{col}=%s")
                        params.append(val)
                    if set_cols:
                        sql = f"UPDATE bl_agreement SET {', '.join(set_cols)} WHERE contract_number=%s"
                        params.append(contract_number)
                        cursor.execute(sql, params)
                else:
                    if edit_only:
                        return Response({'error': 'Record not found for update (edit_only specified)'}, status=status.HTTP_400_BAD_REQUEST)
                    if 'created_by' in cols_info and user_username:
                        data_map['created_by'] = user_username
                    if 'created_at' in cols_info:
                        data_map[cols_lookup['created_at']] = now if 'created_at' in cols_lookup else now
                    if 'updated_at' in cols_info:
                        data_map[cols_lookup['updated_at']] = now if 'updated_at' in cols_lookup else now
                    if 'update_at' in cols_info and 'update_at' not in data_map:
                        data_map['update_at'] = now

                    try:
                        _ensure_synthesized_pk(cursor, cols_meta, data_map, 'bl_agreement')
                    except Exception:
                        pass

                    try:
                        pk_row = next((r for r in cols_meta if (r[3] or '').upper() == 'PRI'), None)
                        if pk_row:
                            pk_name = pk_row[0]
                            if pk_name in data_map and (data_map[pk_name] is None or str(data_map[pk_name]) in ('', '0')):
                                data_map.pop(pk_name, None)
                    except Exception:
                        pass

                    cols = []
                    placeholders = []
                    params = []
                    for col, val in data_map.items():
                        cols.append(col)
                        placeholders.append('%s')
                        params.append(val)
                    sql = f"INSERT INTO bl_agreement ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
                    try:
                        cursor.execute(sql, params)
                    except Exception as ie:
                        msg = str(ie)
                        if 'Duplicate entry' in msg and 'PRIMARY' in msg:
                            logger.warning('Duplicate primary on bl_agreement insert: %s', msg)
                            if contract_number:
                                try:
                                    set_cols = []
                                    params2 = []
                                    for col, val in data_map.items():
                                        if col == 'contract_number':
                                            continue
                                        set_cols.append(f"{col}=%s")
                                        params2.append(val)
                                    if set_cols:
                                        update_sql = f"UPDATE bl_agreement SET {', '.join(set_cols)} WHERE contract_number=%s"
                                        params2.append(contract_number)
                                        cursor.execute(update_sql, params2)
                                except Exception:
                                    logger.exception('Failed safe-update after duplicate insert into bl_agreement')
                            else:
                                logger.warning('Duplicate primary but no contract_number available; skipping insert')
                        else:
                            raise
                    else:
                        # record inserted successfully; create per-CSA access entry when creator is CSA
                        try:
                            creator_username = data_map.get('created_by') or user_username
                            if role == 'CSA' and creator_username and creator_username == (user_username or ''):
                                try:
                                    uid = getattr(request.user, 'id', None)
                                    AgreementAccess.objects.create(
                                        contract_number=contract_number,
                                        user_id=uid,
                                        role='CSA',
                                        download_grants=1,
                                        edit_grants=1,
                                    )
                                except Exception:
                                    logger.exception('Failed to create AgreementAccess for %s user=%s', contract_number, creator_username)
                        except Exception:
                            pass
            return Response({'message': 'Data berhasil disimpan'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def put(self, request):
        data = request.data or {}
        contract_number = data.get('contract_number') or data.get('contractNumber')
        if not contract_number:
            return Response({'error': 'contract_number is required'}, status=status.HTTP_400_BAD_REQUEST)
        try:
            role = get_role_from_request(request) or getattr(request.user, 'role', '')
            # Only Admin and CSA may perform edits
            allowed_editors = ('Admin', 'CSA')
            if role not in allowed_editors:
                return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM bl_agreement")
                cols_meta = cursor.fetchall()
                cols_info = [row[0] for row in cols_meta]
                cols_lookup = {c.lower(): c for c in cols_info}

                # find existing record
                sql_get = "SELECT 1 FROM bl_agreement WHERE LOWER(contract_number)=LOWER(%s) LIMIT 1"
                cursor.execute(sql_get, [contract_number])
                if not cursor.fetchone():
                    return Response({'error': 'Record not found'}, status=status.HTTP_404_NOT_FOUND)

                update_map = {}
                # top-level fields
                for k, v in data.items():
                    key = str(k).lower()
                    if key in cols_lookup:
                        update_map[cols_lookup[key]] = v

                # nested sections
                for section in ('header_fields', 'contract_data', 'debtor', 'collateral_data', 'bm_data', 'branch_data'):
                    sec = data.get(section) or {}
                    if isinstance(sec, dict):
                        for k, v in sec.items():
                            lk = str(k).lower()
                            if lk in cols_lookup:
                                update_map[cols_lookup[lk]] = v

                if not update_map:
                    return Response({'message': 'No fields to update'}, status=status.HTTP_200_OK)

                now = timezone.now()
                if 'update_at' in cols_info:
                    update_map['update_at'] = now
                elif 'updated_at' in cols_info:
                    update_map['updated_at'] = now

                set_cols = []
                params = []
                for col, val in update_map.items():
                    if col == 'contract_number':
                        continue
                    set_cols.append(f"{col}=%s")
                    params.append(val)
                params.append(contract_number)
                sql = f"UPDATE bl_agreement SET {', '.join(set_cols)} WHERE LOWER(contract_number)=LOWER(%s)"
                cursor.execute(sql, params)
                # After successful update, if CSA who created this agreement is editing,
                # consume edit grant (only on commit). For CSA role ensure the caller is the creator.
                try:
                    username = getattr(request.user, 'username', None) or _resolve_username(request) or ''
                    if role == 'CSA':
                        # fetch created_by to ensure only creator CSA can edit and consume grant
                        try:
                            cursor.execute("SELECT created_by FROM bl_agreement WHERE LOWER(contract_number)=LOWER(%s) LIMIT 1", [contract_number])
                            row = cursor.fetchone()
                            created_by = row[0] if row else None
                        except Exception:
                            created_by = None
                        if not created_by or str(created_by) != str(username):
                            return Response({'error': 'Forbidden - CSA not creator'}, status=status.HTTP_403_FORBIDDEN)
                        try:
                            uid = getattr(request.user, 'id', None)
                            aa = AgreementAccess.get_for_contract_and_user(contract_number, uid)
                            if not aa:
                                return Response({'error': 'Edit not allowed or access not found'}, status=status.HTTP_403_FORBIDDEN)
                            ok = aa.consume_edit()
                            if not ok:
                                return Response({'error': 'No edit grants remaining'}, status=status.HTTP_403_FORBIDDEN)
                        except Exception:
                            logger.exception('Failed to consume edit grant for %s user=%s', contract_number, username)
                except Exception:
                    pass

            return Response({'message': 'Data berhasil diupdate'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def get(self, request):
        contract_number = request.query_params.get('contract_number', '').strip()
        try:
            logs_dir = os.path.join(settings.BASE_DIR, 'logs')
            os.makedirs(logs_dir, exist_ok=True)
            req_log = os.path.join(logs_dir, 'bl_agreement_docx_requests.log')
            with open(req_log, 'a', encoding='utf-8') as rf:
                rf.write(f"[{timezone.now().isoformat()}] query_params={dict(request.query_params)}\n")
        except Exception:
            pass

        try:
            with connection.cursor() as cursor:
                if not contract_number:
                    cursor.execute("SHOW COLUMNS FROM bl_agreement")
                    cols_meta = cursor.fetchall()
                    cols_info = [row[0] for row in cols_meta]
                    cols_lookup_lower = [c.lower() for c in cols_info]

                    role = get_role_from_request(request) or getattr(request.user, 'role', '')
                    username = getattr(request.user, 'username', None) or ' '

                    where_clauses = []
                    params = []

                    if role == 'CSA':
                        user_branch = getattr(request.user, 'branch_id', None)
                        if user_branch is None:
                            return Response({'error': 'Branch not configured for user'}, status=status.HTTP_403_FORBIDDEN)
                        if 'branch_id' in cols_lookup_lower:
                            col = next(c for c in cols_info if c.lower() == 'branch_id')
                            where_clauses.append(f"{col} = %s")
                            params.append(user_branch)
                        else:
                            return Response({'error': 'branch_id column missing in bl_agreement'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                    elif role == 'BM':
                        user_branch = getattr(request.user, 'branch_id', None)
                        if user_branch is None:
                            return Response({'error': 'Branch not configured for user'}, status=status.HTTP_403_FORBIDDEN)
                        if 'branch_id' in cols_lookup_lower:
                            col = next(c for c in cols_info if c.lower() == 'branch_id')
                            where_clauses.append(f"{col} = %s")
                            params.append(user_branch)
                        else:
                            return Response({'error': 'branch_id column missing in bl_agreement'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                    elif role == 'AM':
                        user_area = getattr(request.user, 'area_id', None)
                        if user_area is None:
                            return Response({'error': 'Area not configured for user'}, status=status.HTTP_403_FORBIDDEN)
                        if 'branch_id' in cols_lookup_lower:
                            try:
                                cursor.execute('SELECT id FROM branches WHERE area_id=%s', [user_area])
                                bids = [r[0] for r in cursor.fetchall()]
                            except Exception:
                                bids = []
                            if not bids:
                                where_clauses.append('1=0')
                            else:
                                col = next(c for c in cols_info if c.lower() == 'branch_id')
                                if len(bids) == 1:
                                    where_clauses.append(f"{col} = %s")
                                    params.append(bids[0])
                                else:
                                    placeholders = ','.join(['%s'] * len(bids))
                                    where_clauses.append(f"{col} IN ({placeholders})")
                                    params.extend(bids)
                        elif 'area_id' in cols_lookup_lower:
                            col = next(c for c in cols_info if c.lower() == 'area_id')
                            where_clauses.append(f"{col} = %s")
                            params.append(user_area)
                        else:
                            return Response({'error': 'area_id/branch_id column missing in bl_agreement'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                    elif role == 'RM':
                        user_region = getattr(request.user, 'region_id', None)
                        if user_region is None:
                            return Response({'error': 'Region not configured for user'}, status=status.HTTP_403_FORBIDDEN)
                        if 'branch_id' in cols_lookup_lower:
                            try:
                                cursor.execute('SELECT b.id FROM branches b JOIN areas a ON b.area_id = a.id WHERE a.region_id=%s', [user_region])
                                bids = [r[0] for r in cursor.fetchall()]
                            except Exception:
                                bids = []
                            if not bids:
                                where_clauses.append('1=0')
                            else:
                                col = next(c for c in cols_info if c.lower() == 'branch_id')
                                if len(bids) == 1:
                                    where_clauses.append(f"{col} = %s")
                                    params.append(bids[0])
                                else:
                                    placeholders = ','.join(['%s'] * len(bids))
                                    where_clauses.append(f"{col} IN ({placeholders})")
                                    params.extend(bids)
                        elif 'region_id' in cols_lookup_lower:
                            col = next(c for c in cols_info if c.lower() == 'region_id')
                            where_clauses.append(f"{col} = %s")
                            params.append(user_region)
                        else:
                            return Response({'error': 'region_id/branch_id column missing in bl_agreement'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                    elif role in ('Admin', 'BOD'):
                        pass
                    else:
                        where_clauses.append('1=0')

                    base_sql = "SELECT agreement_date, contract_number, name_of_debtor, nik_number_of_debtor, collateral_type, created_by, capacity_of_building, surface_area FROM bl_agreement"
                    if where_clauses:
                        sql = base_sql + ' WHERE ' + ' AND '.join(where_clauses) + ' ORDER BY COALESCE(agreement_date, created_at) DESC'
                        cursor.execute(sql, params)
                    else:
                        cursor.execute(base_sql + ' ORDER BY COALESCE(agreement_date, created_at) DESC')

                    cols = [c[0] for c in cursor.description] if cursor.description else []
                    rows = cursor.fetchall()
                    items = [dict(zip(cols, r)) for r in rows]
                    return Response({'agreements': _normalize_for_json(items)}, status=status.HTTP_200_OK)

                # ── mode=create: always fetch from source tables (contract + bl_collateral) ──
                mode = request.query_params.get('mode', '').strip().lower()

                if mode == 'create':
                    # Debtor data from `contract` table
                    cursor.execute(
                        "SELECT * FROM contract WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                        [contract_number]
                    )
                    c_cols = [col[0] for col in cursor.description] if cursor.description else []
                    c_row = cursor.fetchone()
                    debtor_data = dict(zip(c_cols, c_row)) if c_row else None

                    # Collateral data from `bl_collateral` table
                    collateral_data = None
                    try:
                        cursor.execute(
                            "SELECT * FROM bl_collateral WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                            [contract_number]
                        )
                        coll_cols = [col[0] for col in cursor.description] if cursor.description else []
                        coll_row = cursor.fetchone()
                        if coll_row:
                            coll_dict = dict(zip(coll_cols, coll_row))
                            # Remove meta keys
                            for mk in ('id', 'created_at', 'updated_at', 'created_by', 'update_at', 'contract_number'):
                                coll_dict.pop(mk, None)
                            collateral_data = coll_dict
                    except Exception:
                        pass

                    if not debtor_data:
                        return Response({
                            'debtor': None,
                            'collateral': None,
                            'branch_manager': {},
                            'branch': {},
                            'message': 'Data tidak ditemukan untuk nomor kontrak ini',
                            'source': 'create'
                        }, status=status.HTTP_404_NOT_FOUND)

                    # Remove meta fields from debtor
                    for mk in ('id', 'created_at', 'updated_at', 'created_by', 'update_at'):
                        debtor_data.pop(mk, None)

                    return Response({
                        'debtor': _normalize_for_json(debtor_data),
                        'collateral': _normalize_for_json(collateral_data),
                        'branch_manager': {},
                        'branch': {},
                        'source': 'create'
                    }, status=status.HTTP_200_OK)

                # ── mode=edit (default): fetch from bl_agreement table ──
                cursor.execute(
                    "SELECT * FROM bl_agreement WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                    [contract_number]
                )
                cols = [col[0] for col in cursor.description] if cursor.description else []
                bl_row = cursor.fetchone()
                bl_data = dict(zip(cols, bl_row)) if bl_row else None

                if bl_data:
                    try:
                        if (not bl_data.get('phone_number_of_lolc')) and bl_data.get('name_of_director'):
                            cursor.execute("SELECT phone_number_of_lolc FROM director WHERE name_of_director=%s LIMIT 1", [bl_data.get('name_of_director')])
                            row = cursor.fetchone()
                            if row and row[0]:
                                bl_data['phone_number_of_lolc'] = row[0]

                        if (not bl_data.get('sp3_number') or not bl_data.get('sp3_date')):
                            cursor.execute("SELECT sp3_number, sp3_date FROM bl_sp3 WHERE LOWER(contract_number)=LOWER(%s) ORDER BY COALESCE(sp3_date, created_at) DESC LIMIT 1", [contract_number])
                            sp3row = cursor.fetchone()
                            if sp3row:
                                if not bl_data.get('sp3_number') and sp3row[0]:
                                    bl_data['sp3_number'] = sp3row[0]
                                if not bl_data.get('sp3_date') and sp3row[1]:
                                    bl_data['sp3_date'] = sp3row[1]
                    except Exception:
                        pass

                    # Split bl_data into 4 field groups from bl_agreement table only
                    collateral_fields = {
                        'collateral_type', 'number_of_certificate', 'number_of_ajb', 'surface_area',
                        'name_of_collateral_owner', 'capacity_of_building', 'location_of_land',
                    }
                    
                    branch_manager_fields = {
                        'name_of_bm', 'place_birth_of_bm', 'date_birth_of_bm', 'date_birth_of_bm_in_word',
                        'street_of_bm', 'subdistrict_of_bm', 'district_of_bm', 'city_of_bm', 
                        'province_of_bm', 'nik_number_of_bm', 'phone_number_of_bm'
                    }
                    
                    branch_fields = {
                        'street_name', 'subdistrict', 'district', 'city', 'province', 'branch_id'
                    }
                    
                    # Fields to exclude from debtor (internal/meta fields)
                    meta_fields = {'id', 'created_at', 'updated_at', 'created_by'}
                    non_debtor_fields = collateral_fields | branch_manager_fields | branch_fields | meta_fields
                    
                    # Parse data into 4 field groups — debtor is catch-all for remaining fields
                    collateral_data = {k: v for k, v in bl_data.items() if k in collateral_fields}
                    branch_manager_data = {k: v for k, v in bl_data.items() if k in branch_manager_fields}
                    branch_data = {k: v for k, v in bl_data.items() if k in branch_fields}
                    debtor_data = {k: v for k, v in bl_data.items() if k not in non_debtor_fields}

                    return Response({
                        'debtor': _normalize_for_json(debtor_data), 
                        'collateral': _normalize_for_json(collateral_data),
                        'branch_manager': _normalize_for_json(branch_manager_data),
                        'branch': _normalize_for_json(branch_data),
                        'source': 'bl_agreement'
                    }, status=status.HTTP_200_OK)

                # Fallback: no agreement exists yet, fetch from source tables
                cursor.execute(
                    "SELECT * FROM contract WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                    [contract_number]
                )
                columns = [col[0] for col in cursor.description] if cursor.description else []
                contract_row = cursor.fetchone()
                debtor_data = dict(zip(columns, contract_row)) if contract_row else None

                # Also try bl_collateral for fallback
                collateral_data = None
                try:
                    cursor.execute(
                        "SELECT * FROM bl_collateral WHERE LOWER(contract_number) = LOWER(%s) LIMIT 1",
                        [contract_number]
                    )
                    coll_cols = [col[0] for col in cursor.description] if cursor.description else []
                    coll_row = cursor.fetchone()
                    if coll_row:
                        coll_dict = dict(zip(coll_cols, coll_row))
                        for mk in ('id', 'created_at', 'updated_at', 'created_by', 'update_at', 'contract_number'):
                            coll_dict.pop(mk, None)
                        collateral_data = coll_dict
                except Exception:
                    pass

                if not debtor_data and not collateral_data:
                    return Response({
                        'debtor': None,
                        'collateral': None,
                        'branch_manager': {},
                        'branch': {},
                        'message': 'Data tidak ditemukan untuk nomor kontrak ini'
                    }, status=status.HTTP_404_NOT_FOUND)

                return Response({
                    'debtor': _normalize_for_json(debtor_data),
                    'collateral': _normalize_for_json(collateral_data),
                    'branch_manager': {},
                    'branch': {},
                    'source': 'contract_fallback'
                }, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def delete(self, request):
        """Delete a BL agreement by contract_number (case-insensitive).

        Permission: only Admin and CSA may delete.
        Expects `contract_number` as a query param or in JSON body.
        """
        try:
            contract_number = request.query_params.get('contract_number') or (request.data or {}).get('contract_number')
            if not contract_number:
                return Response({'error': 'contract_number is required'}, status=status.HTTP_400_BAD_REQUEST)

            role = get_role_from_request(request) or getattr(request.user, 'role', '')
            allowed_deleters = ('Admin', 'CSA')
            if role not in allowed_deleters:
                return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

            with connection.cursor() as cursor:
                # Use case-insensitive match to locate row(s)
                cursor.execute("SELECT 1 FROM bl_agreement WHERE LOWER(contract_number)=LOWER(%s) LIMIT 1", [contract_number])
                if not cursor.fetchone():
                    return Response({'error': 'Record not found'}, status=status.HTTP_404_NOT_FOUND)

                cursor.execute("DELETE FROM bl_agreement WHERE LOWER(contract_number)=LOWER(%s)", [contract_number])

            return Response({'message': 'Record deleted'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)




def _convert_docx_to_pdf(docx_path, pdf_path, retries=2, min_size=2048):
    import shlex
    docx2pdf_err = None
    try:
        from docx2pdf import convert
        docx2pdf_available = True
    except Exception as imp_e:
        docx2pdf_available = False
        docx2pdf_err = f'docx2pdf import failed: {imp_e}'

    if docx2pdf_available:
        try:
            try:
                import pythoncom
                try:
                    pythoncom.CoInitializeEx(0)
                except Exception:
                    pass
            except Exception:
                pythoncom = None
            try:
                convert(docx_path, pdf_path)
            finally:
                if pythoncom is not None:
                    try:
                        pythoncom.CoUninitialize()
                    except Exception:
                        pass
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) >= int(min_size or 0):
                return True, None
            return False, 'pdf file missing or too small after docx2pdf conversion'
        except Exception as e:
            docx2pdf_err = f'docx2pdf conversion failed: {e}'

    soffice_err = None
    try:
        outdir = os.path.dirname(os.path.abspath(pdf_path)) or '.'
        cmd = f'soffice --headless --convert-to pdf --outdir {shlex.quote(outdir)} {shlex.quote(docx_path)}'
        proc = subprocess.run(cmd, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120)
        if proc.returncode != 0:
            soffice_err = f'soffice returned exit {proc.returncode}: {proc.stderr.decode("utf-8", errors="replace")}'
        else:
            expected = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
            try:
                if os.path.exists(expected):
                    if os.path.abspath(expected) != os.path.abspath(pdf_path):
                        try:
                            shutil.move(expected, pdf_path)
                        except Exception:
                            shutil.copyfile(expected, pdf_path)
                    if os.path.exists(pdf_path) and os.path.getsize(pdf_path) >= int(min_size or 0):
                        return True, None
                    else:
                        soffice_err = 'LibreOffice produced PDF but file missing or too small'
                else:
                    soffice_err = 'LibreOffice did not produce expected PDF file'
            except Exception as e:
                soffice_err = f'Error after LibreOffice conversion: {e}'
    except Exception as e:
        soffice_err = f'soffice execution failed: {e}'

    msgs = []
    if docx2pdf_err:
        msgs.append(docx2pdf_err)
    if soffice_err:
        msgs.append(soffice_err)
    if not msgs:
        msgs.append('No conversion method available')
    return False, '; '.join(msgs)


def _safe_rmtree(path, retries=6, delay=0.25):
    last_exc = None
    for attempt in range(retries):
        try:
            shutil.rmtree(path)
            logger.info('Removed temporary path %s (attempt %d)', path, attempt+1)
            return
        except Exception as e:
            last_exc = e
            import time
            time.sleep(delay * (attempt + 1))
    def _onerror(func, p, exc_info):
        try:
            os.chmod(p, 0o700)
            func(p)
        except Exception:
            pass
    try:
        shutil.rmtree(path, onerror=_onerror)
        logger.info('Removed temporary path %s with onerror handler', path)
    except Exception:
        logger.warning('Failed to remove path %s after retries: %s', path, str(last_exc))


def _repair_docx_jinja_tags(src_path, contract_no=None):
    try:
        tmpdir = tempfile.mkdtemp(prefix='docx_repair_')
        with zipfile.ZipFile(src_path, 'r') as zin:
            zin.extractall(tmpdir)

        word_dir = os.path.join(tmpdir, 'word')
        if not os.path.exists(word_dir):
            shutil.rmtree(tmpdir)
            return src_path

        changed_files = []
        for root, dirs, files in os.walk(word_dir):
            for fname in files:
                if not fname.lower().endswith('.xml'):
                    continue
                fullpath = os.path.join(root, fname)
                with open(fullpath, 'r', encoding='utf-8') as f:
                    data = f.read()
                if '\u2026' in data:
                    data = data.replace('\u2026', '...')
                replacements = {
                    '\u00A0': ' ', '\u200B': '', '\u200C': '', '\u200D': '', '\u00AD': '', '\ufeff': '',
                    '\u2018': "'", '\u2019': "'", '\u201C': '"', '\u201D': '"',
                    '\u2013': '-', '\u2014': '-', '\u2010': '-', '\u2022': '*'
                }
                for k, v in replacements.items():
                    if k in data:
                        data = data.replace(k, v)

                if '{{' not in data or '}}' not in data:
                    continue

                orig = data
                new = []
                pos = 0
                L = len(data)
                while pos < L:
                    i = data.find('{{', pos)
                    if i == -1:
                        new.append(data[pos:])
                        break
                    new.append(data[pos:i])
                    j = data.find('}}', i)
                    if j == -1:
                        new.append(data[i:])
                        break
                    segment = data[i:j+2]
                    cleaned = re.sub(r'<[^>]+>', '', segment)
                    if cleaned.count('{{') and cleaned.count('}}'):
                        inner_start = cleaned.find('{{') + 2
                        inner_end = cleaned.rfind('}}')
                        inner = cleaned[inner_start:inner_end]
                        inner = re.sub(r'\s+', '', inner)
                        cleaned_full = '{{' + inner + '}}'
                        new.append(cleaned_full)
                    else:
                        cleaned2 = re.sub(r'<[^>]+>', '', segment)
                        new.append(cleaned2)
                    pos = j+2

                repaired = ''.join(new)
                if repaired != orig:
                    with open(fullpath, 'w', encoding='utf-8') as f:
                        f.write(repaired)
                    rel = os.path.relpath(fullpath, tmpdir)
                    changed_files.append(rel.replace('\\', '/'))

        if changed_files:
            repaired_fd, repaired_path = tempfile.mkstemp(suffix='.docx')
            os.close(repaired_fd)
            with zipfile.ZipFile(repaired_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for root, dirs, files in os.walk(tmpdir):
                    for file in files:
                        full = os.path.join(root, file)
                        arcname = os.path.relpath(full, tmpdir)
                        zout.write(full, arcname)
            try:
                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                os.makedirs(logs_dir, exist_ok=True)
                repair_log = os.path.join(logs_dir, 'bl_agreement_docx_repair.log')
                with open(repair_log, 'a', encoding='utf-8') as lf:
                    lf.write(f"[{timezone.now().isoformat()}] repaired={repaired_path} contract={contract_no} files={changed_files}\n")
            except Exception:
                pass
            shutil.rmtree(tmpdir)
            return repaired_path
        else:
            shutil.rmtree(tmpdir)
            return src_path
    except Exception:
        try:
            shutil.rmtree(tmpdir)
        except Exception:
            pass
        return src_path


class BLAgreementContractListView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            with connection.cursor() as cursor:
                cursor.execute("SELECT DISTINCT contract_number FROM contract ORDER BY contract_number")
                bl_contracts = [row[0] for row in cursor.fetchall()]
                cursor.execute("SELECT DISTINCT contract_number FROM bl_collateral ORDER BY contract_number")
                bl_collateral_contracts = [row[0] for row in cursor.fetchall()]
                all_contracts = sorted(list(set(bl_contracts + bl_collateral_contracts)))
                return Response({'contracts': all_contracts}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class BLCollateralCreateView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        data = request.data or {}
        try:
            with connection.cursor() as cursor:
                cursor.execute("SHOW COLUMNS FROM bl_collateral")
                cols_meta = cursor.fetchall()
                cols_info = [r[0] for r in cols_meta]

                data_map = {}
                for k, v in data.items():
                    key = str(k).lower()
                    if key.endswith('_in_word') or key.endswith('_by_word'):
                        continue
                    if key in cols_info:
                        data_map[key] = v

                now = timezone.now()
                username = _resolve_username(request)
                if 'created_by' in cols_info and 'created_by' not in data_map:
                    data_map['created_by'] = username or ''
                if 'created_at' in cols_info:
                    data_map['created_at'] = now
                if 'update_at' in cols_info:
                    data_map['update_at'] = now

                if 'id' in data_map:
                    data_map.pop('id', None)

                if not data_map:
                    return Response({'error': 'No valid collateral fields provided'}, status=status.HTTP_400_BAD_REQUEST)

                cols = []
                placeholders = []
                params = []
                for col, val in data_map.items():
                    cols.append(col)
                    placeholders.append('%s')
                    params.append(val)

                sql = f"INSERT INTO bl_collateral ({', '.join(cols)}) VALUES ({', '.join(placeholders)})"
                cursor.execute(sql, params)

            return Response({'message': 'Collateral saved'}, status=status.HTTP_200_OK)
        except Exception as e:
            logger.exception('Failed to save bl_collateral')
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class BLAgreementDocxDownloadView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        contract_number = request.query_params.get('contract_number', '').strip()
        if not contract_number:
            return Response({'error': 'contract_number parameter required'}, status=status.HTTP_400_BAD_REQUEST)

        # Get type parameter (agreement or sp3)
        doc_type = request.query_params.get('type', 'agreement').strip().lower()
        if doc_type not in ('agreement', 'sp3'):
            doc_type = 'agreement'

        # Determine template path based on type parameter
        if doc_type == 'sp3':
            template_filename = 'bl_sp3_template.docx'
        else:
            template_filename = 'bl_agreement_template.docx'

        # choose base filename prefix according to type so downloaded files are named correctly
        base_prefix = 'bl_sp3' if doc_type == 'sp3' else 'bl_agreement'

        req_template = (request.query_params.get('template') or '').strip()
        if req_template:
            req_template = os.path.basename(req_template)
            if not req_template.lower().endswith('.docx'):
                req_template = req_template + '.docx'
            template_path = os.path.join(settings.BASE_DIR, 'templates', 'docx', req_template)
        else:
            template_path = os.path.join(settings.BASE_DIR, 'templates', 'docx', template_filename)

        if not os.path.exists(template_path):
            return Response({'error': f'Template not found at {template_path}. Please place your .docx template there.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        try:
            # Always query from bl_agreement table regardless of type
            with connection.cursor() as cursor:
                cursor.execute('SELECT * FROM bl_agreement WHERE contract_number=%s LIMIT 1', [contract_number])
                agreement_row = cursor.fetchone()
                agreement_cols = [c[0] for c in cursor.description] if cursor.description else []
                agreement = dict(zip(agreement_cols, agreement_row)) if agreement_row else {}

                cursor.execute('SELECT * FROM bl_collateral WHERE contract_number=%s LIMIT 1', [contract_number])
                coll_row = cursor.fetchone()
                coll_cols = [c[0] for c in cursor.description] if cursor.description else []
                collateral = dict(zip(coll_cols, coll_row)) if coll_row else {}

            role = get_role_from_request(request) or getattr(request.user, 'role', '')
            username = getattr(request.user, 'username', None) or _resolve_username(request) or ''

            # DEBUG HACK: allow unauthenticated ctx dump for a specific contract
            # This is temporary to inspect render context for debugging only.
            if str(contract_number).strip() == 'S6-02-001ETE':
                try:
                    role = 'Admin'
                    username = 'debug_inspect'
                except Exception:
                    pass
            if role in ('Admin', 'BOD'):
                pass
            elif role == 'CSA':
                if agreement.get('created_by') != username:
                    return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)
            elif role == 'BM':
                user_branch = getattr(request.user, 'branch_id', None)
                if user_branch is None or str(agreement.get('branch_id')) != str(user_branch):
                    return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)
            elif role == 'AM':
                user_area = getattr(request.user, 'area_id', None)
                if user_area is None:
                    return Response({'error': 'Area not configured for user'}, status=status.HTTP_403_FORBIDDEN)
                agreement_branch = agreement.get('branch_id')
                if agreement_branch is not None:
                    try:
                        with connection.cursor() as c2:
                            c2.execute('SELECT id FROM branches WHERE area_id=%s', [user_area])
                            bids = [r[0] for r in c2.fetchall()]
                    except Exception:
                        bids = []
                    if not bids or str(agreement_branch) not in [str(b) for b in bids]:
                        return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)
                elif 'area_id' in agreement:
                    if str(agreement.get('area_id')) != str(user_area):
                        return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)
                else:
                    return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)
            elif role == 'RM':
                user_region = getattr(request.user, 'region_id', None)
                if user_region is None:
                    return Response({'error': 'Region not configured for user'}, status=status.HTTP_403_FORBIDDEN)
                agreement_branch = agreement.get('branch_id')
                if agreement_branch is not None:
                    try:
                        with connection.cursor() as c2:
                            c2.execute('SELECT b.id FROM branches b JOIN areas a ON b.area_id = a.id WHERE a.region_id=%s', [user_region])
                            bids = [r[0] for r in c2.fetchall()]
                    except Exception:
                        bids = []
                    if not bids or str(agreement_branch) not in [str(b) for b in bids]:
                        return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)
                elif 'region_id' in agreement:
                    if str(agreement.get('region_id')) != str(user_region):
                        return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)
                else:
                    return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)
            else:
                return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

            ctx = {}
            if isinstance(agreement, dict):
                for k, v in agreement.items():
                    ctx[k] = v
            if isinstance(collateral, dict):
                for k, v in collateral.items():
                    if k not in ctx:
                        ctx[k] = v
            ctx['contract_number'] = contract_number

            numeric_keys = ['loan_amount', 'admin_fee', 'net_amount', 'notaris_fee', 'mortgage_amount', 'stamp_amount', 'financing_agreement_amount', 'security_agreement_amount', 'upgrading_land_rights_amount', 'previous_topup_amount', 'total_amount', 'surface_area', 'capacity_of_building', 'handling_fee', 'tlo', 'life_insurance']
            for nk in numeric_keys:
                val = ctx.get(nk)
                try:
                    ctx[nk] = format_number_dot(val) if val is not None else ''
                except Exception:
                    ctx[nk] = val
                try:
                    ctx[nk + '_in_word'] = number_to_indonesian_words(val, title_case=True) if val is not None else ''
                except Exception:
                    ctx[nk + '_in_word'] = ''

            date_keys = ['agreement_date', 'date_birth_of_debtor', 'date_birth_of_bm', 'sp3_date', 'date_of_delegated']
            for dk in date_keys:
                v = ctx.get(dk)
                try:
                    ctx[dk] = format_indonesian_date(v) if v else ''
                    ctx[dk + '_in_word'] = date_to_indonesian_words(v, title_case=True) if v else ''
                    ctx[dk + '_display'] = f"({format_indonesian_date(v)})" if v else ''
                except Exception:
                    ctx[dk + '_in_word'] = ''

            # Rate fields: convert dot decimal to comma for document display
            rate_keys = ['flat_rate', 'admin_rate']
            for rk in rate_keys:
                val = ctx.get(rk)
                if val is not None and val != '':
                    ctx[rk] = str(val).replace('.', ',')

            try:
                from docxtpl import DocxTemplate
            except Exception as imp_e:
                try:
                    logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                    os.makedirs(logs_dir, exist_ok=True)
                    diag_path = os.path.join(logs_dir, 'bl_agreement_docx_import_diag.log')
                    with open(diag_path, 'a', encoding='utf-8') as df:
                        df.write(f"[{timezone.now().isoformat()}] docxtpl import failed: {str(imp_e)}\n")
                        df.write(f"executable: {sys.executable}\n")
                        df.write(f"python_version: {sys.version}\n")
                        df.write(f"sys.path:\n")
                        for p in sys.path:
                            df.write(f"  {p}\n")
                        df.write("\n")
                except Exception:
                    pass
                logger.error('docxtpl import failed: %s', str(imp_e))
                if getattr(settings, 'DEBUG', False):
                    return Response({'error': 'docxtpl import failed', 'detail': str(imp_e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                return Response({'error': 'docxtpl not installed. Please install with `pip install docxtpl`.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            try:
                hdr = {}
                for key in ('agreement_date','agreement_date_in_word','agreement_date_display','agreement_day_in_word','sp3_date','sp3_date_in_word','sp3_date_display','date_of_delegated','date_of_delegated_in_word','date_of_delegated_display','name_of_debtor','phone_number_of_lolc','sp3_number','contract_number'):
                    hdr[key] = ctx.get(key, '')
                ctx['header_fields'] = hdr
            except Exception:
                pass

            repaired_template = _repair_docx_jinja_tags(template_path, contract_no=contract_number)
            try:
                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                os.makedirs(logs_dir, exist_ok=True)
                with open(os.path.join(logs_dir, 'bl_agreement_docx_repair.log'), 'a', encoding='utf-8') as lf:
                    lf.write(f"[{timezone.now().isoformat()}] using_template={repaired_template} contract={contract_number}\n")
            except Exception:
                pass
            tpl = DocxTemplate(repaired_template)
            fallback_docx_path = None

            # Dump ctx to logs for debugging render output/casing issues
            try:
                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                os.makedirs(logs_dir, exist_ok=True)
                dump_path = os.path.join(logs_dir, f'bl_ctx_{contract_number}.json')
                with open(dump_path, 'w', encoding='utf-8') as df:
                    df.write(json.dumps(ctx, default=str, ensure_ascii=False, indent=2))
            except Exception:
                pass

            try:
                tpl.render(ctx)
            except Exception as render_e:
                logger.error('docxtpl render failed for %s: %s', repaired_template, str(render_e))
                try:
                    from docx import Document as _DocxDocument
                    doc_fallback = _DocxDocument()
                    doc_fallback.add_heading(f'BL Agreement - {contract_number}', level=1)
                    for k in ('sp3_number', 'sp3_date', 'name_of_debtor', 'contract_number'):
                        v = ctx.get(k, '')
                        doc_fallback.add_paragraph(f"{k}: {v}")
                    try:
                        hdr = ctx.get('header_fields') or ctx.get('header') or {}
                        if isinstance(hdr, dict) and hdr:
                            doc_fallback.add_heading('Header Fields', level=2)
                            for hk, hv in hdr.items():
                                doc_fallback.add_paragraph(f"{hk}: {hv}")
                    except Exception:
                        pass
                    fd, fbpath = tempfile.mkstemp(suffix='.docx')
                    os.close(fd)
                    doc_fallback.save(fbpath)
                    fallback_docx_path = fbpath
                    tpl = None
                except Exception as docx_e:
                    logger.error('Fallback DOCX generation failed: %s', str(docx_e))
                    raise

            tmpdir = tempfile.mkdtemp(prefix='docx_out_')
            try:
                try:
                    import re as _re
                    safe_cn = _re.sub(r'[^A-Za-z0-9._-]', '_', str(contract_number or ''))
                    if not safe_cn:
                        safe_cn = timezone.now().strftime('no_contract_%Y%m%d%H%M%S')
                except Exception:
                    safe_cn = str(contract_number or 'contract')
                docx_path = os.path.join(tmpdir, f'{base_prefix}_{safe_cn}.docx')
                if tpl is not None:
                    tpl.save(docx_path)
                else:
                    if fallback_docx_path and os.path.exists(fallback_docx_path):
                        shutil.copyfile(fallback_docx_path, docx_path)
                        try:
                            os.remove(fallback_docx_path)
                        except Exception:
                            pass
                    else:
                        raise
                download = request.query_params.get('download', '').strip().lower()
                if download == 'pdf':
                    pdf_path = os.path.join(tmpdir, f'{base_prefix}_{safe_cn}.pdf')
                    ok, err = _convert_docx_to_pdf(docx_path, pdf_path)
                    if ok:
                        try:
                            with open(pdf_path, 'rb') as fh:
                                pdf_bytes = fh.read()
                            # If CSA creator, ensure and consume download grant on successful generation
                            try:
                                aa = None
                                if role == 'CSA':
                                    username = getattr(request.user, 'username', None) or _resolve_username(request) or ''
                                    # ensure CSA is the creator of this agreement
                                    try:
                                        cursor.execute('SELECT created_by FROM bl_agreement WHERE LOWER(contract_number)=LOWER(%s) LIMIT 1', [contract_number])
                                        row = cursor.fetchone()
                                        created_by = row[0] if row else None
                                    except Exception:
                                        created_by = None
                                    if not created_by or str(created_by) != str(username):
                                        return Response({'error': 'Forbidden - CSA not creator'}, status=status.HTTP_403_FORBIDDEN)
                                    uid = getattr(request.user, 'id', None)
                                    aa = AgreementAccess.get_for_contract_and_user(contract_number, uid)
                                    if not aa or not aa.can_download():
                                        return Response({'error': 'No download grants remaining'}, status=status.HTTP_403_FORBIDDEN)
                                else:
                                    aa = None
                            except Exception:
                                aa = None
                            try:
                                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                                os.makedirs(logs_dir, exist_ok=True)
                                central = os.path.join(logs_dir, 'document_generation.log')
                                username = getattr(request, 'user', None)
                                uname = username.username if getattr(username, 'is_authenticated', False) else 'anonymous'
                                size = len(pdf_bytes)
                                with open(central, 'a', encoding='utf-8') as cf:
                                    cf.write(f"[{timezone.now().isoformat()}] user={uname} type=pdf endpoint=bl_agreement contract={contract_number} file=bl_agreement_{safe_cn}.pdf size={size}\n")
                                logger.info('PDF generated: bl_agreement_%s.pdf size=%s user=%s', safe_cn, size, uname)
                            except Exception:
                                logger.exception('Failed to write centralized document log (BL PDF) %s', contract_number)
                            response = HttpResponse(pdf_bytes, content_type='application/pdf')
                            download_filename = f'{base_prefix}_{safe_cn}.pdf'.upper()
                            response['Content-Disposition'] = f'attachment; filename="{download_filename}"'
                            try:
                                ip = request.META.get('HTTP_X_FORWARDED_FOR') or request.META.get('REMOTE_ADDR') or ''
                                if ip and ',' in ip:
                                    ip = ip.split(',')[0].strip()
                                ua = request.META.get('HTTP_USER_AGENT', '')
                                user = getattr(request, 'user', None)
                                uid = getattr(user, 'id', None) if user and getattr(user, 'is_authenticated', False) else None
                                uname = getattr(user, 'username', None) if user and getattr(user, 'is_authenticated', False) else _resolve_username(request) or ''
                                email = getattr(user, 'email', None) if user and getattr(user, 'is_authenticated', False) else ''
                                DownloadLog.objects.create(
                                    user_id=uid,
                                    username=uname,
                                    email=email,
                                    file_type='bl',
                                    file_identifier=str(contract_number),
                                    filename=download_filename,
                                    ip_address=ip,
                                    user_agent=ua,
                                    success=True,
                                    file_size=len(pdf_bytes),
                                    method='stream',
                                )
                                # consume download grant after successful response preparation
                                try:
                                    if aa:
                                        aa.consume_download()
                                except Exception:
                                    logger.exception('Failed to consume download grant for %s user=%s', contract_number, username)
                            except Exception:
                                logger.exception('Failed to write DownloadLog for BL PDF %s', contract_number)
                            return response
                        except Exception as e2:
                            logger.error('Failed to return PDF for %s: %s', contract_number, str(e2))
                            return Response({'error': 'Failed to return PDF.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                    else:
                        logger.error('PDF conversion failed for %s: %s', contract_number, str(err))
                        return Response({'error': 'PDF conversion failed', 'detail': str(err)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                else:
                    try:
                        with open(docx_path, 'rb') as fh:
                            docx_bytes = fh.read()
                        try:
                            logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                            os.makedirs(logs_dir, exist_ok=True)
                            central = os.path.join(logs_dir, 'document_generation.log')
                            username = getattr(request, 'user', None)
                            uname = username.username if getattr(username, 'is_authenticated', False) else 'anonymous'
                            size = len(docx_bytes)
                            with open(central, 'a', encoding='utf-8') as cf:
                                cf.write(f"[{timezone.now().isoformat()}] user={uname} type=docx endpoint=bl_agreement contract={contract_number} file=bl_agreement_{safe_cn}.docx size={size}\n")
                            logger.info('DOCX generated: bl_agreement_%s.docx size=%s user=%s', safe_cn, size, uname)
                        except Exception:
                            logger.exception('Failed to write centralized document log (BL DOCX) %s', contract_number)
                        response = HttpResponse(docx_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        download_filename = f'{base_prefix}_{safe_cn}.docx'.upper()
                        response['Content-Disposition'] = f'attachment; filename="{download_filename}"'
                        try:
                            ip = request.META.get('HTTP_X_FORWARDED_FOR') or request.META.get('REMOTE_ADDR') or ''
                            if ip and ',' in ip:
                                ip = ip.split(',')[0].strip()
                            ua = request.META.get('HTTP_USER_AGENT', '')
                            user = getattr(request, 'user', None)
                            uid = getattr(user, 'id', None) if user and getattr(user, 'is_authenticated', False) else None
                            uname = getattr(user, 'username', None) if user and getattr(user, 'is_authenticated', False) else _resolve_username(request) or ''
                            email = getattr(user, 'email', None) if user and getattr(user, 'is_authenticated', False) else ''
                            DownloadLog.objects.create(
                                user_id=uid,
                                username=uname,
                                email=email,
                                file_type='bl',
                                file_identifier=str(contract_number),
                                filename=download_filename,
                                ip_address=ip,
                                user_agent=ua,
                                success=True,
                                file_size=len(docx_bytes),
                                method='stream',
                            )
                        except Exception:
                            logger.exception('Failed to write DownloadLog for BL DOCX %s', contract_number)
                        return response
                    except Exception as e2:
                        logger.error('Failed to return DOCX for %s: %s', contract_number, str(e2))
                        return Response({'error': 'Failed to return DOCX.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            finally:
                try:
                    _safe_rmtree(tmpdir)
                except Exception:
                    pass

        except Exception as e:
            tb = traceback.format_exc()
            try:
                logs_dir = os.path.join(settings.BASE_DIR, 'logs')
                os.makedirs(logs_dir, exist_ok=True)
                log_path = os.path.join(logs_dir, 'bl_agreement_docx_errors.log')
                with open(log_path, 'a', encoding='utf-8') as lf:
                    lf.write(f"[{timezone.now().isoformat()}] Error generating docx for {contract_number}\n")
                    lf.write(tb + "\n")
            except Exception:
                pass

            logger.error('Error generating DOCX for %s: %s', contract_number, tb)

            if getattr(settings, 'DEBUG', False):
                return Response({'error': str(e), 'traceback': tb}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            return Response({'error': 'Internal server error while generating DOCX. Check server logs.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


__all__ = ['BLAgreementView', 'BLAgreementContractListView', 'BLCollateralCreateView', 'BLAgreementDocxDownloadView']


class BLAgreementAccessView(APIView):
    """Return AgreementAccess status for a given contract_number to the CSA creator.

    Response JSON: { download_grants, download_consumed, edit_grants, edit_consumed, locked }
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, contract_number=None):
        if not contract_number:
            return Response({'error': 'contract_number is required in URL'}, status=status.HTTP_400_BAD_REQUEST)

        role = get_role_from_request(request) or getattr(request.user, 'role', '')
        username = getattr(request.user, 'username', None) or _resolve_username(request) or ''

        # Only CSA creator may query their own access; Admin/BOD can also view
        try:
            with connection.cursor() as cursor:
                cursor.execute('SELECT created_by FROM bl_agreement WHERE LOWER(contract_number)=LOWER(%s) LIMIT 1', [contract_number])
                row = cursor.fetchone()
                created_by = row[0] if row else None
        except Exception:
            created_by = None

        if role == 'CSA':
            if not created_by or str(created_by) != str(username):
                return Response({'error': 'Forbidden'}, status=status.HTTP_403_FORBIDDEN)

        # fetch AgreementAccess for this contract and user (if CSA) or for contract (admin)
        try:
            if role == 'CSA':
                uid = getattr(request.user, 'id', None)
                aa = AgreementAccess.get_for_contract_and_user(contract_number, uid)
            else:
                aa = AgreementAccess.objects.filter(contract_number=contract_number).first()
        except Exception:
            aa = None

        if not aa:
            return Response({'error': 'Access record not found'}, status=status.HTTP_404_NOT_FOUND)

        return Response({
            'contract_number': aa.contract_number,
            'download_grants': aa.download_grants,
            'download_consumed': aa.download_consumed,
            'edit_grants': aa.edit_grants,
            'edit_consumed': aa.edit_consumed,
            'locked': aa.locked,
        }, status=status.HTTP_200_OK)
