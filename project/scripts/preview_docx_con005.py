from docx import Document
import os

f = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'tmp_docx', 'bl_agreement_CON005.docx'))
print('FILE:' + f)
if not os.path.exists(f):
    print('MISSING')
    raise SystemExit(1)

try:
    doc = Document(f)
except Exception as e:
    print('ERROR_LOADING:' + str(e))
    raise

paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
print('PARA_COUNT:' + str(len(paras)))
for i, t in enumerate(paras[:200], 1):
    s = t.replace('\n', ' ')[:400]
    print(f'P{i}:' + s)

# Print table cells too
cells = []
for table in doc.tables:
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                cells.append(((r,c), txt))

print('TABLE_CELLS:' + str(len(cells)))
for i, ((r,c), txt) in enumerate(cells[:200], 1):
    print(f'T{i} (r{r}c{c}):' + txt.replace('\n',' '))
